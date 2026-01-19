"""
책쓰기 코칭 자동화 웹앱 v2.1
============================
완전 초보자도 6만자 책을 완성할 수 있도록 설계
- 목차 순서대로 1장씩 안내
- 실시간 진행률 표시
- 데이터 저장/불러오기
- 강화된 에러 핸들링 및 자동 저장 기능
"""
import streamlit as st
import json
import re
from datetime import datetime
from prompts.templates import WRITING_TONES
from utils.claude_client import (
    generate_titles,
    generate_toc,
    generate_draft,
    get_feedback,
    refine_text,
    add_storytelling,
    chat_with_coach,
    edit_draft_with_instruction,
    generate_proposal,
    generate_landing_page,
    analyze_youtube_transcript,
    generate_titles_from_transcript,
    generate_toc_from_transcript,
    generate_draft_from_transcript,
)
from utils.voice_handler import (
    render_voice_mode_ui,
    render_transcription_editor,
    clear_voice_session,
)
from utils.youtube_handler import (
    extract_video_id,
    get_video_info,
    get_transcript,
    validate_youtube_url,
    process_multiple_videos,
    merge_transcripts_for_book,
)
from utils.contact_handler import render_contact_section, get_pending_messages_count
from utils.help_chatbot import (
    render_enhanced_chatbot,
    render_help_sidebar_button,
    init_help_chatbot_state,
    render_floating_chatbot_button,
    FAQ_QUESTIONS,
    get_step_context,
)
from utils.achievement_system import (
    init_achievement_state,
    check_and_award_badges,
    check_milestone,
    update_streak,
    update_daily_progress,
    get_random_motivation,
    get_motivation_by_progress,
    render_progress_header,
    render_badge_popup,
    render_milestone_popup,
    render_badges_display,
    render_daily_goal_section,
    on_chapter_complete,
    get_achievement_css,
    get_completed_chapters,
    get_total_chars,
    BADGES,
    MILESTONE_MESSAGES,
)
from utils.achievement_css import ACHIEVEMENT_CSS
# 에러 핸들링 및 자동 저장 모듈
from utils.error_handler import (
    safe_session_init,
    perform_autosave_if_needed,
    check_autosave_reminder,
    validate_session_state,
    show_error_with_recovery,
    get_autosave_files,
    restore_from_autosave,
    render_autosave_indicator,
)
from utils.autosave_handler import (
    init_autosave_state,
    render_autosave_status,
    render_save_buttons,
    render_backup_list,
    render_recovery_prompt,
    trigger_important_save,
    save_progress,
    get_time_since_last_save,
)
try:
    from utils.mode_transition import (
        init_mode_transition_state,
        handle_api_error,
        determine_next_step,
    )
except ImportError:
    # 모드 전환 모듈이 없을 경우 빈 함수로 대체
    def init_mode_transition_state():
        pass
    def handle_api_error(error, context=""):
        return f"오류가 발생했습니다: {str(error)}"
    def determine_next_step():
        return 1

# 페이지 설정
st.set_page_config(
    page_title="작가의집 - AI 책쓰기 코칭",
    page_icon="🏠",
    layout="wide",
    initial_sidebar_state="expanded",
)

# CSS 스타일 (초등학생 친화적 + 파스텔 테마 + 글씨 크기 조절)
def get_font_size_css():
    """글씨 크기에 따른 CSS 변수 반환 - 컴팩트 버전"""
    size = st.session_state.get("font_size", "normal")
    sizes = {
        "small": {"base": "14px", "header": "1.6rem", "step": "1.2rem", "progress": "1.4rem"},
        "normal": {"base": "15px", "header": "1.8rem", "step": "1.4rem", "progress": "1.6rem"},
        "large": {"base": "18px", "header": "2.2rem", "step": "1.7rem", "progress": "2rem"},
    }
    return sizes.get(size, sizes["normal"])

# 기본 글씨 크기 설정 (컴팩트 레이아웃 기본값: 보통)
if "font_size" not in st.session_state:
    st.session_state.font_size = "normal"

font_sizes = get_font_size_css()

st.markdown(f"""
<style>
    /* ============================================ */
    /* 작가의집 책쓰기 코칭 - 전문 테마 CSS         */
    /* 초등학생도 쉽게 + 전문적인 디자인            */
    /* ============================================ */

    /* Google Fonts 로드 */
    @import url('https://fonts.googleapis.com/css2?family=Noto+Sans+KR:wght@300;400;500;600;700;800&display=swap');
    /* ============================================ */

    /* CSS 변수 정의 */
    :root {{
        --font-base: {font_sizes['base']};
        --font-header: {font_sizes['header']};
        --font-step: {font_sizes['step']};
        --font-progress: {font_sizes['progress']};

        /* 파스텔 색상 팔레트 (WCAG AA 대비율 준수) */
        /* 전문 코칭 메인 색상 (네이비 + 골드) */
        --primary-navy: #1e3a5f;
        --primary-navy-light: #2d5a87;
        --primary-navy-dark: #0f2540;
        --accent-gold: #c9a962;
        --accent-gold-light: #e0c989;
        --accent-gold-dark: #a68b4b;

        /* 보조 색상 팔레트 */
        --pastel-pink: #fff5f7;
        --pastel-blue: #ebf8ff;
        --pastel-green: #f0fff4;
        --pastel-yellow: #fffff0;
        --pastel-purple: #faf5ff;
        --pastel-orange: #fffaf0;

        /* 접근성 향상된 텍스트 색상 (WCAG AA 4.5:1 이상) */
        --text-primary: #1a202c;
        --text-secondary: #4a5568;
        --text-on-green: #22543d;
        --text-on-blue: #2a4365;
        --text-on-orange: #c05621;
        --text-muted: #718096;

        /* 무지개 색상 (진행 표시용) */
        --rainbow-1: #1e3a5f;
        --rainbow-2: #2d5a87;
        --rainbow-3: #4299e1;
        --rainbow-4: #c9a962;
        --rainbow-5: #e0c989;

        /* 터치 타겟 최소 크기 (컴팩트) */
        --touch-target-min: 44px;
        --touch-target-comfortable: 48px;

        /* 포커스 스타일 */
        --focus-color: #2d5a87;

        /* 그림자 시스템 */
        --shadow-sm: 0 1px 2px 0 rgba(0, 0, 0, 0.05);
        --shadow-md: 0 4px 6px -1px rgba(0, 0, 0, 0.1), 0 2px 4px -1px rgba(0, 0, 0, 0.06);
        --shadow-lg: 0 10px 15px -3px rgba(0, 0, 0, 0.1), 0 4px 6px -2px rgba(0, 0, 0, 0.05);
        --shadow-xl: 0 20px 25px -5px rgba(0, 0, 0, 0.1), 0 10px 10px -5px rgba(0, 0, 0, 0.04);
        --shadow-gold: 0 8px 30px rgba(201, 169, 98, 0.25);
        --shadow-navy: 0 8px 30px rgba(30, 58, 95, 0.2);
        --focus-width: 3px;

        /* 모바일 safe-area (노치/홈 인디케이터 대응) */
        --safe-area-inset-top: env(safe-area-inset-top, 0px);
        --safe-area-inset-bottom: env(safe-area-inset-bottom, 0px);
        --safe-area-inset-left: env(safe-area-inset-left, 0px);
        --safe-area-inset-right: env(safe-area-inset-right, 0px);
    }}

    /* 기본 폰트 설정 - 한글 및 특수문자/이모지 지원 강화 (컴팩트) */
    html, body, [class*="st-"] {{
        font-family: 'Noto Sans KR', 'Apple SD Gothic Neo', 'Malgun Gothic', '맑은 고딕',
                     -apple-system, BlinkMacSystemFont, 'Segoe UI', 'Segoe UI Emoji',
                     'Segoe UI Symbol', 'Noto Color Emoji', sans-serif !important;
        font-size: var(--font-base) !important;
        line-height: 1.5 !important;
        color: var(--text-primary);
        -webkit-font-smoothing: antialiased;
        -moz-osx-font-smoothing: grayscale;
        text-rendering: optimizeLegibility;
    }}

    /* 이모지 전용 폰트 설정 */
    .emoji, [class*="emoji"] {{
        font-family: 'Apple Color Emoji', 'Segoe UI Emoji', 'Segoe UI Symbol',
                     'Noto Color Emoji', 'Android Emoji', 'EmojiSymbols', sans-serif !important;
    }}

    /* 스크롤 동작 최적화 */
    html {{
        scroll-behavior: smooth;
        -webkit-overflow-scrolling: touch;
    }}

    /* 텍스트 선택 색상 (접근성) */
    ::selection {{
        background-color: rgba(201, 169, 98, 0.3);
        color: inherit;
    }}

    /* 메인 헤더 - 전문적 스타일 (컴팩트) */
    .main-header {{
        font-size: var(--font-header);
        font-weight: 700;
        color: white;
        margin-bottom: 0.8rem;
        letter-spacing: -0.5px;
        background: linear-gradient(135deg, var(--primary-navy) 0%, var(--primary-navy-light) 100%);
        padding: 0.8rem 1.2rem;
        border-radius: 12px;
        text-align: center;
        box-shadow: var(--shadow-navy);
        border-bottom: 3px solid var(--accent-gold);
    }}

    /* 단계 헤더 - 전문적 스타일 (컴팩트) */
    .step-header {{
        font-size: var(--font-step);
        font-weight: 700;
        color: var(--primary-navy);
        margin: 0;
        margin-top: 0.8rem;
        margin-bottom: 0.6rem;
        padding: 0.6rem 1rem;
        background: linear-gradient(90deg, var(--pastel-blue) 0%, rgba(235, 248, 255, 0.3) 100%);
        border-radius: 10px;
        border-left: 4px solid var(--accent-gold);
        box-shadow: var(--shadow-sm);
    }}
    h2.step-header {{
        margin-top: 0.8rem;
        margin-bottom: 0.6rem;
    }}

    /* ============================================ */
    /* 접근성: 키보드 포커스 강조 (WCAG 2.4.7)      */
    /* ============================================ */
    *:focus {{
        outline: none;
    }}
    *:focus-visible {{
        outline: var(--focus-width) solid var(--focus-color) !important;
        outline-offset: 2px !important;
        box-shadow: 0 0 0 6px rgba(25, 118, 210, 0.25) !important;
    }}
    button:focus-visible, a:focus-visible, input:focus-visible, textarea:focus-visible, select:focus-visible {{
        outline: var(--focus-width) solid var(--focus-color) !important;
        outline-offset: 2px !important;
        box-shadow: 0 0 0 6px rgba(25, 118, 210, 0.25) !important;
    }}
    .stButton > button:focus-visible {{
        box-shadow: 0 0 0 4px var(--focus-color), 0 0 0 8px rgba(25, 118, 210, 0.25) !important;
        outline: none !important;
    }}

    /* 스킵 네비게이션 (스크린 리더 및 키보드 사용자용) */
    .skip-link {{
        position: absolute;
        top: -100px;
        left: 0;
        background: var(--focus-color);
        color: white;
        padding: 12px 24px;
        z-index: 10000;
        font-weight: bold;
        font-size: 1rem;
        text-decoration: none;
        border-radius: 0 0 8px 0;
    }}
    .skip-link:focus {{
        top: 0;
    }}

    /* 스크린 리더 전용 텍스트 */
    .sr-only {{
        position: absolute;
        width: 1px;
        height: 1px;
        padding: 0;
        margin: -1px;
        overflow: hidden;
        clip: rect(0, 0, 0, 0);
        white-space: nowrap;
        border: 0;
    }}

    /* 라이브 리전 (실시간 업데이트 알림) */
    .live-region {{
        position: absolute;
        left: -10000px;
        width: 1px;
        height: 1px;
        overflow: hidden;
    }}

    /* ============================================ */
    /* 버튼 스타일 (터치 친화적 44px+ 타겟, 컴팩트)  */
    /* ============================================ */
    .stButton > button {{
        transition: all 0.2s ease-in-out;
        min-height: 44px !important;
        min-width: 44px !important;
        font-size: 0.95rem !important;
        font-weight: 600 !important;
        border-radius: 10px !important;
        padding: 0.5rem 1rem !important;
        border: 2px solid transparent !important;
        cursor: pointer;
        touch-action: manipulation; /* 더블탭 줌 방지 */
        -webkit-tap-highlight-color: transparent; /* 모바일 탭 하이라이트 제거 */
        user-select: none;
    }}
    .stButton > button:hover {{
        transform: translateY(-2px);
        box-shadow: 0 4px 16px rgba(0,0,0,0.12);
    }}
    .stButton > button:active {{
        transform: translateY(1px);
        box-shadow: 0 2px 8px rgba(0,0,0,0.1);
    }}
    /* 버튼 비활성화 상태 */
    .stButton > button:disabled {{
        opacity: 0.5;
        cursor: not-allowed;
        transform: none !important;
    }}

    /* 프라이머리 버튼 - 네이비 테마 */
    .stButton > button[kind="primary"] {{
        background: linear-gradient(135deg, var(--primary-navy) 0%, var(--primary-navy-light) 100%) !important;
        color: white !important;
        font-weight: 700 !important;
        border-bottom: 3px solid var(--accent-gold) !important;
        box-shadow: var(--shadow-md) !important;
    }}
    .stButton > button[kind="primary"]:hover {{
        background: linear-gradient(135deg, var(--primary-navy-light) 0%, var(--primary-navy) 100%) !important;
        box-shadow: var(--shadow-lg) !important;
        transform: translateY(-2px);
    }}

    /* 세컨더리 버튼 - 골드 테마 */
    .stButton > button[kind="secondary"] {{
        background: white !important;
        color: var(--primary-navy) !important;
        border: 2px solid var(--accent-gold) !important;
        font-weight: 600 !important;
    }}
    .stButton > button[kind="secondary"]:hover {{
        background: var(--pastel-yellow) !important;
        border-color: var(--accent-gold-dark) !important;
    }}

    /* 터치 피드백 (모바일) */
    @media (hover: none) and (pointer: coarse) {{
        .stButton > button:hover {{
            transform: none;
            box-shadow: none;
        }}
        .stButton > button:active {{
            transform: scale(0.98);
            opacity: 0.9;
        }}
    }}

    /* 로딩 상태 */
    .loading-state {{
        opacity: 0.6;
        pointer-events: none;
    }}

    /* ============================================ */
    /* UX 개선: 향상된 로딩 애니메이션 (컴팩트)     */
    /* ============================================ */
    .ux-loading-container {{
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        border-radius: 15px;
        padding: 1rem;
        margin: 0.8rem 0;
        text-align: center;
        color: white;
        box-shadow: 0 10px 30px rgba(102, 126, 234, 0.3);
        animation: ux-loading-pulse 2s ease-in-out infinite;
    }}
    @keyframes ux-loading-pulse {{
        0%, 100% {{ transform: scale(1); box-shadow: 0 10px 30px rgba(102, 126, 234, 0.3); }}
        50% {{ transform: scale(1.02); box-shadow: 0 15px 40px rgba(102, 126, 234, 0.5); }}
    }}
    .ux-loading-spinner {{
        display: inline-block;
        width: 40px;
        height: 40px;
        border: 4px solid rgba(255,255,255,0.3);
        border-top-color: #FFD700;
        border-radius: 50%;
        animation: ux-spin 1s linear infinite;
        margin-bottom: 0.5rem;
    }}
    @keyframes ux-spin {{
        to {{ transform: rotate(360deg); }}
    }}
    .ux-loading-emoji {{
        font-size: 2rem;
        animation: ux-bounce-emoji 1s ease-in-out infinite;
        display: inline-block;
    }}
    @keyframes ux-bounce-emoji {{
        0%, 100% {{ transform: translateY(0); }}
        50% {{ transform: translateY(-15px); }}
    }}
    .ux-loading-text {{
        font-size: 1.1rem;
        font-weight: bold;
        margin: 0.5rem 0;
        text-shadow: 1px 1px 2px rgba(0,0,0,0.2);
    }}
    .ux-loading-subtext {{
        font-size: 0.9rem;
        opacity: 0.9;
        margin: 0.3rem 0;
    }}
    .ux-loading-progress {{
        background: rgba(255,255,255,0.3);
        border-radius: 10px;
        height: 8px;
        margin: 0.5rem auto;
        max-width: 250px;
        overflow: hidden;
    }}
    .ux-loading-progress-bar {{
        background: linear-gradient(90deg, #FFD700, #FFA500, #FF6B6B);
        height: 100%;
        border-radius: 15px;
        transition: width 0.5s ease;
        animation: ux-progress-glow 1.5s ease-in-out infinite;
    }}
    @keyframes ux-progress-glow {{
        0%, 100% {{ box-shadow: 0 0 10px rgba(255, 215, 0, 0.5); }}
        50% {{ box-shadow: 0 0 20px rgba(255, 215, 0, 0.8); }}
    }}
    .ux-loading-percent {{
        font-size: 1.5rem;
        font-weight: bold;
        color: #FFD700;
        text-shadow: 2px 2px 4px rgba(0,0,0,0.3);
    }}
    .ux-loading-tips {{
        background: rgba(255,255,255,0.15);
        border-radius: 8px;
        padding: 0.5rem 0.8rem;
        margin-top: 0.5rem;
        font-size: 0.85rem;
    }}

    /* ============================================ */
    /* UX 개선: 토스트 알림 스타일                  */
    /* ============================================ */
    .ux-toast {{
        position: fixed;
        top: 20px;
        right: 20px;
        z-index: 10000;
        padding: 1rem 1.5rem;
        border-radius: 15px;
        box-shadow: 0 8px 25px rgba(0,0,0,0.2);
        animation: ux-toast-slide-in 0.5s ease forwards;
        max-width: 350px;
    }}
    @keyframes ux-toast-slide-in {{
        from {{ transform: translateX(100%); opacity: 0; }}
        to {{ transform: translateX(0); opacity: 1; }}
    }}
    .ux-toast-success {{
        background: linear-gradient(135deg, #4CAF50, #2E7D32);
        color: white;
    }}
    .ux-toast-error {{
        background: linear-gradient(135deg, #f44336, #c62828);
        color: white;
    }}
    .ux-toast-info {{
        background: linear-gradient(135deg, #2196F3, #1565C0);
        color: white;
    }}
    .ux-toast-warning {{
        background: linear-gradient(135deg, #FF9800, #E65100);
        color: white;
    }}
    .ux-toast-icon {{
        font-size: 1.5rem;
        margin-right: 0.5rem;
    }}
    .ux-toast-message {{
        font-size: 1.1rem;
        font-weight: 600;
    }}

    /* ============================================ */
    /* UX 개선: 화려한 축하 효과 (컴팩트)           */
    /* ============================================ */
    .ux-celebration {{
        background: linear-gradient(135deg, #FFD700 0%, #FFA500 50%, #FF6B6B 100%);
        border-radius: 15px;
        padding: 1.2rem;
        margin: 0.8rem 0;
        text-align: center;
        color: white;
        box-shadow: 0 15px 40px rgba(255, 107, 107, 0.4);
        animation: ux-celebration-glow 1.5s ease-in-out infinite;
        position: relative;
        overflow: hidden;
    }}
    @keyframes ux-celebration-glow {{
        0%, 100% {{ box-shadow: 0 15px 40px rgba(255, 107, 107, 0.4); }}
        50% {{ box-shadow: 0 20px 60px rgba(255, 215, 0, 0.6); }}
    }}
    .ux-celebration::before {{
        content: '';
        position: absolute;
        top: -50%;
        left: -50%;
        width: 200%;
        height: 200%;
        background: repeating-linear-gradient(
            45deg,
            transparent,
            transparent 10px,
            rgba(255,255,255,0.1) 10px,
            rgba(255,255,255,0.1) 20px
        );
        animation: ux-confetti-move 3s linear infinite;
    }}
    @keyframes ux-confetti-move {{
        from {{ transform: translate(0, 0); }}
        to {{ transform: translate(50px, 50px); }}
    }}
    .ux-celebration-content {{
        position: relative;
        z-index: 1;
    }}
    .ux-celebration-emoji {{
        font-size: 2.5rem;
        animation: ux-celebration-bounce 0.5s ease-in-out infinite alternate;
        display: inline-block;
        margin: 0 0.2rem;
    }}
    .ux-celebration-emoji:nth-child(2) {{ animation-delay: 0.1s; }}
    .ux-celebration-emoji:nth-child(3) {{ animation-delay: 0.2s; }}
    @keyframes ux-celebration-bounce {{
        from {{ transform: translateY(0) scale(1); }}
        to {{ transform: translateY(-10px) scale(1.1); }}
    }}
    .ux-celebration-title {{
        font-size: 1.4rem;
        font-weight: bold;
        margin: 0.5rem 0;
        text-shadow: 2px 2px 4px rgba(0,0,0,0.2);
    }}
    .ux-celebration-message {{
        font-size: 1rem;
        margin: 0.3rem 0;
    }}

    /* ============================================ */
    /* UX 개선: 네비게이션 현재 위치 표시 (컴팩트)  */
    /* ============================================ */
    .ux-breadcrumb {{
        background: linear-gradient(90deg, var(--pastel-blue), var(--pastel-purple));
        padding: 0.5rem 0.8rem;
        border-radius: 10px;
        margin-bottom: 0.8rem;
        display: flex;
        align-items: center;
        flex-wrap: wrap;
        gap: 0.3rem;
    }}
    .ux-breadcrumb-item {{
        display: inline-flex;
        align-items: center;
        font-size: 0.85rem;
        color: #424242;
        padding: 0.1rem 0;
    }}
    .ux-breadcrumb-item.completed {{
        color: #2E7D32;
    }}
    .ux-breadcrumb-item.active {{
        font-weight: bold;
        color: #1976D2;
        background: white;
        padding: 0.3rem 0.8rem;
        border-radius: 20px;
        box-shadow: 0 2px 8px rgba(0,0,0,0.1);
    }}
    .ux-breadcrumb-separator {{
        color: #9E9E9E;
        margin: 0 0.2rem;
        font-size: 0.85rem;
    }}

    /* ============================================ */
    /* UX 개선: 실시간 글자 수 카운터               */
    /* ============================================ */
    .ux-char-counter {{
        display: flex;
        justify-content: space-between;
        align-items: center;
        padding: 0.5rem 1rem;
        background: #F5F5F5;
        border-radius: 10px;
        margin-top: 0.5rem;
        font-size: 0.95rem;
    }}
    .ux-char-count {{
        font-weight: bold;
        color: #1976D2;
    }}
    .ux-char-count.warning {{
        color: #FF9800;
    }}
    .ux-char-count.danger {{
        color: #f44336;
    }}
    .ux-char-count.success {{
        color: #4CAF50;
    }}
    .ux-char-bar {{
        flex: 1;
        height: 8px;
        background: #E0E0E0;
        border-radius: 4px;
        margin: 0 1rem;
        overflow: hidden;
    }}
    .ux-char-bar-fill {{
        height: 100%;
        border-radius: 4px;
        transition: width 0.3s ease, background 0.3s ease;
    }}

    /* ============================================ */
    /* UX 개선: 단계별 팁 박스 (컴팩트)             */
    /* ============================================ */
    .ux-step-tip {{
        background: linear-gradient(135deg, #E8F5E9, #C8E6C9);
        border: 2px solid #81C784;
        border-radius: 10px;
        padding: 0.8rem 1rem 0.6rem 1rem;
        margin: 0.5rem 0;
        position: relative;
    }}
    .ux-step-tip::before {{
        content: '이렇게 해보세요!';
        position: absolute;
        top: -10px;
        left: 10px;
        background: #4CAF50;
        color: white;
        padding: 0.2rem 0.6rem;
        border-radius: 15px;
        font-size: 0.75rem;
        font-weight: bold;
    }}
    .ux-step-tip-content {{
        margin-top: 0.2rem;
        font-size: 0.9rem;
        line-height: 1.5;
        color: #2E7D32;
    }}
    .ux-step-tip ul {{
        margin: 0.5rem 0;
        padding-left: 1.5rem;
    }}
    .ux-step-tip li {{
        margin: 0.3rem 0;
    }}

    /* ============================================ */
    /* UX 개선: 눈에 띄는 FAQ 버튼                  */
    /* ============================================ */
    .ux-faq-floating {{
        position: fixed;
        bottom: 100px;
        right: 30px;
        z-index: 999;
        background: linear-gradient(135deg, #FF6B6B, #FF8E53);
        color: white;
        border: none;
        border-radius: 50px;
        padding: 12px 20px;
        font-size: 1rem;
        font-weight: bold;
        cursor: pointer;
        box-shadow: 0 6px 20px rgba(255, 107, 107, 0.4);
        transition: all 0.3s ease;
        animation: ux-faq-pulse 2s infinite;
    }}
    @keyframes ux-faq-pulse {{
        0%, 100% {{ box-shadow: 0 6px 20px rgba(255, 107, 107, 0.4); }}
        50% {{ box-shadow: 0 8px 30px rgba(255, 107, 107, 0.6); }}
    }}
    .ux-faq-floating:hover {{
        transform: translateY(-3px) scale(1.05);
    }}

    /* ============================================ */
    /* UX 개선: 툴팁 스타일                         */
    /* ============================================ */
    .ux-tooltip {{
        position: relative;
        display: inline-block;
        cursor: help;
    }}
    .ux-tooltip-icon {{
        display: inline-flex;
        align-items: center;
        justify-content: center;
        width: 20px;
        height: 20px;
        background: #64B5F6;
        color: white;
        border-radius: 50%;
        font-size: 0.8rem;
        font-weight: bold;
        margin-left: 5px;
    }}
    .ux-tooltip-text {{
        visibility: hidden;
        position: absolute;
        bottom: 125%;
        left: 50%;
        transform: translateX(-50%);
        background: #333;
        color: white;
        padding: 0.8rem 1rem;
        border-radius: 10px;
        font-size: 0.9rem;
        width: 250px;
        text-align: center;
        z-index: 100;
        opacity: 0;
        transition: opacity 0.3s, visibility 0.3s;
    }}
    .ux-tooltip:hover .ux-tooltip-text {{
        visibility: visible;
        opacity: 1;
    }}
    .ux-tooltip-text::after {{
        content: '';
        position: absolute;
        top: 100%;
        left: 50%;
        transform: translateX(-50%);
        border: 8px solid transparent;
        border-top-color: #333;
    }}

    /* ============================================ */
    /* UX 개선: 단계 건너뛰기 방지 경고              */
    /* ============================================ */
    .ux-step-locked {{
        background: linear-gradient(135deg, #FFCDD2, #EF9A9A);
        border: 3px dashed #E57373;
        border-radius: 15px;
        padding: 1.5rem;
        margin: 1rem 0;
        text-align: center;
    }}
    .ux-step-locked-icon {{
        font-size: 3rem;
        margin-bottom: 0.5rem;
    }}
    .ux-step-locked-text {{
        font-size: 1.1rem;
        color: #C62828;
        font-weight: 600;
    }}
    .ux-step-locked-hint {{
        font-size: 0.95rem;
        color: #D32F2F;
        margin-top: 0.5rem;
    }}

    /* ============================================ */
    /* UX 개선: 자동 스크롤 대상 강조               */
    /* ============================================ */
    .ux-scroll-target {{
        animation: ux-highlight-flash 1s ease-in-out;
        scroll-margin-top: 100px;
    }}
    @keyframes ux-highlight-flash {{
        0%, 100% {{ background-color: transparent; }}
        50% {{ background-color: rgba(255, 235, 59, 0.5); }}
    }}

    /* ============================================ */
    /* UX 개선: 입력 검증 상태 표시                 */
    /* ============================================ */
    .ux-input-valid {{
        border-color: #4CAF50 !important;
        box-shadow: 0 0 0 3px rgba(76, 175, 80, 0.2) !important;
    }}
    .ux-input-invalid {{
        border-color: #f44336 !important;
        box-shadow: 0 0 0 3px rgba(244, 67, 54, 0.2) !important;
    }}
    .ux-validation-message {{
        font-size: 0.9rem;
        margin-top: 0.3rem;
        display: flex;
        align-items: center;
        gap: 0.3rem;
    }}
    .ux-validation-message.valid {{
        color: #4CAF50;
    }}
    .ux-validation-message.invalid {{
        color: #f44336;
    }}

    /* 진행률 박스 - 전문적 네이비/골드 테마 (컴팩트) */
    .progress-box {{
        background: linear-gradient(135deg, var(--primary-navy) 0%, var(--primary-navy-light) 50%, var(--accent-gold) 100%);
        color: white;
        padding: 0.8rem 1.2rem;
        border-radius: 12px;
        margin: 0.6rem 0;
        text-align: center;
        box-shadow: var(--shadow-md);
        border: 1px solid rgba(201, 169, 98, 0.3);
        position: relative;
        overflow: hidden;
    }}
    .progress-box::before {{
        content: '';
        position: absolute;
        top: 0;
        left: 0;
        right: 0;
        height: 3px;
        background: linear-gradient(90deg, var(--accent-gold), var(--accent-gold-light), var(--accent-gold));
    }}
    .progress-text {{
        font-size: var(--font-progress);
        font-weight: 700;
        text-shadow: 1px 1px 2px rgba(0,0,0,0.3);
        letter-spacing: 0.3px;
    }}

    /* 별 진행률 표시 (컴팩트) */
    .star-progress {{
        font-size: 1.3rem;
        letter-spacing: 3px;
        margin: 0.3rem 0;
    }}

    /* 현재 섹션 박스 (컴팩트) */
    .current-section-box {{
        background: var(--pastel-yellow);
        border-left: 5px solid #FFA000;
        padding: 0.6rem 1rem;
        margin: 0.6rem 0;
        border-radius: 0 12px 12px 0;
        font-size: 1rem;
    }}

    /* 완료 섹션 (컴팩트) */
    .completed-section {{
        background: var(--pastel-green);
        padding: 0.5rem 0.8rem;
        border-radius: 10px;
        margin: 0.3rem 0;
        color: #1B5E20;
        font-weight: 600;
    }}

    /* 대기 섹션 (컴팩트) */
    .pending-section {{
        background: #F5F5F5;
        padding: 0.5rem 0.8rem;
        border-radius: 10px;
        margin: 0.3rem 0;
        color: #757575;
    }}

    .big-button {{
        font-size: 1.2rem !important;
        padding: 0.8rem 1.5rem !important;
    }}

    /* 도움말 박스 (컴팩트) */
    .help-box {{
        background: var(--pastel-blue);
        border: 2px solid #64B5F6;
        padding: 0.8rem 1rem;
        border-radius: 12px;
        margin: 0.6rem 0;
        font-size: 0.95rem;
        line-height: 1.5;
    }}

    /* 힌트 박스 (이렇게 써보세요!) (컴팩트) */
    .hint-box {{
        background: var(--pastel-purple);
        border: 2px dashed #9C27B0;
        padding: 0.6rem 0.8rem;
        border-radius: 10px;
        margin: 0.5rem 0;
        font-size: 0.9rem;
    }}
    .hint-box::before {{
        content: "이렇게 써보세요!";
        display: block;
        font-weight: bold;
        color: #7B1FA2;
        margin-bottom: 0.3rem;
        font-size: 0.85rem;
    }}

    /* 경고 박스 (컴팩트) */
    .warning-box {{
        background: var(--pastel-orange);
        border: 2px solid #FF9800;
        padding: 0.8rem 1rem;
        border-radius: 12px;
        margin: 0.6rem 0;
    }}

    /* 성공 박스 (컴팩트) */
    .success-box {{
        background: var(--pastel-green);
        border: 2px solid #4CAF50;
        padding: 0.8rem 1rem;
        border-radius: 12px;
        margin: 0.6rem 0;
    }}

    /* 저장 상태 표시 */
    .save-status {{
        background: var(--pastel-green);
        color: #1B5E20;
        padding: 0.5rem 1rem;
        border-radius: 25px;
        font-size: 0.9rem;
        font-weight: bold;
        display: inline-block;
        margin: 0.5rem 0;
    }}

    /* 축하 애니메이션 */
    @keyframes celebrate {{
        0%, 100% {{ transform: scale(1); }}
        50% {{ transform: scale(1.1); }}
    }}
    .celebrate {{
        animation: celebrate 0.5s ease-in-out 3;
    }}

    @keyframes bounce {{
        0%, 100% {{ transform: translateY(0); }}
        50% {{ transform: translateY(-10px); }}
    }}
    .bounce {{
        animation: bounce 0.5s ease-in-out infinite;
    }}

    /* 뱃지/스티커 스타일 - 전문적 */
    .badge {{
        display: inline-block;
        background: linear-gradient(135deg, var(--accent-gold), var(--accent-gold-light));
        color: var(--primary-navy-dark);
        padding: 0.5rem 1.2rem;
        border-radius: 25px;
        font-weight: 700;
        font-size: 1rem;
        margin: 0.3rem;
        box-shadow: var(--shadow-gold);
        border: 2px solid rgba(255, 255, 255, 0.3);
        letter-spacing: 0.5px;
    }}

    /* 글씨 크기 조절 버튼 */
    .font-size-btn {{
        padding: 8px 16px !important;
        border-radius: 10px !important;
        margin: 0 5px !important;
    }}
    .font-size-btn.active {{
        background: #4CAF50 !important;
        color: white !important;
    }}

    /* 선택 모드 카드 - 개선된 버전 (컴팩트) */
    .mode-card {{
        background: white;
        border: 3px solid #E0E0E0;
        border-radius: 15px;
        padding: 1rem 0.8rem;
        margin: 0.5rem 0;
        text-align: center;
        cursor: pointer;
        transition: all 0.3s ease;
        position: relative;
        overflow: hidden;
        min-height: 140px;
        display: flex;
        flex-direction: column;
        justify-content: center;
        align-items: center;
    }}
    .mode-card::before {{
        content: '';
        position: absolute;
        top: 0;
        left: 0;
        right: 0;
        bottom: 0;
        background: linear-gradient(135deg, rgba(255,255,255,0) 0%, rgba(255,255,255,0.8) 100%);
        opacity: 0;
        transition: opacity 0.3s ease;
    }}
    .mode-card:hover {{
        border-color: var(--accent-gold);
        transform: translateY(-8px) scale(1.02);
        box-shadow: var(--shadow-xl);
    }}
    .mode-card:hover::before {{
        opacity: 1;
    }}
    .mode-card.chat-card {{ background: linear-gradient(135deg, #ebf8ff 0%, #bee3f8 100%); border-color: var(--primary-navy-light); }}
    .mode-card.chat-card:hover {{ border-color: var(--primary-navy); box-shadow: var(--shadow-navy); }}
    .mode-card.voice-card {{ background: linear-gradient(135deg, #fffff0 0%, #fefcbf 100%); border-color: var(--accent-gold); }}
    .mode-card.voice-card:hover {{ border-color: var(--accent-gold-dark); box-shadow: var(--shadow-gold); }}
    .mode-card.youtube-card {{ background: linear-gradient(135deg, #fff5f5 0%, #fed7d7 100%); border-color: #fc8181; }}
    .mode-card.youtube-card:hover {{ border-color: #e53e3e; box-shadow: 0 15px 40px rgba(229, 62, 62, 0.2); }}
    .mode-card .emoji {{
        font-size: 2.5rem;
        margin-bottom: 0.5rem;
        animation: float 3s ease-in-out infinite;
    }}
    @keyframes float {{
        0%, 100% {{ transform: translateY(0); }}
        50% {{ transform: translateY(-8px); }}
    }}
    .mode-card .title {{
        font-size: 1.1rem;
        font-weight: bold;
        color: #333;
        margin-bottom: 0.3rem;
    }}
    .mode-card .desc {{
        font-size: 0.9rem;
        color: #555;
        margin-top: 0.3rem;
        line-height: 1.4;
    }}
    .mode-card .feature-list {{
        font-size: 0.85rem;
        color: #666;
        margin-top: 0.5rem;
        text-align: left;
        padding: 0 0.3rem;
    }}
    .mode-card .feature-list li {{
        margin: 0.2rem 0;
    }}

    /* 환영 화면 헤더 - 작가의집 브랜딩 (컴팩트) */
    .welcome-header {{
        text-align: center;
        padding: 1.2rem 1.5rem;
        background: linear-gradient(135deg, var(--primary-navy-dark) 0%, var(--primary-navy) 50%, var(--primary-navy-light) 100%);
        color: white;
        border-radius: 15px;
        margin-bottom: 1rem;
        box-shadow: var(--shadow-lg);
        position: relative;
        overflow: hidden;
    }}
    .welcome-header::before {{
        content: '';
        position: absolute;
        bottom: 0;
        left: 0;
        right: 0;
        height: 4px;
        background: linear-gradient(90deg, var(--accent-gold-dark), var(--accent-gold), var(--accent-gold-light), var(--accent-gold), var(--accent-gold-dark));
    }}
    .welcome-header::after {{
        content: '';
        position: absolute;
        top: 0;
        left: 0;
        right: 0;
        bottom: 0;
        background: radial-gradient(ellipse at top right, rgba(201, 169, 98, 0.15) 0%, transparent 50%);
        pointer-events: none;
    }}
    .welcome-header .logo {{
        font-size: 2.5rem;
        margin-bottom: 0.4rem;
        filter: drop-shadow(2px 2px 4px rgba(0,0,0,0.3));
    }}
    .welcome-header .brand-badge {{
        display: inline-block;
        background: var(--accent-gold);
        color: var(--primary-navy-dark);
        padding: 0.2rem 0.8rem;
        border-radius: 15px;
        font-size: 0.75rem;
        font-weight: 700;
        margin-bottom: 0.5rem;
        letter-spacing: 0.5px;
        text-transform: uppercase;
    }}
    .welcome-header .app-name {{
        font-size: 1.8rem;
        font-weight: 800;
        text-shadow: 2px 2px 4px rgba(0,0,0,0.3);
        margin-bottom: 0.2rem;
        letter-spacing: -0.5px;
    }}
    .welcome-header .app-subtitle {{
        font-size: 1rem;
        font-weight: 600;
        color: var(--accent-gold-light);
        margin-bottom: 0.4rem;
    }}
    .welcome-header .tagline {{
        font-size: 0.9rem;
        opacity: 0.9;
        font-weight: 400;
        line-height: 1.4;
    }}

    /* 사용자 타입 선택 카드 (컴팩트) */
    .user-type-card {{
        background: white;
        border: 3px solid #E0E0E0;
        border-radius: 15px;
        padding: 1rem;
        margin: 0.5rem 0;
        text-align: center;
        cursor: pointer;
        transition: all 0.3s ease;
        min-height: 120px;
    }}
    .user-type-card:hover {{
        transform: translateY(-5px);
        box-shadow: 0 12px 30px rgba(0,0,0,0.15);
    }}
    .user-type-card.new-user {{
        background: linear-gradient(135deg, var(--pastel-blue) 0%, #bee3f8 100%);
        border-color: var(--primary-navy-light);
    }}
    .user-type-card.new-user:hover {{
        border-color: var(--primary-navy);
        box-shadow: var(--shadow-navy);
    }}
    .user-type-card.returning-user {{
        background: linear-gradient(135deg, var(--pastel-yellow) 0%, #fefcbf 100%);
        border-color: var(--accent-gold);
    }}
    .user-type-card.returning-user:hover {{
        border-color: var(--accent-gold-dark);
        box-shadow: var(--shadow-gold);
    }}
    .user-type-card .card-emoji {{
        font-size: 2.2rem;
        margin-bottom: 0.4rem;
    }}
    .user-type-card .card-title {{
        font-size: 1.1rem;
        font-weight: bold;
        color: #333;
    }}
    .user-type-card .card-desc {{
        font-size: 0.85rem;
        color: #666;
        margin-top: 0.3rem;
    }}

    /* 온보딩 튜토리얼 (컴팩트) */
    .onboarding-container {{
        background: linear-gradient(135deg, #f5f7fa 0%, #c3cfe2 100%);
        border-radius: 15px;
        padding: 1rem;
        margin: 0.8rem 0;
        box-shadow: 0 8px 20px rgba(0,0,0,0.1);
    }}
    .onboarding-title {{
        text-align: center;
        font-size: 1.2rem;
        font-weight: bold;
        color: #333;
        margin-bottom: 0.8rem;
    }}
    .onboarding-step {{
        display: flex;
        align-items: center;
        background: white;
        border-radius: 12px;
        padding: 0.8rem;
        margin: 0.5rem 0;
        box-shadow: 0 3px 10px rgba(0,0,0,0.08);
        transition: all 0.3s ease;
        animation: slide-in 0.5s ease forwards;
        opacity: 0;
        transform: translateX(-20px);
    }}
    .onboarding-step:nth-child(2) {{ animation-delay: 0.1s; }}
    .onboarding-step:nth-child(3) {{ animation-delay: 0.3s; }}
    .onboarding-step:nth-child(4) {{ animation-delay: 0.5s; }}
    @keyframes slide-in {{
        to {{
            opacity: 1;
            transform: translateX(0);
        }}
    }}
    .onboarding-step:hover {{
        transform: translateY(-3px);
        box-shadow: 0 8px 25px rgba(0,0,0,0.12);
    }}
    .onboarding-step .step-number {{
        width: 36px;
        height: 36px;
        background: linear-gradient(135deg, var(--primary-navy) 0%, var(--primary-navy-light) 100%);
        color: white;
        border-radius: 50%;
        display: flex;
        align-items: center;
        justify-content: center;
        font-size: 1rem;
        font-weight: bold;
        margin-right: 0.8rem;
        flex-shrink: 0;
        border: 2px solid var(--accent-gold);
        box-shadow: var(--shadow-sm);
    }}
    .onboarding-step .step-content {{
        flex: 1;
    }}
    .onboarding-step .step-title {{
        font-size: 0.95rem;
        font-weight: bold;
        color: #333;
        margin-bottom: 0.2rem;
    }}
    .onboarding-step .step-desc {{
        font-size: 0.85rem;
        color: #666;
    }}
    .onboarding-step .step-icon {{
        font-size: 1.8rem;
        margin-left: 0.5rem;
    }}

    /* 저장된 작업 목록 (컴팩트) */
    .saved-work-container {{
        background: linear-gradient(135deg, #FFF9E6 0%, #FFE0B2 100%);
        border-radius: 12px;
        padding: 0.8rem;
        margin: 0.6rem 0;
        border: 2px solid #FFB74D;
    }}
    .saved-work-title {{
        font-size: 1rem;
        font-weight: bold;
        color: #E65100;
        margin-bottom: 0.5rem;
        display: flex;
        align-items: center;
        gap: 0.3rem;
    }}
    .saved-work-item {{
        background: white;
        border-radius: 10px;
        padding: 0.6rem 0.8rem;
        margin: 0.4rem 0;
        display: flex;
        align-items: center;
        justify-content: space-between;
        box-shadow: 0 2px 6px rgba(0,0,0,0.08);
        transition: all 0.3s ease;
    }}
    .saved-work-item:hover {{
        transform: translateX(5px);
        box-shadow: 0 5px 15px rgba(0,0,0,0.12);
    }}
    .saved-work-info {{
        flex: 1;
    }}
    .saved-work-info .work-title {{
        font-size: 0.95rem;
        font-weight: bold;
        color: #333;
    }}
    .saved-work-info .work-meta {{
        font-size: 0.8rem;
        color: #666;
        margin-top: 0.2rem;
    }}
    .saved-work-info .work-progress {{
        font-size: 0.75rem;
        color: #4CAF50;
        font-weight: 600;
        margin-top: 0.15rem;
    }}

    /* 빠른 이어쓰기 버튼 */
    .continue-btn {{
        background: linear-gradient(135deg, #FF6B6B 0%, #FFE66D 100%);
        color: #333;
        border: none;
        border-radius: 12px;
        padding: 0.6rem 1.2rem;
        font-weight: bold;
        cursor: pointer;
        transition: all 0.3s ease;
    }}
    .continue-btn:hover {{
        transform: scale(1.05);
        box-shadow: 0 5px 15px rgba(255, 107, 107, 0.4);
    }}

    /* 타임라인 스타일 */
    .timeline-container {{
        position: relative;
        padding-left: 30px;
        margin: 1.5rem 0;
    }}
    .timeline-item {{
        position: relative;
        padding: 1rem 0;
        border-left: 3px solid #E0E0E0;
        padding-left: 25px;
        margin-left: 10px;
    }}
    .timeline-item.completed {{
        border-left-color: #2E7D32;
    }}
    .timeline-item.current {{
        border-left-color: #1565C0;
    }}
    .timeline-dot {{
        position: absolute;
        left: -10px;
        top: 1.2rem;
        width: 18px;
        height: 18px;
        border-radius: 50%;
        background: #E0E0E0;
    }}
    .timeline-dot.completed {{
        background: #2E7D32;
    }}
    .timeline-dot.current {{
        background: #1565C0;
        box-shadow: 0 0 0 4px rgba(21, 101, 192, 0.3);
    }}

    /* ============================================ */
    /* 입력 필드 (모바일 키보드 최적화, 컴팩트)    */
    /* ============================================ */
    .stTextArea textarea, .stTextInput input {{
        font-size: var(--font-base) !important;
        line-height: 1.5 !important;
        padding: 10px 12px !important;
        border-radius: 10px !important;
        border: 2px solid #E0E0E0 !important;
        min-height: 44px !important;
        /* 모바일 키보드 가림 방지 */
        scroll-margin-bottom: 150px;
    }}
    .stTextArea textarea:focus, .stTextInput input:focus {{
        border-color: var(--focus-color) !important;
        box-shadow: 0 0 0 4px rgba(25, 118, 210, 0.2) !important;
    }}
    /* 입력 필드 레이블 접근성 */
    .stTextArea label, .stTextInput label {{
        font-weight: 600 !important;
        color: var(--text-primary) !important;
        margin-bottom: 8px !important;
    }}

    /* 셀렉트박스 (컴팩트) */
    .stSelectbox > div > div {{
        font-size: var(--font-base) !important;
        min-height: 44px !important;
        border-radius: 10px !important;
    }}
    .stSelectbox label {{
        font-weight: 600 !important;
        color: var(--text-primary) !important;
    }}

    /* 체크박스와 라디오 버튼 터치 타겟 확대 (컴팩트) */
    .stCheckbox > label, .stRadio > label {{
        min-height: 44px !important;
        display: flex !important;
        align-items: center !important;
        padding: 4px 0 !important;
    }}
    .stCheckbox > label > span, .stRadio > label > span {{
        padding-left: 12px !important;
    }}

    /* ARIA 라이브 리전 */
    [role="status"], [aria-live="polite"] {{
        position: relative;
    }}
    [aria-busy="true"] {{
        opacity: 0.7;
        pointer-events: none;
    }}

    /* 빠른 시작 배지 */
    .quick-badge {{
        display: inline-block;
        background: #FF6F00;
        color: white;
        padding: 4px 12px;
        border-radius: 20px;
        font-size: 0.85rem;
        font-weight: 700;
        margin-left: 8px;
        animation: pulse 2s infinite;
    }}
    @keyframes pulse {{
        0%, 100% {{ opacity: 1; }}
        50% {{ opacity: 0.7; }}
    }}

    /* 장르 프리셋 카드 */
    .preset-card {{
        background: white;
        border: 2px solid #E0E0E0;
        border-radius: 12px;
        padding: 1rem;
        margin: 0.5rem 0;
        cursor: pointer;
        transition: all 0.2s;
    }}
    .preset-card:hover {{
        border-color: var(--accent-gold);
        box-shadow: var(--shadow-md);
    }}
    .preset-card.selected {{
        border-color: var(--primary-navy);
        background: var(--pastel-blue);
        border-width: 3px;
    }}

    /* ============================================ */
    /* 반응형 디자인: 태블릿 (768px 이하)          */
    /* ============================================ */
    @media (max-width: 768px) {{
        /* 기본 설정 */
        html, body, [class*="st-"] {{
            font-size: 15px !important;
        }}

        /* 사이드바 숨김/축소 */
        [data-testid="stSidebar"] {{
            transform: translateX(-100%);
            transition: transform 0.3s ease;
        }}
        [data-testid="stSidebar"][aria-expanded="true"] {{
            transform: translateX(0);
        }}
        [data-testid="stSidebarNav"] {{
            padding-top: 0.5rem;
        }}

        /* 메인 콘텐츠 */
        .main .block-container {{
            padding: 0.5rem !important;
            max-width: 100% !important;
        }}

        /* 헤더 */
        .main-header {{
            font-size: 1.4rem !important;
            padding: 0.6rem 0.8rem;
            margin-bottom: 0.5rem;
        }}
        .step-header {{
            font-size: 1.1rem !important;
            padding: 0.5rem 0.8rem;
        }}

        /* 진행률 박스 */
        .progress-box {{
            padding: 0.6rem;
            margin: 0.5rem 0;
        }}
        .progress-text {{
            font-size: 1rem !important;
        }}

        /* 버튼 - 전체 너비 */
        .stButton > button {{
            min-height: 44px !important;
            font-size: 0.9rem !important;
            width: 100% !important;
            margin: 0.3rem 0 !important;
        }}

        /* 입력 필드 - iOS 확대 방지 (최소 16px) */
        .stTextArea textarea, .stTextInput input {{
            font-size: 16px !important;
            min-height: 44px !important;
        }}

        /* 모드 카드 */
        .mode-card {{
            padding: 0.8rem;
            margin: 0.4rem 0;
        }}
        .mode-card .emoji {{
            font-size: 2rem;
        }}
        .mode-card .title {{
            font-size: 1rem;
        }}

        /* 도움말/힌트 박스 */
        .help-box, .hint-box, .warning-box, .success-box {{
            padding: 0.6rem 0.8rem;
            margin: 0.5rem 0;
        }}

        /* 현재 섹션 박스 */
        .current-section-box {{
            padding: 0.5rem 0.8rem;
        }}

        /* 결과 카드 */
        .result-card {{
            padding: 0.8rem;
            margin: 0.5rem 0;
        }}

        /* 채팅 버블 */
        .chat-bubble {{
            max-width: 95%;
            padding: 0.6rem 0.8rem;
            font-size: 0.95rem;
        }}

        /* 플로팅 버튼 위치 조정 */
        .floating-help-btn {{
            bottom: 15px;
            right: 15px;
            padding: 8px 14px;
            font-size: 0.9rem;
        }}

        /* 컬럼 스택 */
        [data-testid="column"] {{
            width: 100% !important;
            flex: 0 0 100% !important;
            min-width: 100% !important;
        }}

        /* 탭 스크롤 가능 */
        .stTabs [data-baseweb="tab-list"] {{
            overflow-x: auto;
            -webkit-overflow-scrolling: touch;
        }}
    }}

    /* ============================================ */
    /* 반응형 디자인: 스마트폰 (480px 이하)         */
    /* ============================================ */
    @media (max-width: 480px) {{
        html, body, [class*="st-"] {{
            font-size: 14px !important;
        }}

        /* 메인 콘텐츠 패딩 최소화 */
        .main .block-container {{
            padding: 0.3rem !important;
        }}

        /* 헤더 */
        .main-header {{
            font-size: 1.2rem !important;
            padding: 0.5rem 0.6rem;
            border-radius: 10px;
        }}
        .step-header {{
            font-size: 1rem !important;
            padding: 0.4rem 0.6rem;
            border-radius: 8px;
        }}

        /* 진행률 */
        .progress-box {{
            padding: 0.5rem;
            border-radius: 10px;
        }}
        .progress-text {{
            font-size: 0.95rem !important;
        }}
        .star-progress {{
            font-size: 1.1rem;
            letter-spacing: 2px;
        }}

        /* 버튼 */
        .stButton > button {{
            font-size: 0.85rem !important;
            padding: 0.4rem 0.8rem !important;
            border-radius: 8px !important;
        }}
        .big-chat-button {{
            font-size: 1rem !important;
            min-height: 48px !important;
        }}

        /* 입력 필드 */
        .stTextArea textarea {{
            min-height: 80px !important;
        }}

        /* 모드 카드 */
        .mode-card {{
            padding: 0.6rem;
            border-radius: 10px;
        }}
        .mode-card .emoji {{
            font-size: 1.6rem;
        }}
        .mode-card .title {{
            font-size: 0.95rem;
        }}
        .mode-card .desc {{
            font-size: 0.8rem;
        }}

        /* 채팅 모드 */
        .chat-mode-header {{
            font-size: 1.1rem;
            padding: 0.6rem;
        }}
        .chat-progress {{
            font-size: 0.9rem;
            padding: 0.5rem;
        }}
        .chat-bubble {{
            padding: 0.5rem 0.7rem;
            font-size: 0.9rem;
            border-radius: 12px;
        }}

        /* 타임라인 */
        .timeline-container {{
            padding-left: 15px;
        }}
        .timeline-item {{
            padding-left: 15px;
        }}

        /* 배지 */
        .badge {{
            padding: 0.3rem 0.6rem;
            font-size: 0.75rem;
        }}

        /* 빠른 배지 */
        .quick-badge {{
            font-size: 0.65rem;
            padding: 2px 6px;
        }}

        /* 플로팅 버튼 - safe-area 적용 */
        .floating-help-btn {{
            bottom: calc(10px + var(--safe-area-inset-bottom));
            right: calc(10px + var(--safe-area-inset-right));
            padding: 6px 12px;
            font-size: 0.8rem;
            border-radius: 30px;
        }}

        /* 입력 컨테이너 - 키보드/safe-area 대응 */
        .chat-input-container {{
            padding-bottom: calc(0.5rem + var(--safe-area-inset-bottom));
        }}
    }}

    /* ============================================ */
    /* iOS Safari 하단 바 대응                      */
    /* ============================================ */
    @supports (padding-bottom: env(safe-area-inset-bottom)) {{
        .floating-help-btn {{
            bottom: calc(20px + env(safe-area-inset-bottom));
        }}
        .chat-input-container {{
            padding-bottom: calc(1.5rem + env(safe-area-inset-bottom));
        }}
    }}

    /* ============================================ */
    /* 세로 모드 최적화 (높이 기반)                 */
    /* ============================================ */
    @media (max-height: 600px) and (orientation: landscape) {{
        .main-header {{
            padding: 0.5rem;
            margin-bottom: 0.5rem;
        }}
        .progress-box {{
            padding: 0.8rem;
            margin: 0.5rem 0;
        }}
        .chat-input-container {{
            padding: 0.8rem;
        }}
    }}

    /* ============================================ */
    /* 고대비 모드 지원 (접근성)                    */
    /* ============================================ */
    @media (prefers-contrast: high) {{
        :root {{
            --pastel-pink: #FFCDD2;
            --pastel-blue: #BBDEFB;
            --pastel-green: #C8E6C9;
            --pastel-yellow: #FFF9C4;
            --pastel-purple: #E1BEE7;
            --pastel-orange: #FFE0B2;
        }}
        .stButton > button {{
            border: 3px solid currentColor !important;
        }}
        .help-box, .hint-box, .warning-box, .success-box {{
            border-width: 4px !important;
        }}
    }}

    /* ============================================ */
    /* 애니메이션 감소 모드 (접근성)                */
    /* ============================================ */
    @media (prefers-reduced-motion: reduce) {{
        *, *::before, *::after {{
            animation-duration: 0.01ms !important;
            animation-iteration-count: 1 !important;
            transition-duration: 0.01ms !important;
            scroll-behavior: auto !important;
        }}
        .celebrate, .bounce {{
            animation: none !important;
        }}
        @keyframes pulse {{
            0%, 100% {{ opacity: 1; }}
        }}
        @keyframes pulse-recording {{
            0%, 100% {{ opacity: 1; box-shadow: none; }}
        }}
    }}

    /* ============================================ */
    /* 다크 모드 대응 (시스템 설정 기반)            */
    /* ============================================ */
    @media (prefers-color-scheme: dark) {{
        .mode-card {{
            background: #2d2d2d;
            border-color: #444;
        }}
        .mode-card .title {{
            color: #fff;
        }}
        .mode-card .desc {{
            color: #bbb;
        }}
    }}

    /* ============================================ */
    /* 인쇄 스타일                                  */
    /* ============================================ */
    @media print {{
        .floating-help-btn,
        .stButton,
        [data-testid="stSidebar"] {{
            display: none !important;
        }}
        .main .block-container {{
            padding: 0 !important;
            max-width: 100% !important;
        }}
        .progress-box {{
            background: #f0f0f0 !important;
            color: #000 !important;
            -webkit-print-color-adjust: exact;
        }}
    }}

    /* ============================================ */
    /* 채팅 모드 전용 스타일 (컴팩트)               */
    /* ============================================ */
    .chat-mode-container {{
        max-width: 800px;
        margin: 0 auto;
        padding: 0.5rem;
    }}
    .chat-mode-header {{
        text-align: center;
        padding: 0.8rem 1rem;
        background: linear-gradient(135deg, #FF6B6B 0%, #FFE66D 50%, #4ECDC4 100%);
        color: white;
        border-radius: 15px;
        margin-bottom: 0.8rem;
        font-size: 1.3rem;
        font-weight: bold;
        text-shadow: 2px 2px 4px rgba(0,0,0,0.2);
        box-shadow: 0 6px 20px rgba(255, 107, 107, 0.3);
    }}
    .chat-progress {{
        background: linear-gradient(135deg, #E8F5E9 0%, #C8E6C9 100%);
        padding: 0.6rem 0.8rem;
        border-radius: 12px;
        margin-bottom: 0.8rem;
        text-align: center;
        font-size: 1rem;
        font-weight: bold;
        color: #2E7D32;
        border: 2px solid #81C784;
        box-shadow: 0 3px 10px rgba(76, 175, 80, 0.2);
    }}
    .chat-bubble {{
        padding: 0.8rem 1rem;
        border-radius: 15px;
        margin: 0.5rem 0;
        font-size: 1rem;
        line-height: 1.5;
        max-width: 90%;
        box-shadow: 0 3px 8px rgba(0,0,0,0.1);
    }}
    .chat-bubble-ai {{
        background: linear-gradient(135deg, #E3F2FD 0%, #BBDEFB 100%);
        border-bottom-left-radius: 8px;
        margin-right: auto;
        border: 2px solid #90CAF9;
    }}
    .chat-bubble-user {{
        background: linear-gradient(135deg, #C8E6C9 0%, #A5D6A7 100%);
        border-bottom-right-radius: 8px;
        margin-left: auto;
        text-align: right;
        border: 2px solid #81C784;
    }}
    .chat-input-container {{
        position: sticky;
        bottom: 0;
        background: white;
        padding: 0.8rem;
        border-top: 2px solid #E0E0E0;
        border-radius: 15px 15px 0 0;
        box-shadow: 0 -4px 15px rgba(0,0,0,0.1);
    }}
    .big-chat-button {{
        font-size: 1.1rem !important;
        padding: 0.8rem 1.5rem !important;
        border-radius: 12px !important;
        min-height: 50px !important;
        font-weight: bold !important;
    }}
    .restart-button {{
        background: linear-gradient(135deg, #FF5722 0%, #FF7043 100%) !important;
        color: white !important;
        font-size: 1rem !important;
        padding: 0.6rem 1.2rem !important;
        border-radius: 10px !important;
        font-weight: bold !important;
    }}
    .result-card {{
        background: linear-gradient(135deg, var(--pastel-yellow) 0%, #fefcbf 100%);
        border: 2px solid var(--accent-gold);
        border-radius: 12px;
        padding: 1rem;
        margin: 0.6rem 0;
        font-size: 1rem;
        box-shadow: var(--shadow-md);
        position: relative;
    }}
    .result-card::before {{
        content: '';
        position: absolute;
        top: 0;
        left: 0;
        right: 0;
        height: 3px;
        background: linear-gradient(90deg, var(--accent-gold-dark), var(--accent-gold), var(--accent-gold-light));
        border-radius: 12px 12px 0 0;
    }}
    .result-card h3, .result-card h4 {{
        color: var(--primary-navy);
        margin-bottom: 0.5rem;
        font-weight: 700;
    }}
    .result-card p {{
        color: var(--text-secondary);
        margin: 0.3rem 0;
        line-height: 1.5;
    }}
    /* 예시 버튼 스타일 (컴팩트) */
    div[data-testid="stButton"] button {{
        font-size: 0.95rem !important;
        border-radius: 10px !important;
        transition: all 0.2s ease !important;
    }}
    div[data-testid="stButton"] button:hover {{
        transform: scale(1.02) !important;
    }}

    /* ============================================ */
    /* 음성 모드 전용 스타일 (컴팩트)                */
    /* ============================================ */
    .voice-mode-container {{
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        border-radius: 15px;
        padding: 1rem;
        margin: 0.8rem 0;
        color: white;
        box-shadow: 0 6px 25px rgba(102, 126, 234, 0.3);
    }}
    .voice-mode-header {{
        font-size: 1.4rem;
        font-weight: bold;
        text-align: center;
        margin-bottom: 0.6rem;
    }}
    .voice-mode-description {{
        text-align: center;
        opacity: 0.9;
        font-size: 0.95rem;
        margin-bottom: 0.8rem;
    }}
    .voice-mic-button {{
        display: flex;
        justify-content: center;
        align-items: center;
        width: 80px;
        height: 80px;
        margin: 0.8rem auto;
        background: white;
        border-radius: 50%;
        cursor: pointer;
        transition: all 0.3s ease;
        box-shadow: 0 4px 15px rgba(0,0,0,0.2);
    }}
    .voice-mic-button:hover {{
        transform: scale(1.1);
        box-shadow: 0 6px 30px rgba(0,0,0,0.3);
    }}
    .voice-mic-icon {{
        font-size: 2rem;
    }}
    .recording-indicator {{
        display: inline-flex;
        align-items: center;
        justify-content: center;
        background: #ff4444;
        color: white;
        padding: 0.5rem 1rem;
        border-radius: 20px;
        font-weight: bold;
        font-size: 0.9rem;
        animation: pulse-recording 1.5s infinite;
    }}
    @keyframes pulse-recording {{
        0%, 100% {{ opacity: 1; box-shadow: 0 0 0 0 rgba(255, 68, 68, 0.7); }}
        50% {{ opacity: 0.8; box-shadow: 0 0 0 15px rgba(255, 68, 68, 0); }}
    }}
    .voice-transcription-box {{
        background: white;
        border-radius: 12px;
        padding: 0.8rem;
        margin: 0.8rem 0;
        color: #333;
        border: 2px solid #667eea;
        min-height: 100px;
    }}
    .voice-transcription-preview {{
        font-size: 1rem;
        line-height: 1.5;
        color: #333;
    }}
    .voice-action-button {{
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%) !important;
        color: white !important;
        font-size: 1rem !important;
        padding: 0.6rem 1.2rem !important;
        border-radius: 12px !important;
        font-weight: bold !important;
        margin: 0.5rem 0 !important;
    }}
    .voice-status-badge {{
        display: inline-block;
        background: #4CAF50;
        color: white;
        padding: 0.4rem 1rem;
        border-radius: 20px;
        font-size: 0.9rem;
        font-weight: bold;
    }}
    .voice-file-info {{
        background: #F5F5F5;
        border-radius: 10px;
        padding: 1rem;
        margin: 1rem 0;
        border-left: 4px solid #667eea;
    }}

    /* 플로팅 도움 챗봇 버튼 - 전문적 */
    .floating-help-btn {{
        position: fixed;
        bottom: 30px;
        right: 30px;
        z-index: 1000;
        background: linear-gradient(135deg, var(--primary-navy) 0%, var(--primary-navy-light) 100%);
        color: white;
        border: none;
        border-radius: 50px;
        padding: 15px 25px;
        font-size: 1.1rem;
        font-weight: 700;
        cursor: pointer;
        box-shadow: var(--shadow-navy);
        transition: all 0.3s ease;
        border-bottom: 3px solid var(--accent-gold);
    }}
    .floating-help-btn:hover {{
        transform: translateY(-3px);
        box-shadow: 0 8px 30px rgba(30, 58, 95, 0.4);
        background: linear-gradient(135deg, var(--primary-navy-light) 0%, var(--primary-navy) 100%);
    }}

    /* 연락하기 박스 */
    .contact-section {{
        background: linear-gradient(135deg, #E3F2FD 0%, #F3E5F5 100%);
        border-radius: 16px;
        padding: 1.5rem;
        margin: 1rem 0;
        border: 2px solid #667eea;
    }}
    .contact-section h4 {{
        color: #333;
        margin-bottom: 1rem;
    }}

    /* 도움 챗봇 섹션 */
    .help-chatbot-section {{
        background: white;
        border-radius: 16px;
        padding: 1rem;
        margin: 1rem 0;
        box-shadow: 0 4px 15px rgba(0,0,0,0.1);
        border: 1px solid #E0E0E0;
    }}

    /* FAQ 버튼 스타일 */
    .faq-button {{
        background: #F0F4FF;
        border: 1px solid #667eea;
        border-radius: 20px;
        padding: 8px 16px;
        margin: 4px 2px;
        font-size: 0.9rem;
        cursor: pointer;
        transition: all 0.2s;
    }}
    .faq-button:hover {{
        background: #667eea;
        color: white;
    }}

    /* ============================================ */
    /* 모달 팝업 스타일 (도움말/연락하기용)          */
    /* ============================================ */
    .modal-overlay {{
        position: fixed;
        top: 0;
        left: 0;
        right: 0;
        bottom: 0;
        background: rgba(0, 0, 0, 0.5);
        z-index: 9998;
        display: flex;
        align-items: center;
        justify-content: center;
        padding: 20px;
        backdrop-filter: blur(4px);
    }}

    .modal-container {{
        background: white;
        border-radius: 20px;
        max-width: 600px;
        width: 100%;
        max-height: 80vh;
        overflow-y: auto;
        box-shadow: 0 20px 60px rgba(0, 0, 0, 0.3);
        animation: modal-slide-up 0.3s ease-out;
        position: relative;
    }}

    @keyframes modal-slide-up {{
        from {{
            opacity: 0;
            transform: translateY(30px);
        }}
        to {{
            opacity: 1;
            transform: translateY(0);
        }}
    }}

    .modal-header {{
        background: linear-gradient(135deg, var(--primary-navy) 0%, var(--primary-navy-light) 100%);
        color: white;
        padding: 1.5rem;
        border-radius: 20px 20px 0 0;
        display: flex;
        justify-content: space-between;
        align-items: center;
        position: sticky;
        top: 0;
        z-index: 10;
    }}

    .modal-header h3 {{
        margin: 0;
        font-size: 1.3rem;
        font-weight: 700;
    }}

    .modal-close-btn {{
        background: rgba(255, 255, 255, 0.2);
        border: none;
        color: white;
        width: 40px;
        height: 40px;
        border-radius: 50%;
        font-size: 1.5rem;
        cursor: pointer;
        display: flex;
        align-items: center;
        justify-content: center;
        transition: all 0.2s ease;
    }}

    .modal-close-btn:hover {{
        background: rgba(255, 255, 255, 0.3);
        transform: scale(1.1);
    }}

    .modal-body {{
        padding: 1.5rem;
    }}

    /* 도움 버튼 강화 스타일 */
    .help-contact-buttons {{
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        border-radius: 16px;
        padding: 1.2rem;
        margin: 1rem 0;
        box-shadow: 0 4px 20px rgba(102, 126, 234, 0.3);
    }}

    .help-contact-buttons h4 {{
        color: white;
        margin: 0 0 1rem 0;
        font-size: 1.1rem;
        text-align: center;
    }}

    .help-btn-row {{
        display: flex;
        gap: 0.8rem;
    }}

    .help-action-btn {{
        flex: 1;
        padding: 1rem;
        border-radius: 12px;
        border: none;
        font-size: 1rem;
        font-weight: 600;
        cursor: pointer;
        transition: all 0.2s ease;
        display: flex;
        flex-direction: column;
        align-items: center;
        gap: 0.5rem;
    }}

    .help-action-btn.chatbot {{
        background: linear-gradient(135deg, #4CAF50 0%, #45a049 100%);
        color: white;
    }}

    .help-action-btn.contact {{
        background: linear-gradient(135deg, #FF9800 0%, #F57C00 100%);
        color: white;
    }}

    .help-action-btn:hover {{
        transform: translateY(-2px);
        box-shadow: 0 4px 15px rgba(0, 0, 0, 0.2);
    }}

    .help-action-btn .btn-icon {{
        font-size: 1.5rem;
    }}

    .help-action-btn .btn-text {{
        font-size: 0.9rem;
    }}

    /* 연락처 방법 표시 스타일 */
    .contact-method-info {{
        background: linear-gradient(135deg, #FFF3E0 0%, #FFE0B2 100%);
        border-radius: 12px;
        padding: 1rem;
        margin: 0.8rem 0;
        border-left: 4px solid #FF9800;
    }}

    .contact-method-info p {{
        margin: 0.3rem 0;
        font-size: 0.95rem;
        color: #333;
    }}

    .contact-method-info .method-title {{
        font-weight: 700;
        color: #E65100;
        margin-bottom: 0.5rem;
    }}

    /* 메시지 전송 성공 팝업 스타일 */
    .success-popup {{
        background: linear-gradient(135deg, #E8F5E9 0%, #C8E6C9 100%);
        border: 2px solid #4CAF50;
        border-radius: 16px;
        padding: 1.5rem;
        text-align: center;
        margin: 1rem 0;
    }}

    .success-popup h4 {{
        color: #2E7D32;
        margin: 0 0 0.8rem 0;
    }}

    .success-popup p {{
        color: #333;
        margin: 0.3rem 0;
    }}

    /* 답변 받을 방법 선택 스타일 */
    .reply-method-select {{
        background: #F5F5F5;
        border-radius: 12px;
        padding: 1rem;
        margin-top: 1rem;
    }}

    .reply-method-select label {{
        display: block;
        padding: 0.8rem;
        margin: 0.5rem 0;
        background: white;
        border-radius: 8px;
        cursor: pointer;
        border: 2px solid transparent;
        transition: all 0.2s ease;
    }}

    .reply-method-select label:hover {{
        border-color: #667eea;
        background: #F0F4FF;
    }}

    .reply-method-select input[type="radio"]:checked + span {{
        color: #667eea;
        font-weight: 600;
    }}
</style>
""", unsafe_allow_html=True)

# Achievement CSS (별도 로드 - f-string 충돌 방지)
st.markdown("<style>" + ACHIEVEMENT_CSS + "</style>", unsafe_allow_html=True)


def init_session_state():
    """세션 상태 초기화 - 안전한 초기화 + 에러 핸들링"""
    try:
        defaults = {
            "current_step": 1,
            "book_info": {},
            "generated_titles": "",
            "selected_title": "",
            "generated_toc": "",
            "parsed_toc": [],  # 파싱된 목차 구조
            "drafts": {},
            "current_section_index": 0,  # 현재 작성 중인 장 인덱스
            "chat_messages": [],
            "show_chatbot": False,
            "generated_proposal": "",
            "generated_landing_page": "",
            "author_info": {},
            "webinar_info": {},
            "button_loading_state": {},  # 버튼 로딩 상태 추적
            "last_action_feedback": None,  # 마지막 작업 피드백
            "last_save_time": None,  # 마지막 저장 시간
            "last_autosave_time": None,  # 자동 저장 시간
            "session_start_time": datetime.now().isoformat(),  # 세션 시작 시간
            # 채팅 모드 관련 상태
            "chat_mode_active": False,  # 채팅 모드 활성화 여부
            "chat_mode_step": 0,  # 채팅 모드 현재 단계 (0-5)
            "chat_mode_history": [],  # 채팅 모드 대화 기록
            "chat_mode_data": {},  # 채팅 모드에서 수집한 데이터
            # 음성 모드 관련 상태
            "voice_mode_active": False,  # 음성 모드 활성화 여부
            "voice_transcribed_text": None,  # 음성에서 변환된 텍스트
            "voice_edited_text": None,  # 사용자가 수정한 텍스트
            # 유튜브 모드 관련 상태
            "youtube_mode_active": False,  # 유튜브 모드 활성화 여부
            "youtube_urls": [],  # 입력된 유튜브 URL 리스트
            "youtube_videos": [],  # 처리된 영상 정보 리스트
            "youtube_transcripts": {},  # 영상별 자막 저장
            "youtube_merged_transcript": "",  # 통합된 자막
            "youtube_analysis": "",  # 자막 분석 결과
            "youtube_step": 1,  # 유튜브 모드 내 단계 (1-4)
            # 강화된 도움 챗봇 관련 상태
            "show_help_chatbot": False,  # 도움 챗봇 표시 여부
            "show_contact_section": False,  # 연락 섹션 표시 여부
            "help_chat_messages": [],  # 도움 챗봇 대화 기록
        }
        for key, value in defaults.items():
            try:
                if key not in st.session_state:
                    st.session_state[key] = value
            except Exception:
                pass  # 개별 키 초기화 실패는 무시

        # 도움 챗봇 상태 초기화
        try:
            init_help_chatbot_state()
        except Exception:
            pass

        # 자동 저장 상태 초기화
        try:
            init_autosave_state()
        except Exception:
            pass

        # 성취 시스템 상태 초기화
        try:
            init_achievement_state()
        except Exception:
            pass

        # 모드 전환 관련 상태 초기화
        try:
            init_mode_transition_state()
        except Exception:
            # 모듈이 없을 경우 수동 초기화
            if "previous_mode" not in st.session_state:
                st.session_state.previous_mode = None
            if "mode_transition_data" not in st.session_state:
                st.session_state.mode_transition_data = {}
            if "last_error" not in st.session_state:
                st.session_state.last_error = None
            if "retry_count" not in st.session_state:
                st.session_state.retry_count = 0

        # 세션 상태 유효성 검사
        is_valid, issues = validate_session_state()
        if not is_valid and issues:
            # 자동 복구 시도 완료 (validate_session_state에서 처리)
            pass

        # 자동 저장 수행 (필요시)
        try:
            perform_autosave_if_needed()
        except Exception:
            pass  # 자동 저장 실패는 무시

    except Exception as e:
        # 치명적 오류 시 기본 초기화
        safe_session_init()
        st.warning("세션을 초기화하는 중 문제가 발생했어요. 기본값으로 시작합니다.")


def check_autosave_reminder():
    """자동 저장 알림 체크 (5분마다)"""
    if not st.session_state.drafts:
        return False

    last_save = st.session_state.get("last_save_time")
    if last_save is None:
        return True  # 한번도 저장 안 함

    try:
        last_save_dt = datetime.fromisoformat(last_save)
        time_diff = datetime.now() - last_save_dt
        return time_diff.total_seconds() > 300  # 5분 경과
    except:
        return True


def get_estimated_time_remaining():
    """남은 작업 예상 시간 계산"""
    parsed_toc = st.session_state.parsed_toc
    drafts = st.session_state.drafts

    if not parsed_toc:
        return None

    remaining = len(parsed_toc) - len(drafts)
    if remaining <= 0:
        return "완료!"

    # 장당 약 1.5분 (API 호출 + 검토 시간)
    estimated_minutes = remaining * 1.5

    if estimated_minutes < 60:
        return f"약 {int(estimated_minutes)}분"
    else:
        hours = int(estimated_minutes // 60)
        minutes = int(estimated_minutes % 60)
        return f"약 {hours}시간 {minutes}분"


def show_action_feedback(success=True, message="", duration=2):
    """표준화된 작업 피드백 표시"""
    if success:
        with st.container():
            st.success(f"✅ {message}")
        import time
        time.sleep(duration)
    else:
        st.error(f"❌ {message}")


def parse_toc(toc_text):
    """목차 텍스트를 파싱하여 구조화된 리스트로 변환"""
    sections = []
    current_part = None
    current_part_title = ""

    lines = toc_text.strip().split('\n')

    for line in lines:
        line = line.strip()
        if not line:
            continue

        # Part 감지 (다양한 형식 지원)
        part_match = re.match(r'^(?:Part|PART|파트|제)\s*(\d+)[.:]\s*(.+)', line, re.IGNORECASE)
        if part_match:
            current_part = int(part_match.group(1))
            current_part_title = part_match.group(2).strip()
            continue

        # 장 감지 (다양한 형식 지원)
        # 1-1, 1.1, 01., 1., 제1장 등
        section_patterns = [
            r'^(\d+)[-.:](\d+)[.:)]\s*(.+)',  # 1-1. 제목, 1.1 제목
            r'^(\d+)[.:)]\s*(.+)',  # 1. 제목
            r'^[-*•]\s*(.+)',  # - 제목, * 제목
            r'^제?(\d+)(?:장|절|화)?[.:)]\s*(.+)',  # 제1장 제목
        ]

        for pattern in section_patterns:
            match = re.match(pattern, line)
            if match:
                groups = match.groups()
                if len(groups) == 3:  # 1-1 형식
                    section_num = f"{groups[0]}-{groups[1]}"
                    section_title = groups[2].strip()
                elif len(groups) == 2:  # 1. 형식
                    section_num = groups[0]
                    section_title = groups[1].strip()
                else:  # 불릿 형식
                    section_num = str(len(sections) + 1)
                    section_title = groups[0].strip()

                sections.append({
                    "part": current_part or 1,
                    "part_title": current_part_title or f"Part {current_part or 1}",
                    "section_num": section_num,
                    "section_title": section_title,
                    "full_title": f"Part {current_part or 1} - {section_num}. {section_title}"
                })
                break

    # 파싱된 섹션이 없으면 줄 단위로 처리
    if not sections:
        for i, line in enumerate(lines):
            line = line.strip()
            if line and not line.startswith('#') and len(line) > 3:
                sections.append({
                    "part": 1,
                    "part_title": "Part 1",
                    "section_num": str(i + 1),
                    "section_title": line,
                    "full_title": f"{i + 1}. {line}"
                })

    return sections


def get_progress_stats():
    """진행 상황 통계 계산"""
    parsed_toc = st.session_state.parsed_toc
    drafts = st.session_state.drafts

    total_sections = len(parsed_toc)
    completed_sections = len(drafts)

    total_chars = sum(
        len(d.replace(" ", "").replace("\n", ""))
        for d in drafts.values()
    )

    target_chars = 60000  # 목표 6만자

    return {
        "total_sections": total_sections,
        "completed_sections": completed_sections,
        "progress_percent": (completed_sections / total_sections * 100) if total_sections > 0 else 0,
        "total_chars": total_chars,
        "target_chars": target_chars,
        "chars_percent": (total_chars / target_chars * 100) if target_chars > 0 else 0,
    }


def generate_quick_manuscript():
    """현재 상태로 빠른 원고 생성 (사이드바용)"""
    title = st.session_state.selected_title or "무제"
    author = st.session_state.book_info.get("name", "작성자")

    content = f"# {title}\n\n"
    content += f"**저자:** {author}\n\n---\n\n"

    if st.session_state.parsed_toc:
        current_part = None
        for section in st.session_state.parsed_toc:
            key = f"{section['section_num']}_{section['section_title']}"

            if section['part'] != current_part:
                current_part = section['part']
                content += f"\n# Part {current_part}. {section['part_title']}\n\n"

            if key in st.session_state.drafts:
                content += f"## {section['section_num']}. {section['section_title']}\n\n"
                content += f"{st.session_state.drafts[key]}\n\n---\n\n"

    return content


def render_progress_bar():
    """진행률 표시 바 (별 이모지 버전)"""
    stats = get_progress_stats()

    if stats["total_sections"] == 0:
        return

    # 별 이모지로 진행률 표시 (5단계)
    star_count = int(stats['progress_percent'] / 20)
    stars = "⭐" * star_count + "☆" * (5 - star_count)

    st.markdown(f"""
    <div class="progress-box" role="region" aria-label="진행률 정보" aria-live="polite">
        <div class="progress-text" aria-label="완료된 장 수">
            <span aria-hidden="true">📝</span> {stats['completed_sections']} / {stats['total_sections']} 장 완료!
        </div>
        <div class="star-progress" role="img" aria-label="진행률 {stats['progress_percent']:.0f}% - 별 {star_count}개 중 5개">
            {stars}
        </div>
        <div style="margin-top: 0.5rem; font-size: 1.1rem;">
            진행률: <span aria-live="polite">{stats['progress_percent']:.0f}%</span>
        </div>
        <div style="margin-top: 0.5rem; font-size: 1rem;">
            현재 <span aria-label="현재 글자 수">{stats['total_chars']:,}자</span> / 목표 <span aria-label="목표 글자 수">{stats['target_chars']:,}자</span>
        </div>
    </div>
    """, unsafe_allow_html=True)

    # Streamlit 프로그레스 바
    st.progress(stats['progress_percent'] / 100)

    # 마일스톤 메시지 (매 5장 완료 시)
    completed = stats['completed_sections']
    total = stats['total_sections']
    if completed > 0 and completed % 5 == 0 and f"milestone_{completed}" not in st.session_state.get("shown_milestones", []):
        if "shown_milestones" not in st.session_state:
            st.session_state.shown_milestones = []
        st.session_state.shown_milestones.append(f"milestone_{completed}")

        if completed == total:
            st.balloons()
            st.success("🎉 축하해! 모든 장을 완성했어!")
        elif completed >= total / 2:
            st.success(f"🌟 와! 벌써 절반이나 썼어! ({completed}장 완료)")
        else:
            st.success(f"✨ 잘하고 있어! {completed}장 완료!")


def navigate_to_step(step_num):
    """일관된 네비게이션 처리"""
    st.session_state.current_step = step_num
    st.rerun()


def render_sidebar():
    """사이드바 네비게이션 - 전문적 디자인"""
    with st.sidebar:
        # 브랜드 헤더
        st.markdown("""
        <div style="text-align: center; padding: 1rem 0.5rem; margin-bottom: 1rem; background: linear-gradient(135deg, #1e3a5f 0%, #2d5a87 100%); border-radius: 12px; border-bottom: 3px solid #c9a962;">
            <div style="font-size: 1.5rem; color: white; font-weight: 700;">📚 작가의집</div>
            <div style="font-size: 0.85rem; color: #c9a962; margin-top: 0.3rem;">AI 책쓰기 코칭</div>
        </div>
        """, unsafe_allow_html=True)

        # 채팅 모드 버튼 (상단에 눈에 띄게 배치)
        if st.session_state.chat_mode_active:
            if st.button("일반 모드로 돌아가기", use_container_width=True, type="secondary"):
                st.session_state.chat_mode_active = False
                st.rerun()
            st.markdown("""
            <div style="background: linear-gradient(135deg, #1e3a5f 0%, #2d5a87 100%); color: white; padding: 12px; border-radius: 10px; margin: 10px 0; text-align: center; font-weight: 600; border-left: 4px solid #c9a962;">
                코칭 대화 진행 중
            </div>
            """, unsafe_allow_html=True)
        else:
            if st.button("대화형 코칭 시작", use_container_width=True, type="primary"):
                st.session_state.chat_mode_active = True
                st.session_state.chat_mode_step = 0
                st.session_state.chat_mode_history = []
                st.session_state.chat_mode_data = {}
                st.rerun()
            st.caption("1:1 코칭처럼 대화로 책 만들기")

        st.markdown("---")

        # 유튜브 모드 버튼
        if st.session_state.get("youtube_mode_active", False):
            if st.button("일반 모드로 돌아가기", key="yt_exit", use_container_width=True, type="secondary"):
                st.session_state.youtube_mode_active = False
                st.rerun()
            st.markdown("""
            <div style="background: linear-gradient(135deg, #c05621 0%, #dd6b20 100%); color: white; padding: 12px; border-radius: 10px; margin: 10px 0; text-align: center; font-weight: 600;">
                유튜브 변환 모드
            </div>
            """, unsafe_allow_html=True)
        else:
            if st.button("유튜브 영상 변환", use_container_width=True, type="secondary"):
                st.session_state.youtube_mode_active = True
                st.session_state.youtube_step = 1
                st.session_state.youtube_urls = []
                st.session_state.youtube_videos = []
                st.session_state.youtube_transcripts = {}
                st.session_state.youtube_merged_transcript = ""
                st.session_state.youtube_analysis = ""
                st.rerun()
            st.caption("유튜브 영상을 책으로 변환")

        st.markdown("---")

        # 시각적 타임라인 (전문적 용어)
        steps = [
            ("1", "기본 정보", "책 콘셉트 정의"),
            ("2", "제목 생성", "AI 제목 추천"),
            ("3", "목차 구성", "5부 40장 구조"),
            ("4", "원고 작성", "챕터별 집필"),
            ("5", "출간기획서", "출판사 제출용"),
            ("6", "마케팅 자료", "랜딩페이지"),
            ("7", "다운로드", "최종 원고"),
        ]

        st.markdown("### 코칭 단계")

        # ARIA 라이브 리전 (스크린리더 지원)
        current_step_name = steps[st.session_state.current_step - 1][1] if st.session_state.current_step <= 7 else "완료"
        st.markdown(f'<div role="status" aria-live="polite" style="position:absolute;left:-9999px;">현재 단계: {current_step_name}</div>', unsafe_allow_html=True)

        for i, (icon, name, desc) in enumerate(steps, 1):
            if i == st.session_state.current_step:
                st.markdown(f"""
                <div style="background: linear-gradient(90deg, #ebf8ff 0%, rgba(235, 248, 255, 0.3) 100%); padding:10px 14px; border-radius:8px; margin:4px 0; border-left:4px solid #c9a962;">
                    <div style="display:flex; align-items:center; gap:8px;">
                        <span style="background:#1e3a5f; color:white; width:24px; height:24px; border-radius:50%; display:inline-flex; align-items:center; justify-content:center; font-size:0.8rem; font-weight:700;">{icon}</span>
                        <b style="color:#1e3a5f;">{name}</b>
                    </div>
                    <div style="font-size:0.8rem; color:#4a5568; margin-left:32px;">{desc}</div>
                </div>
                """, unsafe_allow_html=True)
            elif i < st.session_state.current_step:
                st.markdown(f"""
                <div style="padding:8px 14px; margin:4px 0;">
                    <div style="display:flex; align-items:center; gap:8px;">
                        <span style="background:#22543d; color:white; width:24px; height:24px; border-radius:50%; display:inline-flex; align-items:center; justify-content:center; font-size:0.8rem;">&#10003;</span>
                        <span style="color:#22543d; font-weight:500;">{name}</span>
                    </div>
                </div>
                """, unsafe_allow_html=True)
            else:
                st.markdown(f"""
                <div style="padding:8px 14px; margin:4px 0;">
                    <div style="display:flex; align-items:center; gap:8px;">
                        <span style="background:#e2e8f0; color:#718096; width:24px; height:24px; border-radius:50%; display:inline-flex; align-items:center; justify-content:center; font-size:0.8rem; font-weight:600;">{icon}</span>
                        <span style="color:#718096;">{name}</span>
                    </div>
                </div>
                """, unsafe_allow_html=True)

        st.markdown("---")

        # 진행률 미니 표시
        if st.session_state.parsed_toc:
            stats = get_progress_stats()
            st.markdown("### 집필 현황")
            # 프로그레스 바 스타일
            progress_color = "#c9a962" if stats['progress_percent'] < 100 else "#22543d"
            st.markdown(f"""
            <div style="margin-bottom: 0.5rem;">
                <span style="font-weight: 600; color: #1e3a5f;">{stats['completed_sections']}/{stats['total_sections']}</span>
                <span style="color: #718096;"> 장 완료</span>
            </div>
            """, unsafe_allow_html=True)
            st.markdown(f"**{stats['total_chars']:,}**자 / 60,000자")
            st.progress(stats['progress_percent'] / 100)

            # 예상 남은 시간
            est_time = get_estimated_time_remaining()
            if est_time:
                st.markdown(f"⏱️ 남은 시간: **{est_time}**")

            # 연속 작성 일수 표시
            streak = st.session_state.get("achievement_streak_days", 0)
            if streak > 1:
                st.markdown(f"🔥 **{streak}일** 연속 작성 중!")

            # 획득한 뱃지 수 표시
            earned_badges = st.session_state.get("achievement_earned_badges", [])
            if earned_badges:
                st.markdown(f"🏆 획득 뱃지: **{len(earned_badges)}개**")

        # ===== 자동 저장 섹션 =====
        st.markdown("### 💾 자동 저장")

        # 저장 상태 표시
        render_autosave_status()

        # 저장/불러오기 버튼
        render_save_buttons()

        # 저장된 작업 목록 표시 (활성화된 경우)
        if st.session_state.get("show_backup_list", False):
            render_backup_list()

        st.markdown("---")

        # 즉시 다운로드 (급한 사용자용)
        if st.session_state.drafts:
            st.markdown("### ⚡ 빠른 다운로드")
            stats = get_progress_stats()

            # 현재 상태로 즉시 다운로드
            quick_manuscript = generate_quick_manuscript()
            st.download_button(
                label=f"📥 현재 원고 ({stats['total_chars']:,}자)",
                data=quick_manuscript,
                file_name=f"원고_{datetime.now().strftime('%m%d_%H%M')}.md",
                mime="text/markdown",
                use_container_width=True,
                help="현재까지 작성된 내용을 즉시 다운로드"
            )
            st.caption(f"✅ {stats['completed_sections']}/{stats['total_sections']} 장 완료")

        st.markdown("---")

        # 단계 건너뛰기 (급한 사용자용)
        if st.session_state.current_step in [4, 5, 6] and st.session_state.drafts:
            st.markdown("### 🏃 빠른 진행")
            if st.session_state.current_step < 7:
                if st.button("⏭️ 바로 다운로드로", use_container_width=True, help="책 소개서/책 홍보 페이지를 건너뛰고 다운로드"):
                    st.session_state.current_step = 7
                    st.rerun()

        st.markdown("---")

        # 챗봇 토글 버튼
        if st.button("AI 코치와 대화", use_container_width=True):
            st.session_state.show_chatbot = not st.session_state.show_chatbot
            st.rerun()

        # 음성 모드 버튼
        voice_btn_label = "말로 쓰기 끝내기" if st.session_state.get("voice_mode_active", False) else "말로 책 쓰기"
        if st.button(voice_btn_label, use_container_width=True, help="타자 대신 말로 입력해요!"):
            st.session_state.voice_mode_active = not st.session_state.get("voice_mode_active", False)
            if not st.session_state.voice_mode_active:
                clear_voice_session()
            st.rerun()

        st.markdown("---")

        # 도움 & 연락 섹션 - 개선된 UI
        st.markdown("""
        <div style="background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
                    border-radius: 16px; padding: 1rem; margin: 0.5rem 0;
                    box-shadow: 0 4px 15px rgba(102, 126, 234, 0.3);">
            <h4 style="color: white; margin: 0 0 0.8rem 0; text-align: center; font-size: 1rem;">
                도움이 필요하신가요?
            </h4>
        </div>
        """, unsafe_allow_html=True)

        col_help1, col_help2 = st.columns(2)
        with col_help1:
            help_is_open = st.session_state.get("show_help_chatbot", False)
            help_btn_label = "닫기" if help_is_open else "AI 도움"
            help_btn_type = "secondary" if help_is_open else "primary"
            if st.button(help_btn_label, use_container_width=True, help="화면 중앙에 도움말 팝업이 열립니다", type=help_btn_type, key="sidebar_help_btn"):
                st.session_state.show_help_chatbot = not help_is_open
                st.rerun()

        with col_help2:
            contact_is_open = st.session_state.get("show_contact_section", False)
            contact_btn_label = "닫기" if contact_is_open else "선생님 연락"
            contact_btn_type = "secondary" if contact_is_open else "primary"
            if st.button(contact_btn_label, use_container_width=True, help="화면 중앙에 연락하기 팝업이 열립니다", type=contact_btn_type, key="sidebar_contact_btn"):
                st.session_state.show_contact_section = not contact_is_open
                st.rerun()

        # 설명 텍스트
        st.caption("버튼 클릭 시 화면 중앙에 팝업으로 표시됩니다")

        # 답변 대기 중인 질문이 있으면 알림
        try:
            pending_count = get_pending_messages_count()
            if pending_count > 0:
                st.markdown(f"""
                <div style="background: #FFF3E0; border-radius: 10px; padding: 0.8rem;
                            margin-top: 0.5rem; border-left: 4px solid #FF9800;">
                    <p style="margin: 0; font-size: 0.9rem; color: #E65100;">
                        <strong>내 질문 {pending_count}개</strong>가 답변을 기다리고 있어요!
                    </p>
                </div>
                """, unsafe_allow_html=True)
        except:
            pass


# 장르별 프리셋 데이터
GENRE_PRESETS = {
    "자기계발": {
        "topic": "직장인을 위한 시간 관리와 생산성 향상",
        "target_reader": "20-40대 직장인, 업무 효율화를 원하는 사람",
        "core_message": "하루 2시간의 집중 시간만으로도 인생을 바꿀 수 있다",
        "experience": "10년간 100개 이상의 프로젝트를 성공적으로 완수한 경험",
        "tone": "친절하고 따뜻한"
    },
    "재테크/투자": {
        "topic": "월급쟁이를 위한 부동산 투자 첫걸음",
        "target_reader": "30-50대 직장인, 재테크에 관심 있는 초보자",
        "core_message": "소액으로 시작해도 10년 안에 경제적 자유를 얻을 수 있다",
        "experience": "월급 300만원에서 시작해 자산 10억 달성",
        "tone": "전문가적이고 신뢰감 있는"
    },
    "육아/교육": {
        "topic": "아이의 자존감을 키우는 대화법",
        "target_reader": "3-12세 자녀를 둔 부모, 소통에 어려움을 느끼는 부모",
        "core_message": "매일 10분의 대화로 아이의 평생 자존감이 결정된다",
        "experience": "교육 전문가로 1,000가정 이상 상담 경험",
        "tone": "친절하고 따뜻한"
    },
    "건강/다이어트": {
        "topic": "50대부터 시작하는 건강한 체중 관리",
        "target_reader": "40-60대, 건강한 노후를 준비하는 사람",
        "core_message": "나이는 숫자일 뿐, 지금부터 관리하면 20년 젊게 살 수 있다",
        "experience": "영양학 박사, 1만명 이상의 체중 관리 성공 사례",
        "tone": "전문가적이고 신뢰감 있는"
    },
    "에세이/수필": {
        "topic": "평범한 일상에서 발견한 행복의 조각들",
        "target_reader": "20-50대, 삶의 의미를 찾고 싶은 사람",
        "core_message": "작은 것에 감사하면 큰 행복이 찾아온다",
        "experience": "10년간의 일기와 감사 기록",
        "tone": "감성적이고 공감가는"
    },
}


def render_welcome():
    """첫 방문 환영 화면 - 개선된 온보딩 UX"""

    # 온보딩 상태 초기화
    if "onboarding_completed" not in st.session_state:
        st.session_state.onboarding_completed = False
    if "show_onboarding" not in st.session_state:
        st.session_state.show_onboarding = True
    if "welcome_stage" not in st.session_state:
        st.session_state.welcome_stage = "user_type"  # user_type, mode_select, onboarding

    # === 환영 헤더 (앱 이름과 로고) - 전문적 브랜딩 ===
    st.markdown("""
    <div class="welcome-header">
        <div class="brand-badge">AI Book Coaching</div>
        <div class="logo">📚</div>
        <div class="app-name">작가의집</div>
        <div class="app-subtitle">책쓰기 코칭 전문</div>
        <div class="tagline">당신 안의 이야기를 세상 밖으로<br>10년 경력 출판 코치의 노하우와 AI가 함께합니다</div>
    </div>
    """, unsafe_allow_html=True)

    # 글씨 크기 조절 버튼 (접힌 상태)
    with st.expander("📏 글씨 크기 조절", expanded=False):
        col_s, col_n, col_l = st.columns(3)
        with col_s:
            if st.button("가 작게", key="font_small", use_container_width=True,
                        type="primary" if st.session_state.get("font_size") == "small" else "secondary"):
                st.session_state.font_size = "small"
                st.rerun()
        with col_n:
            if st.button("가 보통", key="font_normal", use_container_width=True,
                        type="primary" if st.session_state.get("font_size") == "normal" else "secondary"):
                st.session_state.font_size = "normal"
                st.rerun()
        with col_l:
            if st.button("가 크게", key="font_large", use_container_width=True,
                        type="primary" if st.session_state.get("font_size") == "large" else "secondary"):
                st.session_state.font_size = "large"
                st.rerun()

    # === 저장된 작업 확인 ===
    saved_files = get_autosave_files()
    has_saved_work = len(saved_files) > 0

    # === 1단계: 사용자 타입 선택 (처음/이어서) ===
    if st.session_state.welcome_stage == "user_type":
        st.markdown("---")
        st.markdown("""
        <div style="text-align: center; margin: 1.5rem 0;">
            <h2 style="color: #333; font-size: 1.8rem;">작가의집에 오신 것을 환영합니다!</h2>
            <p style="color: #666; font-size: 1.1rem;">10년 경력 출판 코치의 노하우로, 당신의 책 출간을 도와드립니다.</p>
        </div>
        """, unsafe_allow_html=True)

        col1, col2 = st.columns(2)

        with col1:
            st.markdown("""
            <div class="user-type-card new-user">
                <div class="card-emoji">🌱</div>
                <div class="card-title">처음 오셨나요?</div>
                <div class="card-desc">걱정 마세요! 체계적인 7단계 코칭으로<br>누구나 책을 완성할 수 있습니다.</div>
            </div>
            """, unsafe_allow_html=True)
            if st.button("🌱 새 프로젝트 시작", key="btn_new_user", use_container_width=True, type="primary"):
                st.session_state.welcome_stage = "onboarding" if st.session_state.show_onboarding else "mode_select"
                st.rerun()

        with col2:
            st.markdown(f"""
            <div class="user-type-card returning-user">
                <div class="card-emoji">📂</div>
                <div class="card-title">다시 오셨군요!</div>
                <div class="card-desc">{'저장된 프로젝트 ' + str(len(saved_files)) + '개가 기다리고 있어요.' if has_saved_work else '이전 작업을 불러오거나 새로 시작하세요.'}</div>
            </div>
            """, unsafe_allow_html=True)
            if st.button("📂 이어서 쓰기", key="btn_returning_user", use_container_width=True,
                        type="primary" if has_saved_work else "secondary"):
                if has_saved_work:
                    st.session_state.welcome_stage = "saved_works"
                else:
                    st.session_state.welcome_stage = "mode_select"
                st.rerun()

    # === 2단계: 온보딩 튜토리얼 (처음 사용자) ===
    elif st.session_state.welcome_stage == "onboarding":
        st.markdown("""
        <div class="onboarding-container">
            <div class="onboarding-title">🏠 작가의집 코칭 프로세스</div>
            <p style="text-align: center; color: #666; margin-bottom: 1.5rem;">
                수천 명의 작가 지망생이 이 방법으로 책을 완성했습니다.
            </p>

            <div class="onboarding-step">
                <div class="step-number">1</div>
                <div class="step-content">
                    <div class="step-title">코칭 상담: 당신의 이야기 발굴</div>
                    <div class="step-desc">전문 코치가 질문을 통해 책의 핵심 콘셉트를 함께 발굴합니다.<br><b>결과물:</b> 명확한 책 주제와 타겟 독자 정의</div>
                </div>
                <div class="step-icon">💬</div>
            </div>

            <div class="onboarding-step">
                <div class="step-number">2</div>
                <div class="step-content">
                    <div class="step-title">체계적 구조화: AI와 함께 완성</div>
                    <div class="step-desc">베스트셀러 공식을 적용한 제목, 5부 40장 목차, 챕터별 초안을 생성합니다.<br><b>결과물:</b> 출판 가능한 6만자 원고</div>
                </div>
                <div class="step-icon">✍️</div>
            </div>

            <div class="onboarding-step">
                <div class="step-number">3</div>
                <div class="step-content">
                    <div class="step-title">출판 준비: 완성 원고 & 기획서</div>
                    <div class="step-desc">출판사 투고용 기획서, 마케팅 랜딩페이지까지 한 번에 준비됩니다.<br><b>결과물:</b> 바로 제출 가능한 출간기획서</div>
                </div>
                <div class="step-icon">📥</div>
            </div>
        </div>
        """, unsafe_allow_html=True)

        # 다시 보지 않기 체크박스
        col_check, col_btn = st.columns([2, 1])
        with col_check:
            dont_show_again = st.checkbox("다시 보지 않기", key="dont_show_onboarding")
            if dont_show_again:
                st.session_state.show_onboarding = False
        with col_btn:
            if st.button("시작하기! 👉", key="btn_start_from_onboarding", use_container_width=True, type="primary"):
                st.session_state.onboarding_completed = True
                st.session_state.welcome_stage = "mode_select"
                st.rerun()

    # === 3단계: 저장된 작업 목록 (이어서 쓰기) ===
    elif st.session_state.welcome_stage == "saved_works":
        st.markdown("---")
        st.markdown("""
        <div class="saved-work-container">
            <div class="saved-work-title">📂 저장된 작업 목록</div>
        """, unsafe_allow_html=True)

        if saved_files:
            for idx, filepath in enumerate(saved_files[:5]):  # 최근 5개만 표시
                try:
                    with open(filepath, 'r', encoding='utf-8') as f:
                        data = json.load(f)

                    book_info = data.get("book_info", {})
                    title = book_info.get("topic", "제목 없음")[:30]
                    if len(book_info.get("topic", "")) > 30:
                        title += "..."
                    saved_at = data.get("saved_at", "알 수 없는 시간")
                    current_step = data.get("current_step", 1)

                    # 진행률 계산
                    drafts = data.get("drafts", {})
                    total_chars = sum(len(d.get("draft", "")) for d in drafts.values())
                    progress_text = f"Step {current_step}/7 | {total_chars:,}자 작성"

                    col_info, col_btn = st.columns([3, 1])
                    with col_info:
                        st.markdown(f"""
                        <div class="saved-work-item">
                            <div class="saved-work-info">
                                <div class="work-title">📄 {title if title else '새 프로젝트'}</div>
                                <div class="work-meta">저장: {saved_at}</div>
                                <div class="work-progress">{progress_text}</div>
                            </div>
                        </div>
                        """, unsafe_allow_html=True)
                    with col_btn:
                        if st.button("이어쓰기 ▶", key=f"continue_work_{idx}", use_container_width=True, type="primary"):
                            success, msg = restore_from_autosave(filepath)
                            if success:
                                st.session_state.welcome_stage = "mode_select"
                                st.success(f"'{title}' 프로젝트를 불러왔어요!")
                                st.rerun()
                            else:
                                st.error(f"불러오기 실패: {msg}")
                except Exception as e:
                    continue

        st.markdown("</div>", unsafe_allow_html=True)

        # 새로 시작하기 버튼
        st.markdown("---")
        col1, col2, col3 = st.columns([1, 2, 1])
        with col2:
            if st.button("🌱 새 프로젝트 시작하기", key="btn_new_project_from_saved", use_container_width=True):
                st.session_state.welcome_stage = "mode_select"
                st.rerun()

    # === 4단계: 모드 선택 ===
    elif st.session_state.welcome_stage == "mode_select":
        st.markdown("---")
        st.markdown("""
        <div style="text-align: center; margin: 1.5rem 0;">
            <h2 style="color: var(--primary-navy, #1e3a5f); font-size: 1.8rem; font-weight: 700;">코칭 방식을 선택하세요</h2>
            <p style="color: var(--text-secondary, #4a5568); font-size: 1.1rem;">어떤 방식을 선택하든 동일한 품질의 결과물을 얻을 수 있습니다.</p>
        </div>
        """, unsafe_allow_html=True)

        col1, col2, col3 = st.columns(3)

        with col1:
            st.markdown("""
            <div class="mode-card chat-card">
                <div class="emoji">💬</div>
                <div class="title">대화형 코칭</div>
                <div class="desc">코치와 1:1 상담하듯<br>질문에 답하며 책을 완성</div>
                <div class="feature-list">
                    <li>가장 추천하는 방식</li>
                    <li>단계별 맞춤 질문 제공</li>
                </div>
            </div>
            """, unsafe_allow_html=True)
            if st.button("💬 대화로 시작", key="mode_chat_welcome", use_container_width=True, type="primary"):
                st.session_state.chat_mode_active = True
                st.session_state.chat_mode_step = 0
                st.session_state.chat_mode_history = []
                st.session_state.chat_mode_data = {}
                st.rerun()

        with col2:
            st.markdown("""
            <div class="mode-card voice-card">
                <div class="emoji">🎤</div>
                <div class="title">음성 녹음</div>
                <div class="desc">말로 녹음하면<br>AI가 텍스트로 변환</div>
                <div class="feature-list">
                    <li>타이핑이 불편한 분께 추천</li>
                    <li>강연/인터뷰 녹음 활용 가능</li>
                </div>
            </div>
            """, unsafe_allow_html=True)
            if st.button("🎤 음성으로 시작", key="mode_voice_welcome", use_container_width=True, type="primary"):
                st.session_state.voice_mode_active = True
                st.rerun()

        with col3:
            st.markdown("""
            <div class="mode-card youtube-card">
                <div class="emoji">🎬</div>
                <div class="title">유튜브 변환</div>
                <div class="desc">기존 영상 콘텐츠를<br>책으로 재탄생</div>
                <div class="feature-list">
                    <li>유튜브 크리에이터에게 추천</li>
                    <li>여러 영상 통합 가능</li>
                </div>
            </div>
            """, unsafe_allow_html=True)
            if st.button("🎬 영상으로 시작", key="mode_youtube_welcome", use_container_width=True, type="primary"):
                st.session_state.youtube_mode_active = True
                st.session_state.youtube_step = 1
                st.rerun()

        st.markdown("---")

        # 고급 모드 (직접 쓰기)
        with st.expander("📝 직접 쓰기 (고급) - 7단계로 직접 진행하기", expanded=False):
            st.markdown("""
            **7단계로 직접 책을 완성해!**
            1. 정보 입력 - 책 주제와 누가 읽을까 정보
            2. 제목 만들기 - AI가 10가지 제목 추천
            3. 목차 만들기 - 40개 장 자동 구성
            4. 첫 번째 글 - 장별 1,500자 자동 생성
            5. 책 소개서 - 출판사 제출용 기획서
            6. 책 홍보 페이지 - 홍보용 페이지 카피
            7. 다운로드 - 완성된 원고 받기
            """)
            st.info("👇 아래에서 장르를 선택하면 예시가 자동으로 채워져!")

        # 빠른 시작 옵션 (급한 사용자용)
        st.markdown("### 빠른 시작 템플릿")
        st.markdown("장르별 템플릿을 선택하면 예시 데이터가 자동으로 채워집니다.")

        cols = st.columns(len(GENRE_PRESETS))
        for idx, (genre, preset) in enumerate(GENRE_PRESETS.items()):
            with cols[idx]:
                if st.button(f"📚 {genre}", key=f"preset_{genre}", use_container_width=True):
                    # 프리셋 데이터로 자동 채우기
                    st.session_state.book_info = {
                        "name": "홍길동",  # 기본 이름
                        "topic": preset["topic"],
                        "target_reader": preset["target_reader"],
                        "core_message": preset["core_message"],
                        "experience": preset["experience"],
                        "tone": preset["tone"],
                    }
                    st.session_state.selected_preset = genre
                    st.success(f"'{genre}' 프리셋이 적용되었습니다! 아래에서 이름만 수정하고 시작하세요.")
                    st.rerun()

        # 처음으로 돌아가기 버튼
        st.markdown("---")
        col1, col2, col3 = st.columns([1, 2, 1])
        with col2:
            if st.button("⬅️ 처음으로 돌아가기", key="btn_back_to_start", use_container_width=True):
                st.session_state.welcome_stage = "user_type"
                st.rerun()


def render_voice_mode():
    """음성 모드 UI 렌더링 - 타자가 어려운 사람을 위한 음성 입력"""
    st.markdown("""
    <div class="voice-mode-container">
        <div class="voice-mode-header">말로 책 쓰기</div>
        <div class="voice-mode-description">
            타자 치기 어려우면 말로 해도 돼요!<br>
            마이크로 녹음하거나 녹음 파일을 올리면<br>
            AI가 글자로 바꿔줘요!
        </div>
    </div>
    """, unsafe_allow_html=True)

    # 음성 입력 UI 렌더링 (voice_handler 모듈 사용)
    transcribed_text = render_voice_mode_ui()

    # 변환된 텍스트가 있으면 편집기 표시
    if transcribed_text:
        st.markdown("---")
        edited_text = render_transcription_editor(transcribed_text)

        if edited_text:
            st.session_state.voice_edited_text = edited_text

            st.markdown("---")
            st.markdown("### 다음에 뭘 할까요?")

            col1, col2 = st.columns(2)

            with col1:
                if st.button(
                    "이대로 책 만들기!",
                    type="primary",
                    use_container_width=True,
                    help="이 내용으로 바로 책을 만들어요!"
                ):
                    # 음성 텍스트를 책 정보에 저장
                    st.session_state.book_info["topic"] = edited_text[:500] if len(edited_text) > 500 else edited_text
                    st.session_state.book_info["core_message"] = edited_text[:200] if len(edited_text) > 200 else edited_text
                    st.session_state.book_info["experience"] = edited_text

                    # 음성 모드 종료
                    st.session_state.voice_mode_active = False
                    st.session_state.previous_mode = "voice"

                    # 이름이 있으면 Step 2로, 없으면 Step 1로 이동
                    if st.session_state.book_info.get("name"):
                        st.session_state.current_step = 2
                    else:
                        st.session_state.current_step = 1

                    st.success("좋아요! 저장했어요! 다음 단계로 가볼까요?")
                    st.rerun()

            with col2:
                if st.button(
                    "글자만 가져가기",
                    use_container_width=True,
                    help="글자로 바꾼 내용을 복사해서 직접 입력할게요"
                ):
                    # 음성 모드 종료하고 1단계로 이동 (텍스트는 세션에 유지)
                    st.session_state.voice_mode_active = False
                    st.session_state.previous_mode = "voice"
                    st.session_state.current_step = 1
                    st.info("위의 글자를 복사해서 사용해 보세요!")
                    st.rerun()

    # 음성 모드 종료 버튼
    st.markdown("---")
    if st.button("음성 모드 끝내기", use_container_width=True):
        clear_voice_session()
        st.session_state.voice_mode_active = False
        st.rerun()


def render_step1():
    """1단계: 컨설팅 정보 입력"""
    # 첫 방문이면 환영 메시지 표시
    if not st.session_state.book_info:
        render_welcome()
        st.markdown("---")

    # UX 개선: 현재 위치 브레드크럼 표시
    render_ux_breadcrumb(1)

    st.markdown('<h2 class="step-header" id="step-1-header" tabindex="-1">1단계: 책의 방향 설정</h2>', unsafe_allow_html=True)

    # UX 개선: 단계별 팁 표시
    render_ux_step_tip(1)

    st.markdown("""
    <div class="help-box" role="note" aria-label="도움말">
    <span aria-hidden="true">💡</span> <b>코치의 안내</b>
    <br><br>
    좋은 책은 명확한 방향에서 시작됩니다. 아래 4가지 질문에 답해주세요.<br>
    완벽하지 않아도 괜찮습니다. 작성 중 언제든 수정할 수 있습니다.
    <br><br>
    <b>각 항목의 역할:</b>
    <ul style="margin: 0.5rem 0; font-size: 1rem;">
    <li><b>저자명</b> → 책 표지와 출간기획서에 표시됩니다</li>
    <li><b>책 주제</b> → 제목 생성과 목차 구성의 핵심 자료입니다</li>
    <li><b>타겟 독자</b> → 글의 난이도와 톤을 결정합니다</li>
    <li><b>핵심 메시지</b> → 책 전체를 관통하는 한 줄 메시지입니다</li>
    </ul>
    </div>
    """, unsafe_allow_html=True)

    # 개선점 1: 입력 템플릿/예시 제공 (초보자 지원)
    with st.expander("📋 베스트셀러 작가들의 예시 참고하기", expanded=False):
        st.markdown("""
        <div class="hint-box">
        <b>예시 1: 자기계발 (베스트셀러 유형)</b><br>
        「부의 추월차선」, 「돈의 속성」 스타일
        <ul style="margin: 0.5rem 0;">
        <li><b>저자:</b> 김재무 (10년차 자산운용사 애널리스트)</li>
        <li><b>책 주제:</b> 월급쟁이가 10년 안에 경제적 자유를 얻는 현실적 로드맵</li>
        <li><b>타겟 독자:</b> 30-45세 직장인, 연봉 5천만원 이상, 재테크를 시작하고 싶지만 방법을 모르는 사람</li>
        <li><b>핵심 메시지:</b> 월급의 30%만 10년간 올바르게 투자하면 누구나 10억을 모을 수 있다</li>
        </ul>
        </div>

        <div class="hint-box">
        <b>예시 2: 실용서 (문제 해결형)</b><br>
        「아침형 인간」, 「하버드 새벽 4시 반」 스타일
        <ul style="margin: 0.5rem 0;">
        <li><b>저자:</b> 박미라 (워킹맘 10년차, 시간관리 코치)</li>
        <li><b>책 주제:</b> 퇴근 후 2시간으로 인생을 바꾸는 워킹맘 시간 혁명</li>
        <li><b>타겟 독자:</b> 30-40대 워킹맘, 육아와 일의 균형을 찾고 싶은 여성, 자기계발 시간이 부족하다고 느끼는 직장인</li>
        <li><b>핵심 메시지:</b> 하루 2시간 루틴만 바꾸면 1년 후 완전히 다른 삶을 살 수 있다</li>
        </ul>
        </div>

        <div class="hint-box">
        <b>예시 3: 비즈니스/트렌드 (시의성 있는 주제)</b><br>
        「챗GPT 업무 혁명」, 「AI 시대 생존법」 스타일
        <ul style="margin: 0.5rem 0;">
        <li><b>저자:</b> 이정훈 (IT기업 팀장, AI 도입 프로젝트 리더)</li>
        <li><b>책 주제:</b> AI를 내 편으로 만드는 직장인 생존 전략 50</li>
        <li><b>타겟 독자:</b> 25-50대 사무직 직장인, AI에 대체될까 불안한 사람, 업무 효율화에 관심 있는 팀장급</li>
        <li><b>핵심 메시지:</b> AI는 당신의 경쟁자가 아니라 24시간 일하는 비서다</li>
        </ul>
        </div>

        <div class="hint-box">
        <b>예시 4: 에세이/자기 이야기</b><br>
        「나는 나로 살기로 했다」, 「죽고 싶지만 떡볶이는 먹고 싶어」 스타일
        <ul style="margin: 0.5rem 0;">
        <li><b>저자:</b> 최은영 (10년차 직장인, 번아웃 극복 경험자)</li>
        <li><b>책 주제:</b> 서른다섯, 회사를 그만두고 나를 다시 만났다</li>
        <li><b>타겟 독자:</b> 30-40대 직장인, 번아웃을 경험했거나 현재 겪고 있는 사람, 퇴사를 고민하는 사람</li>
        <li><b>핵심 메시지:</b> 멈춤은 끝이 아니라 진짜 나를 찾는 시작이다</li>
        </ul>
        </div>
        """, unsafe_allow_html=True)

    with st.form("book_info_form"):
        col1, col2 = st.columns(2)

        with col1:
            name = st.text_input(
                "저자명 *",
                value=st.session_state.book_info.get("name", ""),
                placeholder="책 표지에 표시될 이름 (예: 홍길동, 필명 가능)",
                help="책 표지, 출간기획서, 저자 소개에 사용됩니다. 본명 또는 필명을 입력하세요."
            )
            topic = st.text_area(
                "책 주제 *",
                value=st.session_state.book_info.get("topic", ""),
                placeholder="어떤 내용의 책을 쓰고 싶으신가요?\n\n예시:\n- 퇴직 후 제2의 인생을 시작하는 50대를 위한 창업 가이드\n- 육아와 커리어를 동시에 잡는 워킹맘 생존 전략\n- 주식 초보가 1년 만에 수익률 30% 달성한 실전 노하우",
                height=120,
                help="구체적일수록 좋습니다. 제목 생성과 목차 구성에 핵심 자료로 활용됩니다."
            )
            target_reader = st.text_area(
                "타겟 독자 *",
                value=st.session_state.book_info.get("target_reader", ""),
                placeholder="이 책을 읽어야 하는 사람은 누구인가요?\n\n예시:\n- 30-40대 직장인 중 재테크를 시작하고 싶은 사람\n- 자녀 교육에 고민이 많은 초등학생 학부모\n- 이직을 준비하는 5년차 이상 개발자",
                height=120,
                help="나이, 직업, 상황, 고민을 구체적으로 적어주세요. 글의 난이도와 톤이 결정됩니다."
            )

        with col2:
            core_message = st.text_area(
                "핵심 메시지 * (책을 한 줄로 요약하면?)",
                value=st.session_state.book_info.get("core_message", ""),
                placeholder="독자가 이 책을 읽고 얻어갈 핵심 가치는?\n\n예시:\n- 누구나 3년이면 경제적 자유를 시작할 수 있다\n- 하루 30분 루틴이 인생을 바꾼다\n- AI를 두려워하지 말고 활용하라",
                height=120,
                help="책 전체를 관통하는 한 줄 메시지입니다. 독자가 기억할 핵심 가치를 적어주세요."
            )
            experience = st.text_area(
                "저자 경험/전문성 (선택)",
                value=st.session_state.book_info.get("experience", ""),
                placeholder="이 주제에 대해 쓸 수 있는 자격은?\n\n예시:\n- 10년간 부동산 투자로 자산 10억 달성\n- 500명 이상의 학생을 가르친 영어 강사\n- 3개 스타트업 창업 및 매각 경험",
                height=120,
                help="없어도 괜찮습니다. 있다면 출간기획서에서 신뢰도를 높이는 데 활용됩니다."
            )
            tone = st.selectbox(
                "글쓰기 톤",
                options=WRITING_TONES,
                index=WRITING_TONES.index(st.session_state.book_info.get("tone", WRITING_TONES[0]))
                if st.session_state.book_info.get("tone") in WRITING_TONES else 0,
                help="책의 전체적인 분위기를 결정합니다. 타겟 독자에게 맞는 톤을 선택하세요."
            )

        submitted = st.form_submit_button(
            "✅ 저장하고 다음으로 →",
            use_container_width=True,
            type="primary"
        )

        if submitted:
            if not all([name, topic, target_reader, core_message]):
                st.error("😊 여기도 채워줘! 이름, 책 주제, 누가 읽을까, 가장 중요한 이야기를 모두 적어줘!")
            else:
                # 시각적 피드백: 진행 표시
                progress_placeholder = st.empty()
                progress_placeholder.info("💾 정보를 저장하고 있습니다...")

                st.session_state.book_info = {
                    "name": name,
                    "topic": topic,
                    "target_reader": target_reader,
                    "core_message": core_message,
                    "experience": experience,
                    "tone": tone,
                }
                st.session_state.current_step = 2

                # 명확한 성공 피드백
                progress_placeholder.success("✅ 저장되었습니다!")
                st.markdown("""
                <div style="background: #C8E6C9; padding: 1rem; border-radius: 8px; margin-top: 1rem; border-left: 4px solid #4CAF50;">
                <b>다음 단계:</b> 이제 Step 2로 진행하여 AI가 추천하는 10가지 책 제목을 선택하세요!
                </div>
                """, unsafe_allow_html=True)
                import time
                time.sleep(2)
                st.rerun()


def render_step2():
    """2단계: 제목 생성"""
    # UX 개선: 현재 위치 브레드크럼 표시
    render_ux_breadcrumb(2)

    st.markdown('<h2 class="step-header" id="step-2-header" tabindex="-1">2단계: 책 제목 결정</h2>', unsafe_allow_html=True)

    # UX 개선: 단계별 팁 표시
    render_ux_step_tip(2)

    st.markdown("""
    <div class="help-box">
    <b>코치의 조언</b>
    <br><br>
    제목은 책의 운명을 결정합니다. 서점에서 독자가 책을 집어드는 첫 번째 이유가 바로 제목입니다.
    <br><br>
    AI가 베스트셀러 제목 공식 10가지를 적용해 제목을 추천합니다.<br>
    마음에 드는 것을 선택하거나, 여러 제목의 장점을 조합해 새로운 제목을 만들어보세요.
    <br><br>
    <b>좋은 제목의 특징:</b> 호기심 유발, 구체적 숫자, 타겟 명확, 10자 이내 간결함
    </div>
    """, unsafe_allow_html=True)

    # book_info 유효성 검사
    book_info = st.session_state.get("book_info", {})
    required_fields = ["topic", "target_reader", "core_message"]
    missing_fields = [f for f in required_fields if not book_info.get(f)]

    if missing_fields:
        field_names = {"topic": "책 주제", "target_reader": "타겟 독자", "core_message": "핵심 메시지"}
        missing_names = [field_names.get(f, f) for f in missing_fields]
        st.warning(f"""
        **아직 입력하지 않은 정보가 있습니다.**

        다음 정보를 1단계에서 먼저 입력해주세요: **{', '.join(missing_names)}**
        """)
        if st.button("← 1단계로 돌아가기", type="primary"):
            st.session_state.current_step = 1
            st.rerun()
        return

    col1, col2 = st.columns([2, 1])

    with col1:
        # 제목 생성 버튼
        if st.button("🎯 제목 10개 생성하기", use_container_width=True, type="primary"):
            with st.spinner("✨ AI가 제목을 생성하고 있습니다..."):
                try:
                    result = generate_titles(st.session_state.book_info)
                    if result:
                        st.session_state.generated_titles = result
                        st.rerun()
                    else:
                        st.error("""
                        **제목 생성에 실패했습니다.**

                        - 인터넷 연결을 확인해주세요.
                        - 잠시 후 다시 시도해보세요.
                        """)
                except Exception as e:
                    st.error("""
                    **제목 생성 중 오류가 발생했습니다.**

                    잠시 후 다시 시도해주세요.
                    """)
                    with st.expander("기술 정보", expanded=False):
                        st.code(str(e)[:300])

        # 생성된 제목 표시
        if st.session_state.generated_titles:
            st.markdown("### 📝 생성된 제목 후보")
            st.markdown(st.session_state.generated_titles)

            st.markdown("---")

            # 제목 선택/수정
            selected = st.text_input(
                "✏️ 최종 제목 (위에서 복사하거나 직접 작성)",
                value=st.session_state.selected_title,
                placeholder="최종 제목을 입력하세요"
            )

            if selected:
                st.session_state.selected_title = selected
                st.session_state.book_info["title"] = selected
                # 제목 선택 시 자동 저장 트리거
                trigger_important_save("title_selected")

    with col2:
        st.markdown("### 베스트셀러 제목 분석")
        st.info("""
        **검증된 제목 공식:**

        **숫자형** - 구체적 결과 제시
        - 「1만 시간의 법칙」
        - 「하루 5분 기적의 습관」

        **질문형** - 독자 고민 자극
        - 「왜 나는 항상 돈이 없을까?」

        **변화형** - Before/After 제시
        - 「월급쟁이에서 건물주로」

        **비밀형** - 희소성 강조
        - 「부자들만 아는 세금의 비밀」

        **행동형** - 즉시 실행 유도
        - 「지금 당장 시작하라」
        """)

        if st.session_state.generated_titles:
            if st.button("🔄 다시 생성하기"):
                with st.spinner("✨ 다시 생성 중..."):
                    result = generate_titles(st.session_state.book_info)
                    if result:
                        st.session_state.generated_titles = result
                        st.rerun()

    # 네비게이션
    st.markdown("---")
    col1, col2 = st.columns(2)
    with col1:
        if st.button("← 이전"):
            navigate_to_step(1)
    with col2:
        if st.session_state.selected_title:
            if st.button("다음 →", use_container_width=True, type="primary"):
                navigate_to_step(3)
        else:
            st.warning("👆 위에서 제목을 골라줘!")


def render_step3():
    """3단계: 목차 생성"""
    # UX 개선: 현재 위치 브레드크럼 표시
    render_ux_breadcrumb(3)

    st.markdown('<h2 class="step-header" id="step-3-header" tabindex="-1">3단계: 목차 구성</h2>', unsafe_allow_html=True)
    st.markdown(f"**선택된 제목:** {st.session_state.selected_title}")

    # UX 개선: 단계별 팁 표시
    render_ux_step_tip(3)

    st.markdown("""
    <div class="help-box">
    <b>코치의 조언</b>
    <br><br>
    목차는 책의 설계도입니다. 잘 짜인 목차가 있으면 글쓰기가 훨씬 수월해집니다.
    <br><br>
    AI가 검증된 <b>5부 40장 구조</b>로 목차를 생성합니다:
    <ul style="margin: 0.5rem 0;">
    <li><b>Part 1. WHY</b> - 왜 이 책을 읽어야 하는지 (동기 부여)</li>
    <li><b>Part 2. WHAT</b> - 무엇을 알아야 하는지 (핵심 개념)</li>
    <li><b>Part 3. HOW</b> - 어떻게 해야 하는지 (실천 방법)</li>
    <li><b>Part 4. DO</b> - 실제로 해보기 (사례와 실습)</li>
    <li><b>Part 5. FUTURE</b> - 앞으로의 여정 (비전과 동기부여)</li>
    </ul>
    각 장 1,500자 x 40장 = <b>총 6만자 분량</b>의 책이 완성됩니다.
    </div>
    """, unsafe_allow_html=True)

    col1, col2 = st.columns([2, 1])

    with col1:
        # 목차 생성 버튼
        if st.button("📋 목차 만들어줘! (책의 5가지 부분, 40장)", use_container_width=True, type="primary"):
            with st.spinner("✨ AI가 목차를 만들고 있어... (약 1분 걸려)"):
                result = generate_toc(st.session_state.book_info)
                if result:
                    st.session_state.generated_toc = result
                    st.session_state.parsed_toc = parse_toc(result)
                    # 목차 생성 시 자동 저장 트리거
                    trigger_important_save("toc_generated")
                    st.rerun()
                else:
                    st.error("""
                    **앗! 잠깐 문제가 생겼어. 다시 해볼까?**

                    - 30초 후 다시 '목차 만들어줘!' 버튼을 눌러봐
                    - 인터넷 연결을 확인해봐
                    """)

        # 생성된 목차 표시
        if st.session_state.generated_toc:
            st.markdown("### 📚 만들어진 목차")

            # 파싱 결과 표시
            if st.session_state.parsed_toc:
                st.success(f"✅ {len(st.session_state.parsed_toc)}개 장이 만들어졌어!")

                # 그룹별로 표시
                current_part = None
                for section in st.session_state.parsed_toc:
                    if section["part"] != current_part:
                        current_part = section["part"]
                        st.markdown(f"**Part {current_part}. {section['part_title']}**")
                    st.markdown(f"  - {section['section_num']}. {section['section_title']}")

            # 편집 가능한 텍스트 영역
            with st.expander("📝 목차 직접 수정 (고급)", expanded=False):
                edited_toc = st.text_area(
                    "목차 수정",
                    value=st.session_state.generated_toc,
                    height=400,
                    label_visibility="collapsed"
                )
                if edited_toc != st.session_state.generated_toc:
                    if st.button("🔄 수정 적용"):
                        st.session_state.generated_toc = edited_toc
                        st.session_state.parsed_toc = parse_toc(edited_toc)
                        st.rerun()

    with col2:
        st.markdown("### 📐 책의 5가지 부분 설명")
        st.info("""
        **Part 1. 왜? (WHY)**
        → 왜 이게 중요해?

        **Part 2. 뭐? (WHAT)**
        → 무엇을 알아야 해?

        **Part 3. 어떻게? (HOW)**
        → 어떻게 하면 돼?

        **Part 4. 해보자! (DO)**
        → 실제로 해보기

        **Part 5. 앞으로! (FUTURE)**
        → 앞으로 어떻게 될까?
        """)

        if st.session_state.generated_toc:
            if st.button("🔄 목차 다시 만들기"):
                with st.spinner("✨ 다시 만드는 중..."):
                    result = generate_toc(st.session_state.book_info)
                    if result:
                        st.session_state.generated_toc = result
                        st.session_state.parsed_toc = parse_toc(result)
                        st.rerun()

    # 네비게이션
    st.markdown("---")
    col1, col2 = st.columns(2)
    with col1:
        if st.button("← 이전"):
            navigate_to_step(2)
    with col2:
        if st.session_state.parsed_toc:
            if st.button("다음: 첫 번째 글 쓰기 시작! →", use_container_width=True, type="primary"):
                navigate_to_step(4)
        else:
            st.warning("👆 먼저 3단계에서 목차를 만들어줘!")


def get_section_key(section):
    """섹션 키 생성 (단순화)"""
    return f"{section['section_num']}_{section['section_title']}"


def get_part_transcript(full_transcript: str, part_number: int) -> str:
    """
    다중 영상 자막에서 특정 Part의 자막 추출

    Args:
        full_transcript: 전체 통합 자막 (=== Part N: 제목 === 형식)
        part_number: 추출할 Part 번호

    Returns:
        해당 Part의 자막 또는 전체 자막의 일부
    """
    import re

    if not full_transcript:
        return ""

    # Part 구분자 패턴
    part_pattern = r'=== Part (\d+):'

    # 모든 Part 위치 찾기
    matches = list(re.finditer(part_pattern, full_transcript))

    if not matches:
        # Part 구분이 없으면 전체 자막을 Part 개수로 균등 분할
        total_len = len(full_transcript)
        chunk_size = total_len // 5  # 5부 기준
        start = (part_number - 1) * chunk_size
        end = min(start + chunk_size * 2, total_len)  # 여유 있게 추출
        return full_transcript[start:end]

    # 해당 Part 찾기
    for i, match in enumerate(matches):
        current_part = int(match.group(1))
        if current_part == part_number:
            start = match.start()
            # 다음 Part 시작점 또는 끝
            if i + 1 < len(matches):
                end = matches[i + 1].start()
            else:
                end = len(full_transcript)
            return full_transcript[start:end]

    # 해당 Part를 찾지 못하면 Part 번호 기반으로 추정
    total_parts = len(matches)
    if part_number <= total_parts:
        idx = part_number - 1
        start = matches[idx].start()
        if idx + 1 < len(matches):
            end = matches[idx + 1].start()
        else:
            end = len(full_transcript)
        return full_transcript[start:end]

    # 그래도 없으면 전체 자막의 일부 반환
    return full_transcript[:8000]


def is_section_completed(section, drafts):
    """섹션 완료 여부 확인"""
    return get_section_key(section) in drafts


def get_section_status(section, drafts, current_idx, section_idx):
    """섹션 상태 정보 반환 (prefix, help_text)"""
    if section_idx == current_idx:
        return "➡️ (현재)", "지금 쓰고 있는 장이에요"
    elif is_section_completed(section, drafts):
        return "✅ (완료)", "이 장은 다 썼어요!"
    else:
        return "⬜ (미작성)", "아직 안 쓴 장이에요"


def is_part_completed(part_num, parsed_toc, drafts):
    """특정 Part의 모든 장 완료 여부"""
    part_sections = [s for s in parsed_toc if s['part'] == part_num]
    return all(is_section_completed(s, drafts) for s in part_sections)


def render_step4():
    """4단계: 첫 번째 글 생성 - 순차적 플로우"""
    # UX 개선: 현재 위치 브레드크럼 표시
    render_ux_breadcrumb(4)

    st.markdown('<h2 class="step-header" id="step-4-header" tabindex="-1">4단계: 첫 번째 글 쓰기</h2>', unsafe_allow_html=True)

    # UX 개선: 단계별 팁 표시
    render_ux_step_tip(4)

    # 성취 시스템 - 강화된 진행률 헤더
    render_progress_header()

    # 뱃지 및 마일스톤 팝업 표시
    render_badge_popup()
    render_milestone_popup()

    # 오늘의 목표 표시
    render_daily_goal_section()

    # 획득한 뱃지 표시
    render_badges_display()

    parsed_toc = st.session_state.parsed_toc
    drafts = st.session_state.drafts

    # 현재 상태에 따른 명확한 안내 메시지
    completed_count = len(drafts)
    total_count = len(parsed_toc)

    if total_count > 0:
        if completed_count == 0:
            st.markdown("""
            <div class="help-box">
            👋 <b>시작해볼까?</b> 아래 '✨ AI로 글 만들기' 버튼 하나만 누르면 돼!
            </div>
            """, unsafe_allow_html=True)
        elif completed_count < total_count:
            remaining = total_count - completed_count
            st.markdown(f"""
            <div class="help-box">
            💪 <b>잘하고 있어!</b> {completed_count}장 완료! 남은 장: {remaining}개
            </div>
            """, unsafe_allow_html=True)
        else:
            st.markdown("""
            <div class="help-box" style="background: #E8F5E9; border-color: #4CAF50;">
            🎉 <b>모든 첫 번째 글 완료!</b> 아래에서 다듬거나 다음으로 넘어가!
            </div>
            """, unsafe_allow_html=True)

    if not parsed_toc:
        st.warning("👆 먼저 3단계에서 목차를 만들어줘!")
        if st.button("← 목차 만들러 가기"):
            navigate_to_step(3)
        return

    # 현재 섹션 인덱스 유효성 확인 및 조정
    current_idx = st.session_state.current_section_index
    if current_idx >= len(parsed_toc):
        current_idx = 0
        st.session_state.current_section_index = 0

    current_section = parsed_toc[current_idx]
    section_key = get_section_key(current_section)

    col1, col2 = st.columns([2, 1])

    with col1:
        # 현재 작성할 장 표시
        st.markdown(f"""
        <div class="current-section-box">
        <h3>✍️ 지금 쓸 장</h3>
        <p><b>Part {current_section['part']}.</b> {current_section['part_title']}</p>
        <p style="font-size: 1.3rem;"><b>{current_section['section_num']}. {current_section['section_title']}</b></p>
        </div>
        """, unsafe_allow_html=True)

        # 이미 작성된 경우
        if section_key in drafts:
            st.success("✅ 이 장은 이미 썼어!")

            # 수정 가능 (다듬기)
            edited_draft = st.text_area(
                "쓴 내용 (다듬기 가능)",
                value=drafts[section_key],
                height=400
            )
            if edited_draft != drafts[section_key]:
                drafts[section_key] = edited_draft

            # UX 개선: 향상된 글자 수 카운터
            render_ux_char_counter(edited_draft, target_chars=1500)

            col_a, col_b = st.columns(2)
            with col_a:
                if current_idx > 0:
                    if st.button("⬅️ 이전 장"):
                        st.session_state.current_section_index = current_idx - 1
                        st.rerun()
            with col_b:
                if current_idx < len(parsed_toc) - 1:
                    if st.button("➡️ 다음 장", type="primary"):
                        st.session_state.current_section_index = current_idx + 1
                        st.rerun()
        else:
            # 첫 번째 글 생성
            st.markdown("### 🚀 첫 번째 글 만들기")

            # 추가 정보 입력 (선택)
            with st.expander("💡 추가 정보 넣기 (안 해도 돼!)", expanded=False):
                st.caption("아래 정보를 넣으면 AI가 더 잘 써줘!")
                section_message = st.text_area(
                    "이 장에서 가장 하고 싶은 말",
                    placeholder="예: 완벽하지 않아도 괜찮아!",
                    height=80,
                    help="이 장에서 가장 중요한 말 한 줄!"
                )
                section_examples = st.text_area(
                    "넣고 싶은 예시나 이야기",
                    placeholder="예: 내 친구 이야기",
                    height=80,
                    help="재미있는 예시가 있으면 적어봐!"
                )

            # 생성 버튼
            if st.button("✨ AI로 글 만들기 (약 1,500자)", use_container_width=True, type="primary"):
                section_info = {
                    "part_number": current_section["part"],
                    "part_title": current_section["part_title"],
                    "section_number": current_section["section_num"],
                    "section_title": current_section["section_title"],
                    "core_message": section_message if 'section_message' in dir() else "",
                    "examples": section_examples if 'section_examples' in dir() else "",
                }

                # UX 개선: 향상된 로딩 애니메이션
                loading_placeholder = st.empty()
                with loading_placeholder.container():
                    render_ux_loading(30, "글 만들기", show_tips=True)

                result = generate_draft(st.session_state.book_info, section_info)

                # 로딩 애니메이션 제거
                loading_placeholder.empty()

                if result:
                    st.session_state.drafts[section_key] = result

                    # 초안 완성 시 자동 저장 트리거
                    trigger_important_save("draft_completed")

                    # 성취 시스템 호출 - 장 완료 처리
                    on_chapter_complete()

                    # 마일스톤 성취감 피드백
                    new_completed = len(st.session_state.drafts)
                    total = len(parsed_toc)

                    # 동기부여 메시지 가져오기
                    motivation = get_motivation_by_progress()

                    # UX 개선: 마일스톤별 피드백 메시지 (축하 효과 강화)
                    if new_completed == total:
                        st.balloons()
                        st.snow()
                        render_ux_celebration(
                            "모든 첫 번째 글 완성!",
                            "축하합니다! 이제 책 한 권 분량의 원고가 완성되었어요!",
                            ["🎆", "👑", "🎉"]
                        )
                        # UX 개선: 자동 스크롤
                        inject_auto_scroll_script("ux-celebration-box")
                    elif new_completed in [5, 10, 15, 20, 25, 30, 35]:
                        st.balloons()
                        milestone = MILESTONE_MESSAGES.get(new_completed, {})
                        if milestone:
                            render_ux_celebration(
                                milestone.get('title', f'{new_completed}장 완료!'),
                                milestone.get('message', ''),
                                [milestone.get('emoji', '🎉'), "✨", "🌟"]
                            )
                    elif is_part_completed(current_section['part'], parsed_toc, st.session_state.drafts):
                        st.balloons()
                        render_ux_celebration(
                            f"Part {current_section['part']} 완료!",
                            "한 파트를 모두 완료했어요! 대단해요!",
                            ["🎉", "📖", "✨"]
                        )
                    else:
                        # UX 개선: 토스트 스타일 성공 메시지
                        render_ux_toast(f"글 완성! ({new_completed}/{total})", "success")
                        st.success(f"✅ 글 완성! ({new_completed}/{total})\n\n{motivation}")
                else:
                    # UX 개선: 토스트 스타일 에러 메시지
                    render_ux_toast("앗! 잠깐 문제가 생겼어요.", "error")
                    st.session_state['last_failed_section'] = section_key
                    st.error("""
                    **앗! 잠깐 문제가 생겼어. 다시 해볼까?**

                    **이렇게 해봐:**
                    1. 🔄 잠깐(30초) 기다렸다가 다시 버튼을 눌러봐
                    2. 인터넷 연결을 확인해봐

                    💡 **팁:** '직접 쓰기'에 먼저 글을 쓰고, 나중에 AI한테 도움 받을 수도 있어!
                    """)

                st.rerun()

            # 직접 작성 옵션
            st.markdown("---")
            st.markdown("**또는 직접 쓰기:**")
            manual_draft = st.text_area(
                "직접 글 쓰기",
                height=300,
                placeholder="여기에 직접 써도 돼...",
                label_visibility="collapsed"
            )
            if manual_draft:
                if st.button("💾 저장하기", use_container_width=True, type="primary"):
                    st.session_state.drafts[section_key] = manual_draft
                    # 수동 저장 시 자동 저장 트리거
                    trigger_important_save("manual_draft_saved")
                    # 성취 시스템 호출 - 장 완료 처리
                    on_chapter_complete()
                    char_count = len(manual_draft.replace(" ", "").replace("\n", ""))
                    motivation = get_motivation_by_progress()
                    st.success(f"✅ 저장했어! ({char_count}자)\n\n{motivation}")
                    st.rerun()

    with col2:
        st.markdown("### 📋 진행 현황")

        # 현재 Part만 표시 (인지 부하 감소)
        current_part_num = current_section['part']
        current_part_sections = [s for s in parsed_toc if s['part'] == current_part_num]
        part_completed = sum(1 for s in current_part_sections
                            if f"{s['section_num']}_{s['section_title']}" in drafts)

        st.markdown(f"**Part {current_part_num}** ({part_completed}/{len(current_part_sections)})")

        # 키보드 접근성 개선: 각 버튼에 명확한 상태 설명
        for idx, section in enumerate(parsed_toc):
            if section['part'] != current_part_num:
                continue

            prefix, help_text = get_section_status(section, drafts, current_idx, idx)
            display_text = f"{prefix} {section['section_num']}. {section['section_title'][:12]}..."

            if st.button(
                display_text,
                key=f"jump_{idx}",
                use_container_width=True,
                help=help_text
            ):
                st.session_state.current_section_index = idx
                st.rerun()

        # 다른 Part는 접어서 표시
        with st.expander("📑 다른 Part 보기", expanded=False):
            for part_num in range(1, 6):
                if part_num == current_part_num:
                    continue
                part_sections = [s for s in parsed_toc if s['part'] == part_num]
                if not part_sections:
                    continue
                completed = sum(1 for s in part_sections
                               if f"{s['section_num']}_{s['section_title']}" in drafts)
                if st.button(f"Part {part_num} ({completed}/{len(part_sections)})",
                            key=f"part_{part_num}"):
                    first_section = part_sections[0]
                    st.session_state.current_section_index = parsed_toc.index(first_section)
                    st.rerun()

        st.markdown("---")

        # 전체 미완료 장 수
        all_unfinished = [s for s in parsed_toc
                         if f"{s['section_num']}_{s['section_title']}" not in drafts]

        if all_unfinished:
            st.markdown(f"### ⚡ 빠른 완성")
            st.markdown(f"**남은 장: {len(all_unfinished)}개**")

            # 전체 자동 생성 버튼 (핵심 기능!)
            if st.button(f"🚀 전체 자동 생성", use_container_width=True, type="primary"):
                progress_bar = st.progress(0)
                status_text = st.empty()

                for i, section in enumerate(all_unfinished):
                    key = f"{section['section_num']}_{section['section_title']}"
                    status_text.text(f"📝 {i+1}/{len(all_unfinished)}: {section['section_title'][:20]}...")

                    section_info = {
                        "part_number": section["part"],
                        "part_title": section["part_title"],
                        "section_number": section["section_num"],
                        "section_title": section["section_title"],
                        "core_message": "",
                        "examples": "",
                    }

                    result = generate_draft(st.session_state.book_info, section_info)
                    if result:
                        st.session_state.drafts[key] = result
                        # 성취 시스템 호출 - 장 완료 처리
                        on_chapter_complete()

                    progress_bar.progress((i + 1) / len(all_unfinished))

                st.balloons()
                st.snow()
                st.success("🎉 모든 첫 번째 글 완성! 👑 책 완성!")
                st.rerun()

            st.caption("💡 남은 모든 장을 한번에 자동 생성해요")
        else:
            st.success("🎉 모든 장 완료!")

    # 네비게이션
    st.markdown("---")
    col1, col2 = st.columns(2)
    with col1:
        if st.button("← 목차로"):
            navigate_to_step(3)
    with col2:
        if st.button("다음: 책 소개서 →", use_container_width=True):
            navigate_to_step(5)


def render_step5():
    """5단계: 책 소개서"""
    # UX 개선: 현재 위치 브레드크럼 표시
    render_ux_breadcrumb(5)

    st.markdown('<h2 class="step-header" id="step-5-header" tabindex="-1">5단계: 책 소개서</h2>', unsafe_allow_html=True)

    # UX 개선: 단계별 팁 표시
    render_ux_step_tip(5)

    st.markdown("""
    출판사에 보여줄 책 소개서를 만들어요!

    **저자 정보가 중요한 이유:** 출판사는 저자의 신뢰도와 영향력을 확인합니다. 당신의 경력과 성과를 입력해주세요!
    """)

    col1, col2 = st.columns([2, 1])

    with col1:
        # 개선점 2: 저자 정보 입력 템플릿 제공
        with st.expander("📋 저자 정보 입력 템플릿 (참고용)", expanded=False):
            st.markdown("""
            **예시 1: 컨설턴트**
            - 직업/전문 분야: 부동산 컨설턴트
            - 경력/전문성: 10년간 부동산 투자, 100명 이상의 클라이언트 자산 관리
            - 대표 성과: 서울 강남 프로젝트 성공 사례 다수, 유튜브 구독자 5만명
            - SNS/블로그: 유튜브 채널 5만, 블로그 월 10만 방문

            **예시 2: 기업 임원**
            - 직업/전문 분야: 마케팅 이사 (업력 12년)
            - 경력/전문성: 500억 규모 마케팅 캠페인 성공 주도, 3개 회사 성장 경험
            - 대표 성과: MBC 경제 프로그램 출연, 비즈니스 신문 칼럼니스트
            - SNS/블로그: 링크드인 팔로워 2만
            """)

        # 저자 정보 입력
        with st.expander("📝 저자 정보 입력", expanded=not st.session_state.generated_proposal):
            profession = st.text_input(
                "직업/전문 분야",
                value=st.session_state.author_info.get("profession", ""),
                placeholder="예: 부동산 컨설턴트, 10년차 마케터",
                help="💡 현재 직업과 전문성을 명확히 입력하세요. 업력 정보가 있으면 좋습니다."
            )
            career = st.text_area(
                "경력/전문성",
                value=st.session_state.author_info.get("career", ""),
                placeholder="예: 부동산 투자 15년, 자산 100억 달성",
                height=80,
                help="💡 구체적인 숫자와 성과를 포함하세요. 예) '100명 클라이언트 관리', '500억 프로젝트 주도'"
            )
            achievements = st.text_area(
                "대표 성과",
                value=st.session_state.author_info.get("achievements", ""),
                placeholder="예: 베스트셀러 저자, TV 출연",
                height=80,
                help="💡 미디어 출연, 상장 사건, 저널 기고 등 신뢰도를 높이는 경력사항"
            )
            sns = st.text_input(
                "SNS/블로그",
                value=st.session_state.author_info.get("sns", ""),
                placeholder="예: 인스타 1만, 유튜브 5천",
                help="💡 팔로워 수를 명시해 주세요. 영향력 지표가 됩니다."
            )
            contact = st.text_input(
                "연락처",
                value=st.session_state.author_info.get("contact", ""),
                placeholder="예: email@example.com",
                help="💡 출판사가 연락할 이메일이나 전화번호"
            )

            if st.button("💾 저자 정보 저장", use_container_width=True, type="secondary"):
                st.session_state.author_info = {
                    "name": st.session_state.book_info.get("name", ""),
                    "profession": profession,
                    "career": career,
                    "achievements": achievements,
                    "sns": sns,
                    "contact": contact,
                }
                st.success("✅ 저자 정보가 저장되었습니다!")

        # 기획서 생성
        if st.button("📄 책 소개서 만들기", use_container_width=True, type="primary"):
            if not st.session_state.author_info:
                st.warning("먼저 저자 정보를 저장해줘!")
            else:
                with st.spinner("✨ 책 소개서를 만들고 있어요..."):
                    result = generate_proposal(
                        st.session_state.book_info,
                        st.session_state.author_info
                    )
                    if result:
                        st.session_state.generated_proposal = result
                        st.rerun()

        # 생성된 기획서 표시
        if st.session_state.generated_proposal:
            st.markdown("### 📋 만들어진 책 소개서")
            edited_proposal = st.text_area(
                "책 소개서 수정 (직접 고칠 수 있어요)",
                value=st.session_state.generated_proposal,
                height=500,
                label_visibility="collapsed"
            )
            st.session_state.generated_proposal = edited_proposal

            st.download_button(
                label="📥 책 소개서 저장하기",
                data=edited_proposal,
                file_name=f"{st.session_state.selected_title}_책소개서.md",
                mime="text/markdown",
                use_container_width=True
            )

    with col2:
        st.markdown("### 💡 책 소개서 팁")
        st.info("""
        **꼭 들어가야 할 7가지:**
        1. 제목 & 부제목
        2. 왜 이 책을 썼을까?
        3. 누가 읽을까?
        4. 비슷한 책 분석
        5. 목차 요약
        6. 나(저자) 소개
        7. 홍보 계획
        """)

    # 네비게이션
    st.markdown("---")
    col1, col2 = st.columns(2)
    with col1:
        if st.button("← 이전"):
            navigate_to_step(4)
    with col2:
        if st.button("다음 →", use_container_width=True, type="primary"):
            navigate_to_step(6)


def render_step6():
    """6단계: 책 홍보 페이지"""
    # UX 개선: 현재 위치 브레드크럼 표시
    render_ux_breadcrumb(6)

    st.markdown('<h2 class="step-header" id="step-6-header" tabindex="-1">6단계: 책 홍보 페이지</h2>', unsafe_allow_html=True)

    # UX 개선: 단계별 팁 표시
    render_ux_step_tip(6)

    st.markdown("책을 홍보할 수 있는 페이지 글을 만들어요!")

    col1, col2 = st.columns([2, 1])

    with col1:
        # 개선점 2: 웨비나 정보 입력 템플릿 제공
        with st.expander("📋 웨비나 정보 입력 템플릿 (참고용)", expanded=False):
            st.markdown("""
            **예시: 부동산 투자 웨비나**
            - 웨비나 제목: 월급쟁이도 10년 안에 경제적 자유를 얻는 부동산 투자법
            - 일시: 2025년 2월 15일 (토) 오후 2시 ~ 3시 30분
            - 강사: 김재무
            - 주요 내용:
              - 부동산 초보자가 저지르는 3가지 실수
              - 월 500만원으로 시작하는 포트폴리오 구성
              - 2024년 부동산 시장 분석 & 기회 포착법
              - 실전 사례: 강남, 여의도 프로젝트 경험담
            - 보너스/혜택:
              - 부동산 투자 체크리스트 (PDF)
              - 지역별 시세 분석 자료
              - 참석자 전용 1:1 상담 권리 (1명 선착순)
            """)

        # 웨비나 정보 입력
        with st.expander("📝 웨비나/이벤트 정보", expanded=not st.session_state.generated_landing_page):
            webinar_title = st.text_input(
                "웨비나 제목",
                value=st.session_state.webinar_info.get("webinar_title", ""),
                placeholder="예: 2달 만에 책 쓰는 비밀 공개",
                help="💡 책에서 가장 중요한 이야기를 담은 멋진 제목을 써봐!"
            )
            datetime_input = st.text_input(
                "일시",
                value=st.session_state.webinar_info.get("datetime", ""),
                placeholder="예: 2025년 2월 15일 (토) 오후 2시",
                help="💡 구체적인 날짜와 시간(소요 시간 포함)을 입력하세요."
            )
            speaker = st.text_input(
                "강사",
                value=st.session_state.webinar_info.get("speaker", st.session_state.book_info.get("name", "")),
                help="💡 웨비나를 주도할 강사(저자)의 이름"
            )
            content = st.text_area(
                "주요 내용",
                value=st.session_state.webinar_info.get("content", ""),
                height=80,
                placeholder="- 주제1: 구체적 내용\n- 주제2: 학습할 내용\n- 사례: 성공 사례 소개",
                help="💡 3-5개의 핵심 주제를 불릿 포인트로 작성하세요."
            )
            bonus = st.text_area(
                "보너스/혜택",
                value=st.session_state.webinar_info.get("bonus", ""),
                height=80,
                placeholder="- 제공 자료: PDF 템플릿\n- 혜택: 1:1 상담\n- 할인: 책 출간 기념 30% 할인 쿠폰",
                help="💡 참석자를 끌어당길 수 있는 구체적인 혜택을 나열하세요."
            )

            if st.button("💾 웨비나 정보 저장", use_container_width=True, type="secondary"):
                st.session_state.webinar_info = {
                    "webinar_title": webinar_title,
                    "datetime": datetime_input,
                    "speaker": speaker,
                    "content": content,
                    "bonus": bonus,
                }
                st.success("✅ 웨비나 정보가 저장되었습니다!")

        # 홍보 페이지 생성
        if st.button("🎨 홍보 페이지 글 만들기", use_container_width=True, type="primary"):
            if not st.session_state.webinar_info:
                st.warning("먼저 웨비나 정보를 저장해줘!")
            else:
                with st.spinner("✨ 홍보 페이지 글을 만들고 있어요..."):
                    result = generate_landing_page(
                        st.session_state.book_info,
                        st.session_state.webinar_info
                    )
                    if result:
                        st.session_state.generated_landing_page = result
                        st.rerun()

        # 생성된 홍보 페이지 표시
        if st.session_state.generated_landing_page:
            st.markdown("### 🎨 만들어진 홍보 페이지 글")
            edited_landing = st.text_area(
                "글 수정 (직접 고칠 수 있어요)",
                value=st.session_state.generated_landing_page,
                height=500,
                label_visibility="collapsed"
            )
            st.session_state.generated_landing_page = edited_landing

            st.download_button(
                label="📥 홍보 페이지 저장하기",
                data=edited_landing,
                file_name=f"{st.session_state.selected_title}_홍보페이지.md",
                mime="text/markdown",
                use_container_width=True
            )

    with col2:
        st.markdown("### 💡 홍보 페이지 구조")
        st.info("""
        1. 눈에 띄는 제목
        2. 어떤 문제를 해결해줄까?
        3. 해결 방법 알려주기
        4. 강사(저자) 소개
        5. 강의 내용
        6. 후기/성과
        7. 보너스 혜택
        8. 신청 방법
        9. 자주 묻는 질문
        10. 마지막 한마디
        """)

    # 네비게이션
    st.markdown("---")
    col1, col2 = st.columns(2)
    with col1:
        if st.button("← 이전"):
            navigate_to_step(5)
    with col2:
        if st.button("다음: 최종 다운로드 →", use_container_width=True, type="primary"):
            navigate_to_step(7)


def analyze_reading_level(text):
    """읽기 수준 분석 (강화 버전)"""
    if not text:
        return {
            "level": "알 수 없음",
            "avg_sentence_len": 0,
            "complex_ratio": 0,
            "total_chars": 0,
            "total_words": 0,
            "total_sentences": 0,
            "estimated_pages_a4": 0,
            "reading_time_minutes": 0,
            "difficulty_score": 0,
            "difficulty_label": "알 수 없음"
        }

    # 문장 분리 (간단)
    sentences = [s.strip() for s in re.split(r'[.!?。]', text) if s.strip()]
    if not sentences:
        return {
            "level": "알 수 없음",
            "avg_sentence_len": 0,
            "complex_ratio": 0,
            "total_chars": 0,
            "total_words": 0,
            "total_sentences": 0,
            "estimated_pages_a4": 0,
            "reading_time_minutes": 0,
            "difficulty_score": 0,
            "difficulty_label": "알 수 없음"
        }

    # 기본 통계
    total_chars = len(text.replace('\n', '').replace(' ', ''))
    words = text.replace('\n', ' ').split()
    total_words = len(words)
    total_sentences = len(sentences)

    # 평균 문장 길이
    avg_len = sum(len(s) for s in sentences) / len(sentences)

    # 복잡한 단어 비율 (4음절 이상 추정)
    complex_words = [w for w in words if len(w) >= 8]  # 한글 기준 긴 단어
    complex_ratio = len(complex_words) / len(words) * 100 if words else 0

    # A4 페이지 수 계산 (1800자 기준)
    estimated_pages_a4 = max(1, round(total_chars / 1800, 1))

    # 읽기 시간 계산 (한국어 평균 독서 속도: 분당 400~600자)
    reading_time_minutes = max(1, round(total_chars / 500))

    # 난이도 점수 계산 (0-100)
    difficulty_score = min(100, int((avg_len / 80 * 50) + (complex_ratio * 2.5)))

    # 난이도 레이블
    if difficulty_score < 30:
        difficulty_label = "쉬움"
    elif difficulty_score < 60:
        difficulty_label = "보통"
    else:
        difficulty_label = "어려움"

    # 수준 판정
    if avg_len < 30 and complex_ratio < 10:
        level = "쉬움 (초등~중등)"
    elif avg_len < 50 and complex_ratio < 20:
        level = "보통 (고등~대학)"
    else:
        level = "어려움 (전문가)"

    return {
        "level": level,
        "avg_sentence_len": round(avg_len, 1),
        "complex_ratio": round(complex_ratio, 1),
        "total_chars": total_chars,
        "total_words": total_words,
        "total_sentences": total_sentences,
        "estimated_pages_a4": estimated_pages_a4,
        "reading_time_minutes": reading_time_minutes,
        "difficulty_score": difficulty_score,
        "difficulty_label": difficulty_label
    }


def generate_book_manuscript():
    """책다운 원고 생성 (표지, 저작권, 에필로그 포함)"""
    title = st.session_state.selected_title or "무제"
    author = st.session_state.book_info.get("name", "저자")
    topic = st.session_state.book_info.get("topic", "")
    core_message = st.session_state.book_info.get("core_message", "")
    year = datetime.now().year

    # 표지 페이지
    manuscript = f"""
{'='*60}

# {title}

{'='*60}

**{author} 지음**

{'-'*60}





{'='*60}

## 저작권 안내

{'='*60}

© {year} {author}

이 책의 저작권은 저자에게 있습니다.
무단 전재와 복제를 금합니다.

초판 발행: {year}년

저자: {author}
제작: AI 책쓰기 코칭 시스템

{'-'*60}





{'='*60}

## 프롤로그

{'='*60}

{core_message}

이 책은 {topic}에 대해 다룹니다.
독자 여러분의 인생에 작은 변화가 되기를 바랍니다.

{'-'*60}





{'='*60}

## 목차

{'='*60}

{st.session_state.generated_toc}

{'-'*60}




"""

    # 본문 추가
    if st.session_state.parsed_toc:
        current_part = None
        for section in st.session_state.parsed_toc:
            key = f"{section['section_num']}_{section['section_title']}"

            if section['part'] != current_part:
                current_part = section['part']
                manuscript += f"""

{'='*60}

# Part {current_part}. {section['part_title']}

{'='*60}

"""

            if key in st.session_state.drafts:
                manuscript += f"""
## {section['section_num']}. {section['section_title']}

{st.session_state.drafts[key]}

{'-'*40}

"""

    # 에필로그
    manuscript += f"""


{'='*60}

## 에필로그

{'='*60}

이 책을 끝까지 읽어주셔서 감사합니다.

{core_message}

여러분의 여정을 응원합니다.

**{author} 드림**

{'-'*60}




{'='*60}

## 저자 소개

{'='*60}

**{author}**

{st.session_state.book_info.get('experience', '저자 경력 정보가 없습니다.')}

{'-'*60}
"""

    return manuscript


def generate_html_manuscript():
    """HTML 형식 원고 생성"""
    title = st.session_state.selected_title or "무제"
    author = st.session_state.book_info.get("name", "저자")
    core_message = st.session_state.book_info.get("core_message", "")
    year = datetime.now().year

    html = f"""<!DOCTYPE html>
<html lang="ko">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{title}</title>
    <style>
        body {{
            font-family: 'Noto Sans KR', sans-serif;
            line-height: 1.8;
            max-width: 800px;
            margin: 0 auto;
            padding: 40px 20px;
            color: #333;
        }}
        h1 {{ font-size: 2.5em; text-align: center; margin: 60px 0; border-bottom: 3px solid #333; padding-bottom: 20px; }}
        h2 {{ font-size: 1.8em; margin-top: 50px; color: #1565C0; }}
        h3 {{ font-size: 1.4em; margin-top: 30px; }}
        .cover {{ text-align: center; padding: 100px 0; border: 2px solid #333; margin-bottom: 60px; }}
        .cover h1 {{ border: none; }}
        .copyright {{ background: #f5f5f5; padding: 30px; margin: 40px 0; font-size: 0.9em; }}
        .chapter {{ margin: 40px 0; padding: 20px 0; border-top: 1px solid #eee; }}
        .epilogue {{ background: #E3F2FD; padding: 30px; margin-top: 60px; border-radius: 8px; }}
        p {{ margin: 1.2em 0; }}
    </style>
</head>
<body>
    <div class="cover">
        <h1>{title}</h1>
        <p style="font-size: 1.5em; margin-top: 40px;">{author} 지음</p>
    </div>

    <div class="copyright">
        <p>© {year} {author}</p>
        <p>이 책의 저작권은 저자에게 있습니다.</p>
    </div>

    <h2>프롤로그</h2>
    <p>{core_message}</p>

    <h2>목차</h2>
    <pre>{st.session_state.generated_toc}</pre>
"""

    # 본문 추가
    if st.session_state.parsed_toc:
        current_part = None
        for section in st.session_state.parsed_toc:
            key = f"{section['section_num']}_{section['section_title']}"

            if section['part'] != current_part:
                current_part = section['part']
                html += f"""
    <h1>Part {current_part}. {section['part_title']}</h1>
"""

            if key in st.session_state.drafts:
                content = st.session_state.drafts[key].replace('\n', '</p><p>')
                html += f"""
    <div class="chapter">
        <h2>{section['section_num']}. {section['section_title']}</h2>
        <p>{content}</p>
    </div>
"""

    # 에필로그
    html += f"""
    <div class="epilogue">
        <h2>에필로그</h2>
        <p>이 책을 끝까지 읽어주셔서 감사합니다.</p>
        <p>{core_message}</p>
        <p><strong>{author} 드림</strong></p>
    </div>
</body>
</html>
"""
    return html


def generate_docx_manuscript():
    """DOCX 형식 원고 생성 (python-docx 사용)"""
    try:
        from docx import Document
        from docx.shared import Inches, Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.style import WD_STYLE_TYPE
        import io

        title = st.session_state.selected_title or "무제"
        author = st.session_state.book_info.get("name", "저자")
        core_message = st.session_state.book_info.get("core_message", "")
        year = datetime.now().year

        doc = Document()

        # 페이지 설정
        sections = doc.sections
        for section in sections:
            section.page_width = Cm(21)  # A4
            section.page_height = Cm(29.7)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2.5)

        # 표지 페이지
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(title)
        title_run.font.size = Pt(36)
        title_run.bold = True
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()
        doc.add_paragraph()

        author_para = doc.add_paragraph()
        author_run = author_para.add_run(f"{author} 지음")
        author_run.font.size = Pt(18)
        author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_page_break()

        # 저작권 페이지
        doc.add_heading("저작권 안내", level=1)
        doc.add_paragraph(f"(C) {year} {author}")
        doc.add_paragraph("이 책의 저작권은 저자에게 있습니다.")
        doc.add_paragraph("무단 전재와 복제를 금합니다.")
        doc.add_paragraph()
        doc.add_paragraph(f"초판 발행: {year}년")
        doc.add_paragraph(f"저자: {author}")
        doc.add_paragraph("제작: AI 책쓰기 코칭 시스템")

        doc.add_page_break()

        # 프롤로그
        doc.add_heading("프롤로그", level=1)
        doc.add_paragraph(core_message)

        doc.add_page_break()

        # 목차
        doc.add_heading("목차", level=1)
        if st.session_state.parsed_toc:
            current_part = None
            for section in st.session_state.parsed_toc:
                if section['part'] != current_part:
                    current_part = section['part']
                    part_para = doc.add_paragraph()
                    part_run = part_para.add_run(f"Part {current_part}. {section['part_title']}")
                    part_run.bold = True
                doc.add_paragraph(f"    {section['section_num']}. {section['section_title']}")

        doc.add_page_break()

        # 본문
        if st.session_state.parsed_toc:
            current_part = None
            for section in st.session_state.parsed_toc:
                key = f"{section['section_num']}_{section['section_title']}"

                if section['part'] != current_part:
                    current_part = section['part']
                    doc.add_page_break()
                    part_heading = doc.add_heading(f"Part {current_part}. {section['part_title']}", level=1)
                    part_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

                if key in st.session_state.drafts:
                    doc.add_heading(f"{section['section_num']}. {section['section_title']}", level=2)
                    content = st.session_state.drafts[key]
                    paragraphs = content.split('\n\n')
                    for para_text in paragraphs:
                        if para_text.strip():
                            doc.add_paragraph(para_text.strip())

        # 에필로그
        doc.add_page_break()
        doc.add_heading("에필로그", level=1)
        doc.add_paragraph("이 책을 끝까지 읽어주셔서 감사합니다.")
        doc.add_paragraph(core_message)
        doc.add_paragraph(f"{author} 드림")

        # 저자 소개
        doc.add_page_break()
        doc.add_heading("저자 소개", level=1)
        author_heading = doc.add_paragraph()
        author_name_run = author_heading.add_run(author)
        author_name_run.bold = True
        doc.add_paragraph(st.session_state.book_info.get('experience', '저자 경력 정보가 없습니다.'))

        # 바이트 스트림으로 저장
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    except ImportError:
        return None


def generate_pdf_manuscript():
    """PDF 형식 원고 생성 (reportlab 사용)"""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
        import io
        import os

        title = st.session_state.selected_title or "무제"
        author = st.session_state.book_info.get("name", "저자")
        core_message = st.session_state.book_info.get("core_message", "")
        year = datetime.now().year

        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            leftMargin=2.5*cm,
            rightMargin=2.5*cm,
            topMargin=2.5*cm,
            bottomMargin=2.5*cm
        )

        styles = getSampleStyleSheet()

        # 한글 폰트 등록 시도
        font_registered = False
        font_paths = [
            "C:/Windows/Fonts/malgun.ttf",
            "C:/Windows/Fonts/NanumGothic.ttf",
            "/usr/share/fonts/truetype/nanum/NanumGothic.ttf",
            "/System/Library/Fonts/AppleGothic.ttf"
        ]
        for font_path in font_paths:
            if os.path.exists(font_path):
                try:
                    pdfmetrics.registerFont(TTFont('Korean', font_path))
                    font_registered = True
                    break
                except:
                    continue

        font_name = 'Korean' if font_registered else 'Helvetica'

        # 커스텀 스타일
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Title'],
            fontName=font_name,
            fontSize=28,
            alignment=TA_CENTER,
            spaceAfter=30
        )

        heading1_style = ParagraphStyle(
            'CustomHeading1',
            parent=styles['Heading1'],
            fontName=font_name,
            fontSize=20,
            spaceAfter=20,
            spaceBefore=30
        )

        heading2_style = ParagraphStyle(
            'CustomHeading2',
            parent=styles['Heading2'],
            fontName=font_name,
            fontSize=16,
            spaceAfter=15,
            spaceBefore=20
        )

        body_style = ParagraphStyle(
            'CustomBody',
            parent=styles['Normal'],
            fontName=font_name,
            fontSize=11,
            leading=18,
            alignment=TA_JUSTIFY,
            spaceAfter=12
        )

        story = []

        # 표지
        story.append(Spacer(1, 5*cm))
        story.append(Paragraph(title, title_style))
        story.append(Spacer(1, 2*cm))
        story.append(Paragraph(f"{author} 지음", ParagraphStyle('Author', parent=body_style, alignment=TA_CENTER, fontSize=14)))
        story.append(PageBreak())

        # 저작권
        story.append(Paragraph("저작권 안내", heading1_style))
        story.append(Paragraph(f"(C) {year} {author}", body_style))
        story.append(Paragraph("이 책의 저작권은 저자에게 있습니다.", body_style))
        story.append(Paragraph("무단 전재와 복제를 금합니다.", body_style))
        story.append(Spacer(1, 1*cm))
        story.append(Paragraph(f"초판 발행: {year}년", body_style))
        story.append(Paragraph(f"저자: {author}", body_style))
        story.append(PageBreak())

        # 프롤로그
        story.append(Paragraph("프롤로그", heading1_style))
        story.append(Paragraph(core_message, body_style))
        story.append(PageBreak())

        # 목차
        story.append(Paragraph("목차", heading1_style))
        if st.session_state.parsed_toc:
            current_part = None
            for section in st.session_state.parsed_toc:
                if section['part'] != current_part:
                    current_part = section['part']
                    story.append(Paragraph(f"<b>Part {current_part}. {section['part_title']}</b>", body_style))
                story.append(Paragraph(f"    {section['section_num']}. {section['section_title']}", body_style))
        story.append(PageBreak())

        # 본문
        if st.session_state.parsed_toc:
            current_part = None
            for section in st.session_state.parsed_toc:
                key = f"{section['section_num']}_{section['section_title']}"

                if section['part'] != current_part:
                    current_part = section['part']
                    story.append(PageBreak())
                    story.append(Paragraph(f"Part {current_part}. {section['part_title']}", title_style))

                if key in st.session_state.drafts:
                    story.append(Paragraph(f"{section['section_num']}. {section['section_title']}", heading2_style))
                    content = st.session_state.drafts[key]
                    paragraphs = content.split('\n\n')
                    for para_text in paragraphs:
                        if para_text.strip():
                            # XML 특수문자 이스케이프
                            safe_text = para_text.strip().replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                            story.append(Paragraph(safe_text, body_style))

        # 에필로그
        story.append(PageBreak())
        story.append(Paragraph("에필로그", heading1_style))
        story.append(Paragraph("이 책을 끝까지 읽어주셔서 감사합니다.", body_style))
        story.append(Paragraph(core_message, body_style))
        story.append(Paragraph(f"{author} 드림", body_style))

        doc.build(story)
        buffer.seek(0)
        return buffer.getvalue()

    except ImportError:
        return None
    except Exception as e:
        st.error(f"PDF 생성 오류: {str(e)}")
        return None


def generate_print_html():
    """인쇄 최적화 HTML 생성 (목차 포함, 페이지 번호)"""
    title = st.session_state.selected_title or "무제"
    author = st.session_state.book_info.get("name", "저자")
    core_message = st.session_state.book_info.get("core_message", "")
    year = datetime.now().year

    # 목차 생성
    toc_html = ""
    if st.session_state.parsed_toc:
        current_part = None
        for section in st.session_state.parsed_toc:
            if section['part'] != current_part:
                current_part = section['part']
                toc_html += f'<li class="toc-part"><a href="#part-{current_part}">Part {current_part}. {section["part_title"]}</a></li>'
            toc_html += f'<li class="toc-chapter"><a href="#chapter-{section["section_num"]}">{section["section_num"]}. {section["section_title"]}</a></li>'

    html = f"""<!DOCTYPE html>
<html lang="ko">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{title} - 인쇄용</title>
    <style>
        @media print {{
            @page {{
                size: A4;
                margin: 2cm 2.5cm;
                @bottom-center {{
                    content: counter(page);
                }}
            }}
            body {{
                font-size: 11pt;
                line-height: 1.6;
            }}
            .no-print {{ display: none !important; }}
            .page-break {{ page-break-before: always; }}
            h1, h2, h3 {{ page-break-after: avoid; }}
            p {{ orphans: 3; widows: 3; }}
        }}

        @media screen {{
            body {{
                max-width: 210mm;
                margin: 0 auto;
                padding: 40px 20px;
                background: #f5f5f5;
            }}
            .print-page {{
                background: white;
                padding: 40px 60px;
                margin-bottom: 20px;
                box-shadow: 0 2px 10px rgba(0,0,0,0.1);
            }}
        }}

        body {{
            font-family: 'Noto Serif KR', 'Batang', Georgia, serif;
            color: #333;
            counter-reset: page;
        }}

        .cover {{
            text-align: center;
            padding: 150px 0;
        }}

        .cover h1 {{
            font-size: 32pt;
            margin-bottom: 40px;
            border: none;
        }}

        .cover .author {{
            font-size: 16pt;
            margin-top: 60px;
        }}

        h1 {{
            font-size: 24pt;
            margin: 40px 0 20px 0;
            padding-bottom: 10px;
            border-bottom: 2px solid #333;
        }}

        h2 {{
            font-size: 18pt;
            margin: 30px 0 15px 0;
            color: #1565C0;
        }}

        h3 {{
            font-size: 14pt;
            margin: 20px 0 10px 0;
        }}

        p {{
            text-align: justify;
            margin: 12px 0;
            text-indent: 1em;
        }}

        .copyright {{
            font-size: 10pt;
            margin: 40px 0;
            padding: 20px;
            background: #f9f9f9;
            border: 1px solid #ddd;
        }}

        .toc ul {{
            list-style: none;
            padding: 0;
        }}

        .toc li {{
            margin: 8px 0;
        }}

        .toc-part {{
            font-weight: bold;
            margin-top: 20px !important;
            font-size: 14pt;
        }}

        .toc-chapter {{
            padding-left: 30px;
        }}

        .toc a {{
            color: #333;
            text-decoration: none;
        }}

        .toc a:hover {{
            color: #1565C0;
        }}

        .chapter {{
            margin: 30px 0;
        }}

        .epilogue {{
            margin-top: 60px;
            padding: 30px;
            background: #f0f7ff;
            border-radius: 8px;
        }}

        .footer {{
            text-align: center;
            margin-top: 40px;
            padding-top: 20px;
            border-top: 1px solid #ddd;
            font-size: 10pt;
            color: #666;
        }}

        .print-btn {{
            position: fixed;
            top: 20px;
            right: 20px;
            padding: 15px 30px;
            background: #1565C0;
            color: white;
            border: none;
            border-radius: 8px;
            cursor: pointer;
            font-size: 14pt;
            box-shadow: 0 4px 15px rgba(21, 101, 192, 0.3);
        }}

        .print-btn:hover {{
            background: #0D47A1;
        }}
    </style>
</head>
<body>
    <button class="print-btn no-print" onclick="window.print()">인쇄하기</button>

    <div class="print-page cover">
        <h1>{title}</h1>
        <p class="author">{author} 지음</p>
    </div>

    <div class="print-page page-break">
        <div class="copyright">
            <p><strong>저작권 안내</strong></p>
            <p>(C) {year} {author}</p>
            <p>이 책의 저작권은 저자에게 있습니다. 무단 전재와 복제를 금합니다.</p>
            <p>초판 발행: {year}년 | 저자: {author}</p>
        </div>
    </div>

    <div class="print-page page-break">
        <h1>프롤로그</h1>
        <p>{core_message}</p>
    </div>

    <div class="print-page page-break toc">
        <h1>목차</h1>
        <ul>
            {toc_html}
        </ul>
    </div>
"""

    # 본문 추가
    if st.session_state.parsed_toc:
        current_part = None
        for section in st.session_state.parsed_toc:
            key = f"{section['section_num']}_{section['section_title']}"

            if section['part'] != current_part:
                current_part = section['part']
                html += f"""
    <div class="print-page page-break">
        <h1 id="part-{current_part}">Part {current_part}. {section['part_title']}</h1>
    </div>
"""

            if key in st.session_state.drafts:
                content = st.session_state.drafts[key].replace('\n\n', '</p><p>').replace('\n', '<br>')
                html += f"""
    <div class="print-page chapter">
        <h2 id="chapter-{section['section_num']}">{section['section_num']}. {section['section_title']}</h2>
        <p>{content}</p>
    </div>
"""

    # 에필로그
    html += f"""
    <div class="print-page page-break">
        <div class="epilogue">
            <h1>에필로그</h1>
            <p>이 책을 끝까지 읽어주셔서 감사합니다.</p>
            <p>{core_message}</p>
            <p><strong>{author} 드림</strong></p>
        </div>
    </div>

    <div class="print-page">
        <h1>저자 소개</h1>
        <p><strong>{author}</strong></p>
        <p>{st.session_state.book_info.get('experience', '저자 경력 정보가 없습니다.')}</p>
    </div>

    <div class="footer no-print">
        <p>AI 책쓰기 코칭 시스템으로 제작됨</p>
    </div>
</body>
</html>
"""
    return html


def generate_share_link():
    """공유 링크 생성 (세션 기반 임시 링크)"""
    import hashlib
    import base64

    # 현재 세션 데이터 해시 생성
    data = {
        "title": st.session_state.selected_title,
        "author": st.session_state.book_info.get("name", ""),
        "timestamp": datetime.now().isoformat()
    }
    data_str = json.dumps(data, ensure_ascii=False)
    hash_obj = hashlib.md5(data_str.encode())
    share_id = base64.urlsafe_b64encode(hash_obj.digest()[:8]).decode().rstrip('=')

    return share_id


def render_step7():
    """7단계: 결과물 다운로드 (책다운 출력) - 강화 버전"""
    # UX 개선: 현재 위치 브레드크럼 표시
    render_ux_breadcrumb(7)

    st.markdown('<h2 class="step-header" id="step-7-header" tabindex="-1">7단계: 완성! 다운로드</h2>', unsafe_allow_html=True)

    # UX 개선: 단계별 팁 표시
    render_ux_step_tip(7)

    # UX 개선: 화려한 축하 효과 표시
    render_ux_celebration(
        "책이 완성되었습니다!",
        "축하합니다! 이제 당신도 작가입니다. 아래에서 원고를 다운로드하세요.",
        ["🎉", "📚", "✨"]
    )

    # 최종 진행률
    render_progress_bar()

    # 통계
    stats = get_progress_stats()

    # 원고 분석 (강화된 버전)
    all_text = " ".join(st.session_state.drafts.values())
    reading_analysis = analyze_reading_level(all_text)

    # ===== 1. 강화된 통계 섹션 =====
    st.markdown("### 📊 원고 통계")

    # 메인 통계 (5컬럼)
    col1, col2, col3, col4, col5 = st.columns(5)
    with col1:
        st.metric("총 글자 수", f"{reading_analysis['total_chars']:,}자")
    with col2:
        st.metric("총 단어 수", f"{reading_analysis['total_words']:,}개")
    with col3:
        st.metric("총 문장 수", f"{reading_analysis['total_sentences']:,}개")
    with col4:
        st.metric("예상 페이지 (A4)", f"약 {reading_analysis['estimated_pages_a4']}쪽")
    with col5:
        st.metric("읽기 예상 시간", f"약 {reading_analysis['reading_time_minutes']}분")

    # 난이도 분석 표시
    difficulty_colors = {"쉬움": "#4CAF50", "보통": "#FF9800", "어려움": "#F44336"}
    difficulty_color = difficulty_colors.get(reading_analysis['difficulty_label'], "#9E9E9E")

    st.markdown(f"""
    <div style="background: linear-gradient(90deg, {difficulty_color}22 0%, transparent 100%);
                padding: 15px; border-radius: 10px; margin: 15px 0; border-left: 4px solid {difficulty_color};">
        <strong>난이도:</strong> <span style="color: {difficulty_color}; font-weight: bold;">{reading_analysis['difficulty_label']}</span>
        (점수: {reading_analysis['difficulty_score']}/100) |
        <strong>읽기 수준:</strong> {reading_analysis['level']}
    </div>
    """, unsafe_allow_html=True)

    # 상세 분석 Expander
    with st.expander("상세 분석 보기", expanded=False):
        detail_col1, detail_col2, detail_col3, detail_col4 = st.columns(4)
        with detail_col1:
            st.metric("평균 문장 길이", f"{reading_analysis['avg_sentence_len']}자")
        with detail_col2:
            st.metric("복잡한 단어 비율", f"{reading_analysis['complex_ratio']}%")
        with detail_col3:
            completion_rate = (stats['completed_sections'] / stats['total_sections'] * 100) if stats['total_sections'] > 0 else 0
            st.metric("완성도", f"{completion_rate:.0f}%")
        with detail_col4:
            st.metric("작성한 장", f"{stats['completed_sections']}/{stats['total_sections']}개")

        # 난이도 게이지
        st.markdown("**난이도 분포**")
        st.progress(min(reading_analysis['difficulty_score'] / 100, 1.0))

        # 출판 준비도
        st.markdown("---")
        if stats['total_chars'] >= 55000 and completion_rate >= 90:
            st.success("**출판 준비 완료!** 원고가 출판 가능한 수준입니다.")
        elif stats['total_chars'] >= 40000:
            st.info(f"목표 6만자까지 {60000 - stats['total_chars']:,}자 남았습니다.")
        else:
            st.warning(f"조금만 더 힘내세요! 아직 {60000 - stats['total_chars']:,}자가 더 필요합니다.")

    st.markdown("---")

    # ===== 2. 미리보기 섹션 =====
    st.markdown("### 👁️ 책 미리보기")

    # 미리보기 상태 초기화
    if "preview_page" not in st.session_state:
        st.session_state.preview_page = 0
    if "preview_chapter" not in st.session_state:
        st.session_state.preview_chapter = None

    # 장 바로가기 선택
    chapter_options = ["표지/프롤로그"]
    chapter_keys = [None]
    if st.session_state.parsed_toc:
        for section in st.session_state.parsed_toc:
            key = f"{section['section_num']}_{section['section_title']}"
            if key in st.session_state.drafts:
                chapter_options.append(f"{section['section_num']}. {section['section_title']}")
                chapter_keys.append(key)
    chapter_options.append("에필로그/저자소개")
    chapter_keys.append("epilogue")

    preview_col1, preview_col2 = st.columns([1, 3])
    with preview_col1:
        selected_chapter_idx = st.selectbox(
            "장 바로가기",
            range(len(chapter_options)),
            format_func=lambda x: chapter_options[x],
            key="chapter_selector"
        )

    # 미리보기 내용 생성
    with st.container():
        preview_style = """
        <style>
        .preview-container {{
            border: 1px solid #ddd;
            border-radius: 10px;
            padding: 30px;
            background: white;
            min-height: 400px;
            max-height: 500px;
            overflow-y: auto;
            font-family: 'Noto Serif KR', Georgia, serif;
            line-height: 1.8;
        }}
        .preview-title {{
            text-align: center;
            font-size: 24px;
            font-weight: bold;
            margin-bottom: 20px;
            padding-bottom: 15px;
            border-bottom: 2px solid #333;
        }}
        .preview-chapter-title {{
            font-size: 20px;
            color: #1565C0;
            margin: 20px 0 15px 0;
        }}
        .preview-content {{
            text-align: justify;
            text-indent: 1em;
        }}
        </style>
        """
        st.markdown(preview_style, unsafe_allow_html=True)

        # 선택된 장에 따른 미리보기 내용
        selected_key = chapter_keys[selected_chapter_idx]

        if selected_key is None:
            # 표지/프롤로그
            title = st.session_state.selected_title or "무제"
            author = st.session_state.book_info.get("name", "저자")
            core_message = st.session_state.book_info.get("core_message", "")
            preview_html = f"""
            <div class="preview-container">
                <div class="preview-title">{title}</div>
                <p style="text-align: center; font-size: 16px;">{author} 지음</p>
                <hr style="margin: 30px 0;">
                <div class="preview-chapter-title">프롤로그</div>
                <div class="preview-content">{core_message}</div>
            </div>
            """
        elif selected_key == "epilogue":
            # 에필로그
            author = st.session_state.book_info.get("name", "저자")
            core_message = st.session_state.book_info.get("core_message", "")
            experience = st.session_state.book_info.get('experience', '저자 경력 정보가 없습니다.')
            preview_html = f"""
            <div class="preview-container">
                <div class="preview-chapter-title">에필로그</div>
                <div class="preview-content">
                    <p>이 책을 끝까지 읽어주셔서 감사합니다.</p>
                    <p>{core_message}</p>
                    <p><strong>{author} 드림</strong></p>
                </div>
                <hr style="margin: 30px 0;">
                <div class="preview-chapter-title">저자 소개</div>
                <div class="preview-content">
                    <p><strong>{author}</strong></p>
                    <p>{experience}</p>
                </div>
            </div>
            """
        else:
            # 일반 장
            content = st.session_state.drafts.get(selected_key, "내용이 없습니다.")
            # 줄바꿈 처리
            content_html = content.replace('\n\n', '</p><p class="preview-content">').replace('\n', '<br>')
            chapter_title = chapter_options[selected_chapter_idx]
            preview_html = f"""
            <div class="preview-container">
                <div class="preview-chapter-title">{chapter_title}</div>
                <p class="preview-content">{content_html}</p>
            </div>
            """

        st.markdown(preview_html, unsafe_allow_html=True)

        # 페이지 네비게이션
        nav_col1, nav_col2, nav_col3 = st.columns([1, 2, 1])
        with nav_col1:
            if selected_chapter_idx > 0:
                if st.button("< 이전 장", use_container_width=True):
                    st.session_state.chapter_selector = selected_chapter_idx - 1
                    st.rerun()
        with nav_col3:
            if selected_chapter_idx < len(chapter_options) - 1:
                if st.button("다음 장 >", use_container_width=True):
                    st.session_state.chapter_selector = selected_chapter_idx + 1
                    st.rerun()

    st.markdown("---")

    # ===== 3. 다운로드 섹션 =====
    st.markdown("### 📥 다운로드")

    # 원고 생성
    book_manuscript = generate_book_manuscript()
    html_manuscript = generate_html_manuscript()
    txt_manuscript = book_manuscript.replace('=', '-').replace('#', '')
    print_html = generate_print_html()

    # 다운로드 탭
    download_tab1, download_tab2, download_tab3 = st.tabs(["기본 형식", "문서 형식", "인쇄/출판"])

    with download_tab1:
        st.markdown("**기본 다운로드 형식**")
        col1, col2, col3, col4 = st.columns(4)

        with col1:
            st.download_button(
                label="Markdown (.md)",
                data=book_manuscript,
                file_name=f"{st.session_state.selected_title}_원고.md",
                mime="text/markdown",
                use_container_width=True,
                help="Notion, Obsidian, 블로그 등에서 편집 가능"
            )

        with col2:
            st.download_button(
                label="HTML (.html)",
                data=html_manuscript,
                file_name=f"{st.session_state.selected_title}_원고.html",
                mime="text/html",
                use_container_width=True,
                help="웹브라우저에서 바로 열기 가능"
            )

        with col3:
            st.download_button(
                label="텍스트 (.txt)",
                data=txt_manuscript,
                file_name=f"{st.session_state.selected_title}_원고.txt",
                mime="text/plain",
                use_container_width=True,
                help="메모장, 한글 등에서 바로 편집"
            )

        with col4:
            all_data = {
                "book_info": st.session_state.book_info,
                "selected_title": st.session_state.selected_title,
                "generated_toc": st.session_state.generated_toc,
                "drafts": st.session_state.drafts,
                "stats": stats,
                "reading_analysis": reading_analysis,
            }
            st.download_button(
                label="JSON 백업",
                data=json.dumps(all_data, ensure_ascii=False, indent=2),
                file_name=f"{st.session_state.selected_title}_데이터.json",
                mime="application/json",
                use_container_width=True,
                help="나중에 이어서 작업할 때 사용"
            )

    with download_tab2:
        st.markdown("**문서 형식 다운로드**")
        doc_col1, doc_col2, doc_col3 = st.columns(3)

        with doc_col1:
            # DOCX 다운로드
            docx_data = generate_docx_manuscript()
            if docx_data:
                st.download_button(
                    label="Word 문서 (.docx)",
                    data=docx_data,
                    file_name=f"{st.session_state.selected_title}_원고.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                    help="Microsoft Word에서 편집 가능"
                )
            else:
                st.warning("DOCX 생성을 위해 python-docx 패키지가 필요합니다.")
                st.code("pip install python-docx", language="bash")

        with doc_col2:
            # PDF 다운로드
            pdf_data = generate_pdf_manuscript()
            if pdf_data:
                st.download_button(
                    label="PDF 문서 (.pdf)",
                    data=pdf_data,
                    file_name=f"{st.session_state.selected_title}_원고.pdf",
                    mime="application/pdf",
                    use_container_width=True,
                    help="PDF 뷰어에서 바로 열기 가능"
                )
            else:
                st.warning("PDF 생성을 위해 reportlab 패키지가 필요합니다.")
                st.code("pip install reportlab", language="bash")

        with doc_col3:
            st.info("**TIP**: Word 문서는 출판사 제출용으로 적합합니다.")

    with download_tab3:
        st.markdown("**인쇄 및 출판용**")
        print_col1, print_col2 = st.columns(2)

        with print_col1:
            st.download_button(
                label="인쇄용 HTML",
                data=print_html,
                file_name=f"{st.session_state.selected_title}_인쇄용.html",
                mime="text/html",
                use_container_width=True,
                help="A4 페이지에 맞게 최적화, 목차 및 페이지 번호 포함"
            )

        with print_col2:
            st.markdown("""
            **인쇄용 HTML 특징:**
            - A4 페이지 크기에 최적화
            - 클릭 가능한 목차
            - 인쇄 버튼 내장
            - 페이지 나눔 자동 처리
            """)

    st.markdown("---")

    # ===== 4. 공유 섹션 =====
    st.markdown("### 🔗 공유하기")

    share_col1, share_col2 = st.columns(2)

    with share_col1:
        st.markdown("**SNS 공유**")
        title = st.session_state.selected_title or "내 책"
        author = st.session_state.book_info.get("name", "저자")

        share_text = f"'{title}' - {author} 저 | AI 책쓰기 코칭으로 완성한 나만의 책!"
        encoded_text = share_text.replace(' ', '%20').replace("'", '%27')

        # SNS 공유 버튼들
        sns_html = f"""
        <div style="display: flex; gap: 10px; flex-wrap: wrap;">
            <a href="https://twitter.com/intent/tweet?text={encoded_text}" target="_blank"
               style="background: #1DA1F2; color: white; padding: 10px 20px; border-radius: 5px; text-decoration: none; display: inline-block;">
                X (Twitter)
            </a>
            <a href="https://www.facebook.com/sharer/sharer.php?quote={encoded_text}" target="_blank"
               style="background: #4267B2; color: white; padding: 10px 20px; border-radius: 5px; text-decoration: none; display: inline-block;">
                Facebook
            </a>
            <a href="https://www.linkedin.com/shareArticle?mini=true&title={encoded_text}" target="_blank"
               style="background: #0077B5; color: white; padding: 10px 20px; border-radius: 5px; text-decoration: none; display: inline-block;">
                LinkedIn
            </a>
        </div>
        """
        st.markdown(sns_html, unsafe_allow_html=True)

    with share_col2:
        st.markdown("**공유 정보**")
        share_id = generate_share_link()
        st.text_input(
            "공유 코드 (참고용)",
            value=share_id,
            disabled=True,
            help="향후 클라우드 저장 기능 추가 시 사용 예정"
        )
        st.caption("현재는 파일 다운로드 후 직접 공유해 주세요.")

    st.markdown("---")

    # ===== 5. 축하 메시지 =====
    completion_rate = (stats['completed_sections'] / stats['total_sections'] * 100) if stats['total_sections'] > 0 else 0
    if stats['completed_sections'] >= stats['total_sections'] * 0.8:
        st.balloons()
        st.markdown("""
        <div style="background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
                    color: white; padding: 30px; border-radius: 15px; text-align: center;">
            <h2 style="margin-top: 0;">축하합니다! 책 원고가 거의 완성되었습니다!</h2>
            <p style="font-size: 18px;">당신의 이야기가 세상에 나갈 준비가 되었습니다.</p>
            <div style="margin-top: 20px; text-align: left; background: rgba(255,255,255,0.1); padding: 20px; border-radius: 10px;">
                <p><strong>다음 단계:</strong></p>
                <ol style="margin-bottom: 0;">
                    <li>Word 또는 PDF 파일로 다운로드</li>
                    <li>최종 검토 및 교정</li>
                    <li>책 소개서와 함께 출판사에 제출</li>
                </ol>
            </div>
            <p style="margin-bottom: 0; margin-top: 20px; font-size: 20px;">당신의 책 출간을 진심으로 응원합니다!</p>
        </div>
        """, unsafe_allow_html=True)

    st.markdown("---")

    # ===== 6. 새 책 시작 =====
    if "confirm_new_book" not in st.session_state:
        st.session_state.confirm_new_book = False

    if not st.session_state.confirm_new_book:
        if st.button("새 책 시작하기", use_container_width=True):
            st.session_state.confirm_new_book = True
            st.rerun()
    else:
        st.warning("정말로 모든 데이터를 삭제하고 새로 시작할까요?")
        st.caption("삭제 전에 반드시 다운로드를 먼저 하세요!")
        col_confirm1, col_confirm2 = st.columns(2)
        with col_confirm1:
            if st.button("예, 삭제하고 새로 시작", type="primary"):
                for key in list(st.session_state.keys()):
                    del st.session_state[key]
                st.rerun()
        with col_confirm2:
            if st.button("취소"):
                st.session_state.confirm_new_book = False
                st.rerun()

    if st.button("< 이전 단계"):
        st.session_state.current_step = 6
        st.rerun()


def render_chatbot():
    """AI 코치 챗봇"""
    st.markdown("### 💬 AI 책쓰기 코치")

    # 대화 히스토리
    for msg in st.session_state.chat_messages:
        if msg["role"] == "user":
            st.markdown(f"**🧑 나:** {msg['content']}")
        else:
            st.markdown(f"**🤖 코치:** {msg['content']}")
            st.markdown("---")

    # 입력
    with st.form("chat_form", clear_on_submit=True):
        user_input = st.text_area(
            "메시지",
            placeholder="무엇이든 물어보세요!",
            height=80,
            label_visibility="collapsed"
        )
        col1, col2 = st.columns([3, 1])
        with col1:
            send_btn = st.form_submit_button("보내기", use_container_width=True)
        with col2:
            if st.form_submit_button("초기화"):
                st.session_state.chat_messages = []
                st.rerun()

    if send_btn and user_input:
        st.session_state.chat_messages.append({"role": "user", "content": user_input})

        with st.spinner("답변 중..."):
            response = chat_with_coach(
                st.session_state.chat_messages,
                st.session_state.book_info
            )
            if response:
                st.session_state.chat_messages.append({"role": "assistant", "content": response})
        st.rerun()

    # 빠른 질문 (키보드 접근성 개선)
    st.markdown("#### 💡 자주 묻는 질문")
    st.caption("아래 질문 버튼을 클릭하거나 Tab 키로 이동 후 Enter를 눌러 선택할 수 있습니다.")
    quick_questions = [
        "글이 막힐 때 어떻게 하나요?",
        "제목을 더 매력적으로 만들려면?",
    ]
    for idx, q in enumerate(quick_questions):
        if st.button(
            q,
            key=f"quick_{idx}",
            use_container_width=True,
            help=f"빠른 질문: {q}"  # 스크린리더 지원
        ):
            st.session_state.chat_messages.append({"role": "user", "content": q})
            with st.spinner("답변 중..."):
                response = chat_with_coach(st.session_state.chat_messages, st.session_state.book_info)
                if response:
                    st.session_state.chat_messages.append({"role": "assistant", "content": response})
            st.rerun()


# ============================================================
# 🎬 유튜브 모드 관련 함수들
# ============================================================

def render_youtube_mode():
    """유튜브 모드 메인 렌더링"""
    st.markdown('<p class="main-header">🎬 유튜브 모드</p>', unsafe_allow_html=True)
    st.markdown("유튜브 영상의 자막을 추출하여 책으로 변환합니다!")

    col_exit, col_empty = st.columns([1, 3])
    with col_exit:
        if st.button("← 일반 모드로 돌아가기"):
            st.session_state.youtube_mode_active = False
            st.rerun()

    st.markdown("---")

    youtube_step = st.session_state.youtube_step
    steps_info = [("1. URL 입력", "영상 링크"), ("2. 영상 분석", "자막 추출"), ("3. 제목/목차", "AI 생성"), ("4. 첫 번째 글", "글 작성")]

    cols = st.columns(4)
    for i, (step_name, step_desc) in enumerate(steps_info, 1):
        with cols[i-1]:
            if i < youtube_step:
                st.markdown(f"✅ **{step_name}**")
            elif i == youtube_step:
                st.markdown(f"➡️ **{step_name}**")
            else:
                st.markdown(f"⬜ {step_name}")

    st.markdown("---")

    if youtube_step == 1:
        render_youtube_step1_url_input()
    elif youtube_step == 2:
        render_youtube_step2_analyze()
    elif youtube_step == 3:
        render_youtube_step3_title_toc()
    elif youtube_step == 4:
        render_youtube_step4_drafts()


def render_youtube_step1_url_input():
    """유튜브 모드 1단계: URL 입력"""
    st.markdown("### 📋 유튜브 링크 입력")
    st.markdown('<div class="help-box">💡 <b>사용 방법:</b> 유튜브 영상 링크를 붙여넣으세요. 여러 개 입력 시 각각 하나의 Part가 됩니다. 자막이 있는 영상만 지원됩니다.</div>', unsafe_allow_html=True)

    url_input = st.text_area("🔗 유튜브 링크 (여러 개는 줄바꿈으로 구분)", placeholder="https://www.youtube.com/watch?v=...", height=150)

    col1, col2 = st.columns(2)
    with col1:
        if st.button("🔍 영상 확인하기", type="primary", use_container_width=True):
            if not url_input.strip():
                st.error("👆 유튜브 링크를 위에 붙여넣어줘!")
                return
            urls = [u.strip() for u in url_input.strip().split('\n') if u.strip()]
            st.session_state.youtube_urls = urls

            with st.spinner("🔄 영상 정보를 확인하고 있습니다..."):
                videos = []
                for i, url in enumerate(urls):
                    # URL 유효성 검사
                    is_valid, result = validate_youtube_url(url)
                    if not is_valid:
                        st.warning(f"⚠️ 영상 {i+1}: {result}")
                        continue

                    video_id = result  # 유효한 경우 result는 video_id
                    info = get_video_info(url)
                    if info and 'error' not in info:
                        videos.append({'url': url, 'video_id': video_id, 'info': info, 'part_number': i + 1})
                    else:
                        st.warning(f"⚠️ 영상 {i+1}: {info.get('error', '알 수 없는 오류') if info else '알 수 없는 오류'}")
                if videos:
                    st.session_state.youtube_videos = videos
                    st.success(f"✅ {len(videos)}개 영상 확인 완료!")
                    st.rerun()
                else:
                    st.error("😢 처리할 수 있는 영상이 없어. 다른 영상을 넣어줘!")

    with col2:
        st.markdown("**지원하는 URL 형식:**")
        st.markdown("""
        - `youtube.com/watch?v=...`
        - `youtu.be/...`
        - `youtube.com/shorts/...`
        - `m.youtube.com/...`
        """)

    if st.session_state.youtube_videos:
        st.markdown("---")
        st.markdown("### 📺 확인된 영상 목록")
        for video in st.session_state.youtube_videos:
            info = video.get('info', {})
            col_thumb, col_info = st.columns([1, 3])
            with col_thumb:
                if info.get('thumbnail'):
                    st.image(info['thumbnail'], width=200)
            with col_info:
                st.markdown(f"**Part {video['part_number']}: {info.get('title', '제목 없음')}**")
                st.caption(f"채널: {info.get('channel', '?')} | 길이: {info.get('duration_str', '?')} | 조회수: {info.get('view_count_str', '?')}")
            st.markdown("---")

        if st.button("▶️ 자막 추출 시작하기", type="primary", use_container_width=True):
            st.session_state.youtube_step = 2
            st.rerun()


def render_youtube_step2_analyze():
    """유튜브 모드 2단계: 자막 추출"""
    st.markdown("### 📝 자막 추출 및 분석")
    videos = st.session_state.youtube_videos

    if not videos:
        st.warning("👆 먼저 위에서 영상을 추가해줘!")
        if st.button("← 영상 추가하러 가기"):
            st.session_state.youtube_step = 1
            st.rerun()
        return

    if not st.session_state.youtube_transcripts:
        if st.button("🎯 자막 추출 시작", type="primary", use_container_width=True):
            progress_bar = st.progress(0)
            status_text = st.empty()
            transcripts = {}
            for i, video in enumerate(videos):
                video_id = video.get('video_id')
                title = video.get('info', {}).get('title', f'영상 {i+1}')
                status_text.text(f"📝 자막 추출 중... ({i+1}/{len(videos)})")
                progress_bar.progress((i + 0.5) / len(videos))
                transcript, lang_or_error = get_transcript(video_id)
                if transcript:
                    transcripts[video_id] = {'text': transcript, 'language': lang_or_error, 'title': title, 'part_number': video.get('part_number', i+1)}
                    st.success(f"✅ Part {video.get('part_number')}: 완료 ({lang_or_error})")
                else:
                    # 자막 없는 경우 구분 처리
                    if lang_or_error and lang_or_error.startswith("NO_TRANSCRIPT:"):
                        error_msg = lang_or_error.replace("NO_TRANSCRIPT:", "")
                        st.warning(f"⚠️ Part {video.get('part_number')}: {error_msg}")
                    else:
                        st.error(f"❌ Part {video.get('part_number')}: 실패 - {lang_or_error}")
                progress_bar.progress((i + 1) / len(videos))
            if transcripts:
                st.session_state.youtube_transcripts = transcripts
                merged = "".join([f"\n\n=== Part {d['part_number']}: {d['title']} ===\n\n{d['text']}" for d in transcripts.values()])
                st.session_state.youtube_merged_transcript = merged.strip()
                st.success(f"🎉 총 {len(transcripts)}개 영상 자막 추출 완료!")
                st.rerun()
            else:
                st.error("😢 자막을 찾을 수 없어!")
                st.markdown("""
💡 **자막이 없을 때 해결 방법:**
- [Vrew](https://vrew.voyagerx.com/) 또는 [클로바노트](https://clovanote.naver.com/)로 음성을 텍스트로 변환
- 자막이 있는 다른 영상 선택
                """)

    if st.session_state.youtube_transcripts:
        st.markdown("---")
        st.markdown("### 📄 추출된 자막 미리보기")
        for vid_id, data in st.session_state.youtube_transcripts.items():
            with st.expander(f"📺 Part {data['part_number']}: {data['title'][:50]}...", expanded=False):
                st.caption(f"자막 길이: {len(data['text']):,}자 | 언어: {data['language']}")
                st.text_area("자막", value=data['text'][:2000] + ("..." if len(data['text']) > 2000 else ""), height=200, disabled=True, label_visibility="collapsed")

        st.markdown(f"**📊 총 자막 길이: {len(st.session_state.youtube_merged_transcript):,}자**")
        st.markdown("---")

        if not st.session_state.youtube_analysis:
            if st.button("✨ AI로 내용 분석하기", type="primary", use_container_width=True):
                with st.spinner("🤖 AI가 분석 중..."):
                    first_video = list(st.session_state.youtube_transcripts.values())[0]
                    analysis = analyze_youtube_transcript(st.session_state.youtube_merged_transcript, first_video.get('title', ''))
                    if analysis:
                        st.session_state.youtube_analysis = analysis
                        st.rerun()
                    else:
                        st.error("😢 분석이 안 됐어. 다시 해볼까?")

        if st.session_state.youtube_analysis:
            st.markdown("#### 📋 분석 결과")
            st.markdown(st.session_state.youtube_analysis)

        st.markdown("---")
        col1, col2 = st.columns(2)
        with col1:
            if st.button("← 이전 단계"):
                st.session_state.youtube_step = 1
                st.rerun()
        with col2:
            if st.button("▶️ 제목/목차 생성", type="primary", use_container_width=True):
                st.session_state.youtube_step = 3
                st.rerun()


def render_youtube_step3_title_toc():
    """유튜브 모드 3단계: 제목 및 목차"""
    st.markdown("### 📚 제목 및 목차 생성")

    if not st.session_state.youtube_merged_transcript:
        st.warning("👆 먼저 2단계에서 자막을 뽑아줘!")
        if st.button("← 자막 추출하러 가기"):
            st.session_state.youtube_step = 2
            st.rerun()
        return

    col1, col2 = st.columns([2, 1])
    with col1:
        st.markdown("#### 🎯 Step 1: 책 제목 생성")
        if not st.session_state.generated_titles:
            if st.button("✨ AI 제목 10개 생성", use_container_width=True, type="primary"):
                with st.spinner("✨ AI가 제목 생성 중..."):
                    first_video = list(st.session_state.youtube_transcripts.values())[0]
                    result = generate_titles_from_transcript(st.session_state.youtube_merged_transcript, {'title': first_video.get('title', ''), 'channel': ''})
                    if result:
                        st.session_state.generated_titles = result
                        st.rerun()
                    else:
                        st.error("😢 제목을 만들지 못했어. 다시 해볼까?")

        if st.session_state.generated_titles:
            st.markdown("**📝 생성된 제목 후보:**")
            st.markdown(st.session_state.generated_titles)
            st.markdown("---")
            selected = st.text_input("✏️ 최종 제목 선택/입력", value=st.session_state.selected_title, placeholder="위에서 복사하거나 직접 입력")
            if selected:
                st.session_state.selected_title = selected
                st.session_state.book_info["title"] = selected
            if st.button("🔄 제목 다시 생성"):
                st.session_state.generated_titles = ""
                st.rerun()

        st.markdown("---")
        st.markdown("#### 📋 Step 2: 목차 생성")
        if st.session_state.selected_title:
            if not st.session_state.generated_toc:
                if st.button("📋 목차 생성하기", use_container_width=True, type="primary"):
                    with st.spinner("📋 AI가 목차 생성 중... (약 1분)"):
                        video_count = len(st.session_state.youtube_videos)
                        result = generate_toc_from_transcript(st.session_state.youtube_merged_transcript, st.session_state.book_info, video_count)
                        if result:
                            st.session_state.generated_toc = result
                            st.session_state.parsed_toc = parse_toc(result)
                            st.rerun()
                        else:
                            st.error("😢 목차를 만들지 못했어. 다시 해볼까?")

            if st.session_state.generated_toc:
                st.markdown("**📚 생성된 목차:**")
                if st.session_state.parsed_toc:
                    st.success(f"✅ {len(st.session_state.parsed_toc)}개 장 인식!")
                    current_part = None
                    for section in st.session_state.parsed_toc:
                        if section["part"] != current_part:
                            current_part = section["part"]
                            st.markdown(f"**Part {current_part}. {section['part_title']}**")
                        st.markdown(f"  - {section['section_num']}. {section['section_title']}")
                if st.button("🔄 목차 다시 생성"):
                    st.session_state.generated_toc = ""
                    st.session_state.parsed_toc = []
                    st.rerun()
        else:
            st.info("💡 먼저 제목을 선택해주세요.")

    with col2:
        st.markdown("### 💡 도움말")
        st.info("**유튜브 → 책 변환:** 제목 만들기 → 목차 만들기 → 첫 번째 글 쓰기")
        if st.session_state.selected_title:
            st.markdown(f"**제목:** {st.session_state.selected_title}")
        if st.session_state.parsed_toc:
            st.markdown(f"**목차:** {len(st.session_state.parsed_toc)}개 장")

    st.markdown("---")
    col1, col2 = st.columns(2)
    with col1:
        if st.button("← 이전 단계"):
            st.session_state.youtube_step = 2
            st.rerun()
    with col2:
        if st.session_state.parsed_toc:
            if st.button("▶️ 첫 번째 글 쓰기 시작", type="primary", use_container_width=True):
                st.session_state.book_info["youtube_mode"] = True
                st.session_state.book_info["transcript"] = st.session_state.youtube_merged_transcript
                st.session_state.youtube_step = 4
                st.rerun()


def render_youtube_step4_drafts():
    """유튜브 모드 4단계: 첫 번째 글 쓰기"""
    st.markdown("### ✍️ 첫 번째 글 쓰기")
    render_progress_bar()

    parsed_toc = st.session_state.parsed_toc
    drafts = st.session_state.drafts
    transcript = st.session_state.youtube_merged_transcript

    if not parsed_toc:
        st.warning("👆 먼저 3단계에서 목차를 만들어줘!")
        if st.button("← 목차 생성하러 가기"):
            st.session_state.youtube_step = 3
            st.rerun()
        return

    completed_count = len(drafts)
    total_count = len(parsed_toc)
    st.markdown(f'<div class="help-box">💪 진행: {completed_count}/{total_count}개 완료</div>', unsafe_allow_html=True)

    col1, col2 = st.columns([2, 1])
    with col1:
        current_idx = st.session_state.current_section_index
        if current_idx >= len(parsed_toc):
            current_idx = 0
            st.session_state.current_section_index = 0

        current_section = parsed_toc[current_idx]
        section_key = get_section_key(current_section)

        st.markdown(f'<div class="current-section-box"><h3>✍️ 지금 쓸 장</h3><p><b>Part {current_section["part"]}.</b> {current_section["part_title"]}</p><p style="font-size: 1.3rem;"><b>{current_section["section_num"]}. {current_section["section_title"]}</b></p></div>', unsafe_allow_html=True)

        if section_key in drafts:
            st.success("✅ 이미 작성됨!")
            edited_draft = st.text_area("작성된 내용", value=drafts[section_key], height=400)
            if edited_draft != drafts[section_key]:
                drafts[section_key] = edited_draft
            st.caption(f"📊 글자 수: {len(edited_draft.replace(' ', '').replace(chr(10), ''))}자")
            col_a, col_b = st.columns(2)
            with col_a:
                if current_idx > 0 and st.button("⬅️ 이전 장"):
                    st.session_state.current_section_index = current_idx - 1
                    st.rerun()
            with col_b:
                if current_idx < len(parsed_toc) - 1 and st.button("➡️ 다음 장", type="primary"):
                    st.session_state.current_section_index = current_idx + 1
                    st.rerun()
        else:
            if st.button("✨ AI가 글 써줘!", use_container_width=True, type="primary"):
                section_info = {"part_number": current_section["part"], "part_title": current_section["part_title"], "section_number": current_section["section_num"], "section_title": current_section["section_title"]}
                with st.spinner("✨ 글을 쓰고 있어요..."):
                    # Part별 자막 추출 (다중 영상인 경우 해당 Part 자막 사용)
                    part_transcript = get_part_transcript(transcript, current_section["part"])
                    result = generate_draft_from_transcript(st.session_state.book_info, section_info, part_transcript)
                    if result:
                        st.session_state.drafts[section_key] = result
                        st.success("✅ 완료!")
                        st.rerun()
                    else:
                        st.error("😢 잠깐 문제가 생겼어. 다시 해볼까?")

    with col2:
        st.markdown("### 📋 진행 현황")
        current_part_num = current_section['part']
        current_part_sections = [s for s in parsed_toc if s['part'] == current_part_num]
        part_completed = sum(1 for s in current_part_sections if f"{s['section_num']}_{s['section_title']}" in drafts)
        st.markdown(f"**Part {current_part_num}** ({part_completed}/{len(current_part_sections)})")

        for idx, section in enumerate(parsed_toc):
            if section['part'] != current_part_num:
                continue
            key = f"{section['section_num']}_{section['section_title']}"
            prefix = "➡️" if idx == current_idx else ("✅" if key in drafts else "⬜")
            if st.button(f"{prefix} {section['section_num']}. {section['section_title'][:12]}...", key=f"yt_jump_{idx}", use_container_width=True):
                st.session_state.current_section_index = idx
                st.rerun()

        st.markdown("---")
        all_unfinished = [s for s in parsed_toc if f"{s['section_num']}_{s['section_title']}" not in drafts]
        if all_unfinished:
            st.markdown(f"**남은 장: {len(all_unfinished)}개**")
            if st.button("🚀 전체 자동 생성", use_container_width=True, type="primary"):
                progress_bar = st.progress(0)
                status_text = st.empty()
                for i, section in enumerate(all_unfinished):
                    key = f"{section['section_num']}_{section['section_title']}"
                    section_info = {"part_number": section["part"], "part_title": section["part_title"], "section_number": section["section_num"], "section_title": section["section_title"]}
                    status_text.text(f"✍️ {section['section_num']}. {section['section_title'][:20]}... 작성 중")
                    # Part별 자막 추출
                    part_transcript = get_part_transcript(transcript, section["part"])
                    result = generate_draft_from_transcript(st.session_state.book_info, section_info, part_transcript)
                    if result:
                        st.session_state.drafts[key] = result
                    progress_bar.progress((i + 1) / len(all_unfinished))
                status_text.empty()
                st.balloons()
                st.success("🎉 모든 첫 번째 글 완성!")
                st.rerun()

    st.markdown("---")
    col1, col2 = st.columns(2)
    with col1:
        if st.button("← 이전 단계"):
            st.session_state.youtube_step = 3
            st.rerun()
    with col2:
        if st.session_state.drafts and st.button("✅ 일반 모드로 전환", type="primary", use_container_width=True):
            # 유튜브 모드에서 일반 모드로 전환
            st.session_state.youtube_mode_active = False
            st.session_state.previous_mode = "youtube"

            # 유튜브 모드 플래그와 자막 저장
            st.session_state.book_info["youtube_mode"] = True
            st.session_state.book_info["transcript"] = st.session_state.get("youtube_merged_transcript", "")

            # 초안이 있으면 다운로드 단계로, 없으면 첫 번째 글 쓰기 단계로
            if len(st.session_state.drafts) >= len(st.session_state.parsed_toc) * 0.5:
                st.session_state.current_step = 7
            else:
                st.session_state.current_step = 4

            st.rerun()


# ============================================================
# 📱 채팅 모드 (초등학생용 대화형 책쓰기)
# ============================================================

# 채팅 모드 단계별 질문 정의 (40~50대 CEO/전문직 대상)
CHAT_MODE_STEPS = [
    {
        "step": 0,
        "name": "저자명",
        "emoji": "1️⃣",
        "question": "안녕하세요! 저는 책쓰기 전문 코치 '북코치'입니다.\n\n오늘부터 작가님의 책 집필을 도와드리겠습니다.\n\n먼저 저자명(필명)을 알려주시겠어요?",
        "key": "name",
        "placeholder": "저자명을 입력해 주세요 (예: 홍길동, 김작가)",
        "examples": ["홍길동", "김대표", "이코치", "박박사", "최전문"],
        "help_text": "본명 또는 필명을 입력해 주세요.",
    },
    {
        "step": 1,
        "name": "책 주제",
        "emoji": "2️⃣",
        "question": "{name}님, 반갑습니다! 함께 좋은 책을 만들어 가겠습니다.\n\n어떤 주제로 책을 쓰고 싶으신가요?\n\n그동안 쌓아오신 경험, 전문 지식, 또는 전하고 싶은 메시지가 무엇인지 알려주세요.",
        "key": "topic",
        "placeholder": "책 주제를 입력해 주세요 (예: 30년 경영 노하우를 담은 리더십 이야기)",
        "examples": ["30년 경영 노하우를 담은 리더십 이야기", "의사가 알려주는 건강한 삶의 비결", "변호사의 협상 전략", "투자 전문가의 자산 관리 비법", "창업 10년, 실패와 성공의 기록", "교육자가 전하는 자녀 교육법"],
        "help_text": "전문 분야, 경험, 노하우 등 어떤 주제든 책으로 만들 수 있습니다.",
    },
    {
        "step": 2,
        "name": "타겟 독자",
        "emoji": "3️⃣",
        "question": "'{topic}'이라는 주제, 훌륭합니다!\n\n이 책은 어떤 분들이 읽으면 좋을까요?\n\n타겟 독자를 구체적으로 정의하면 더 효과적인 책이 됩니다.",
        "key": "target_reader",
        "placeholder": "타겟 독자를 입력해 주세요 (예: 중견기업 임원 및 예비 창업가)",
        "examples": ["중견기업 임원 및 예비 창업가", "40-50대 건강 관리에 관심 있는 직장인", "사업 성장을 고민하는 CEO", "재테크를 시작하려는 전문직 종사자", "자녀 교육에 고민이 있는 학부모"],
        "help_text": "나이, 직업, 관심사 등을 구체적으로 설정하시면 좋습니다.",
    },
    {
        "step": 3,
        "name": "제목 생성",
        "emoji": "4️⃣",
        "question": "좋습니다! 이제 준비가 되었습니다.\n\n📚 책 정보 정리:\n• 주제: {topic}\n• 타겟 독자: {target_reader}\n\nAI가 베스트셀러 공식을 적용하여 책 제목 10가지를 추천해 드리겠습니다.\n\n아래 버튼을 클릭해 주세요.",
        "key": "title",
        "action": "generate_titles",
        "loading_message": "📝 베스트셀러 제목 공식을 분석하고 있습니다...\n\n약 30초 정도 소요됩니다. 잠시만 기다려 주세요.",
        "error_message": "죄송합니다. 일시적인 오류가 발생했습니다.\n\n인터넷 연결을 확인하신 후 다시 시도해 주세요.",
    },
    {
        "step": 4,
        "name": "목차 생성",
        "emoji": "5️⃣",
        "question": "제목이 결정되었습니다!\n\n「{title}」 훌륭한 선택입니다.\n\n이제 책의 목차를 구성할 차례입니다.\n\n📖 5부 40장 구조로 체계적인 목차를 설계해 드리겠습니다.\n\nWhy-What-How-Do-Future 프레임워크를 적용합니다.",
        "key": "toc",
        "action": "generate_toc",
        "loading_message": "📋 체계적인 목차 구조를 설계하고 있습니다.\n\n약 1분 정도 소요됩니다. 잠시만 기다려 주세요.",
        "error_message": "죄송합니다. 일시적인 오류가 발생했습니다.\n\n다시 시도해 주세요.",
    },
    {
        "step": 5,
        "name": "첫 번째 챕터",
        "emoji": "6️⃣",
        "question": "목차도 완성되었습니다! 거의 다 왔습니다.\n\n이제 본격적으로 원고를 작성할 시간입니다. ✍️\n\nAI가 첫 번째 챕터 초안을 작성해 드리겠습니다.\n\n작성된 초안을 검토하시고 필요한 부분을 수정해 주세요.",
        "key": "draft",
        "action": "generate_draft",
        "loading_message": "✍️ 첫 번째 챕터를 작성하고 있습니다.\n\n약 1분 정도 소요됩니다. 잠시만 기다려 주세요.",
        "error_message": "죄송합니다. 일시적인 오류가 발생했습니다.\n\n다시 시도해 주세요.",
    },
]


def get_step_progress_text(step):
    """단계별 진행 상황 텍스트 반환 (초등학생용)"""
    progress_texts = {
        0: "🌱 시작! 이름을 알려줘!",
        1: "📝 좋아하는 것을 말해줘!",
        2: "👥 누가 읽을지 정해보자!",
        3: "✨ 제목을 만들자!",
        4: "📋 목차를 만들자!",
        5: "✍️ 첫 글을 써보자!",
    }
    return progress_texts.get(step, "🚀 진행 중...")


def get_chat_mode_greeting(name):
    """이름에 맞는 호칭 반환"""
    if name:
        last_char = name[-1]
        if '가' <= last_char <= '힣':
            code = ord(last_char) - ord('가')
            if code % 28 == 0:
                return f"{name}야"
            else:
                return f"{name}아"
    return name


def render_chat_mode():
    """채팅 모드 화면 렌더링 (초등학생 친화적 개선)"""
    step = st.session_state.chat_mode_step
    history = st.session_state.chat_mode_history
    data = st.session_state.chat_mode_data

    # 헤더 - 더 친근하게
    st.markdown("""
    <div class="chat-mode-header">
        📱 북코치와 함께 책 만들기!
    </div>
    """, unsafe_allow_html=True)

    total_steps = len(CHAT_MODE_STEPS)
    completed = min(step, total_steps)
    progress_text = get_step_progress_text(step) if step < total_steps else "🎉 완성!"

    # 진행 상황을 더 친근하게 표시
    st.markdown(f"""
    <div class="chat-progress">
        {progress_text}<br>
        <small style="font-size: 0.9rem;">⭐ {completed}/{total_steps} 단계 완료!</small>
    </div>
    """, unsafe_allow_html=True)

    # 진행바를 더 크고 예쁘게
    st.progress(completed / total_steps)

    # 단계 표시 아이콘 (가로로 배치)
    step_cols = st.columns(6)
    step_emojis = ["1️⃣", "2️⃣", "3️⃣", "4️⃣", "5️⃣", "6️⃣"]
    step_names = ["이름", "주제", "독자", "제목", "목차", "글쓰기"]
    for i, col in enumerate(step_cols):
        with col:
            if i < completed:
                st.markdown(f"<div style='text-align:center;'>✅<br><small>{step_names[i]}</small></div>", unsafe_allow_html=True)
            elif i == step:
                st.markdown(f"<div style='text-align:center;font-weight:bold;color:#1976D2;'>👉{step_emojis[i]}<br><small>{step_names[i]}</small></div>", unsafe_allow_html=True)
            else:
                st.markdown(f"<div style='text-align:center;opacity:0.4;'>{step_emojis[i]}<br><small>{step_names[i]}</small></div>", unsafe_allow_html=True)

    # 버튼 행 - 처음부터 다시 + 이전 단계
    col_back, col_restart, col_spacer = st.columns([1, 1, 2])
    with col_back:
        if step > 0 and step < total_steps:
            if st.button("⬅️ 이전으로", use_container_width=True, key="back_btn"):
                st.session_state.chat_mode_step = step - 1
                # 이전 단계 데이터 유지하면서 현재 단계 데이터만 삭제
                current_step_key = CHAT_MODE_STEPS[step]["key"] if step < len(CHAT_MODE_STEPS) else None
                if current_step_key and current_step_key in data:
                    del data[current_step_key]
                # 히스토리에서 마지막 대화 제거 (AI 응답 + 사용자 응답)
                if len(history) >= 2:
                    st.session_state.chat_mode_history = history[:-2]
                st.rerun()
    with col_restart:
        if st.button("🔄 처음부터 다시", use_container_width=True, key="restart_btn"):
            st.session_state.chat_mode_step = 0
            st.session_state.chat_mode_history = []
            st.session_state.chat_mode_data = {}
            st.rerun()

    st.markdown("---")

    for msg in history:
        if msg["role"] == "ai":
            st.markdown(f"""
            <div class="chat-bubble chat-bubble-ai">
                🤖 {msg["content"]}
            </div>
            """, unsafe_allow_html=True)
        else:
            st.markdown(f"""
            <div class="chat-bubble chat-bubble-user">
                😊 {msg["content"]}
            </div>
            """, unsafe_allow_html=True)

    if step < len(CHAT_MODE_STEPS):
        current_step_info = CHAT_MODE_STEPS[step]

        # 질문 템플릿 처리 (변수 치환)
        question = current_step_info["question"]
        if "{name}" in question:
            name = data.get("name", "친구")
            greeting = get_chat_mode_greeting(name)
            question = question.replace("{name}아/야", greeting).replace("{name}", name)
        if "{topic}" in question:
            question = question.replace("{topic}", data.get("topic", "주제"))
        if "{target_reader}" in question:
            question = question.replace("{target_reader}", data.get("target_reader", "독자"))
        if "{title}" in question:
            question = question.replace("{title}", data.get("title", "제목"))

        # 현재 질문 표시
        if not history or history[-1]["role"] != "ai" or history[-1]["content"] != question:
            st.markdown(f"""
            <div class="chat-bubble chat-bubble-ai">
                🤖 {question}
            </div>
            """, unsafe_allow_html=True)

        # 도움말 표시 (있는 경우)
        if "help_text" in current_step_info:
            st.info(f"💡 힌트: {current_step_info['help_text']}")

        if "action" in current_step_info:
            action = current_step_info["action"]
            loading_msg = current_step_info.get("loading_message", "잠시만 기다려줘...")
            error_msg = current_step_info.get("error_message", "앗! 오류가 발생했어. 다시 시도해줘!")

            if action == "generate_titles":
                if st.button("✨ 제목 만들어줘!", use_container_width=True, type="primary", key="gen_title_btn"):
                    with st.spinner(loading_msg):
                        try:
                            book_info = {
                                "name": data.get("name", ""),
                                "topic": data.get("topic", ""),
                                "target_reader": data.get("target_reader", ""),
                                "core_message": f"{data.get('topic', '')}에 대한 재미있는 이야기",
                                "experience": "",
                                "tone": "친절하고 따뜻한",
                            }
                            result = generate_titles(book_info)
                            if result:
                                data["generated_titles"] = result
                                history.append({"role": "ai", "content": question})
                                history.append({"role": "user", "content": "제목 만들어줘!"})
                                history.append({"role": "ai", "content": f"짜잔! 🎉 제목 후보들이 나왔어!"})
                                st.rerun()
                            else:
                                st.error(error_msg)
                        except Exception as e:
                            st.error(f"{error_msg}\n\n(오류 내용: {str(e)[:100]})")

                if "generated_titles" in data:
                    st.markdown("""
                    <div class="result-card">
                        <h3>🌟 제목 후보들이야! 마음에 드는 걸 골라봐!</h3>
                    </div>
                    """, unsafe_allow_html=True)
                    st.markdown(data["generated_titles"])

                    # 다시 만들기 버튼
                    if st.button("🔄 다른 제목 더 보고 싶어!", use_container_width=True, key="regenerate_title_btn"):
                        del data["generated_titles"]
                        st.rerun()

                    st.markdown("---")
                    st.markdown("### ✏️ 위에서 마음에 드는 제목을 골라서 아래에 써줘!")
                    st.markdown("**(그대로 복사해도 되고, 네 마음대로 바꿔도 돼!)**")
                    selected_title = st.text_input(
                        "제목 입력",
                        placeholder="예: 마인크래프트 마스터가 되는 법",
                        label_visibility="collapsed",
                        key="title_input"
                    )

                    if st.button("👍 이 제목으로 할래!", use_container_width=True, type="primary", key="select_title_btn"):
                        if selected_title and selected_title.strip():
                            data["title"] = selected_title.strip()
                            history.append({"role": "user", "content": f"제목: {selected_title}"})
                            history.append({"role": "ai", "content": f"와~ '{selected_title}' 정말 멋진 제목이야! 🌟 이 제목으로 책을 쓰면 많은 사람들이 읽고 싶어할 거야!"})
                            st.session_state.chat_mode_step = step + 1
                            st.rerun()
                        else:
                            st.warning("앗! 제목을 써줘야 해! ✍️ 위에서 마음에 드는 제목을 골라서 적어봐!")

            elif action == "generate_toc":
                if st.button("📋 목차 만들어줘!", use_container_width=True, type="primary", key="gen_toc_btn"):
                    with st.spinner(loading_msg):
                        try:
                            book_info = {
                                "name": data.get("name", ""),
                                "topic": data.get("topic", ""),
                                "target_reader": data.get("target_reader", ""),
                                "core_message": f"{data.get('topic', '')}에 대한 재미있는 이야기",
                                "title": data.get("title", ""),
                                "experience": "",
                                "tone": "친절하고 따뜻한",
                            }
                            result = generate_toc(book_info)
                            if result:
                                data["generated_toc"] = result
                                data["parsed_toc"] = parse_toc(result)
                                history.append({"role": "ai", "content": question})
                                history.append({"role": "user", "content": "목차 만들어줘!"})
                                history.append({"role": "ai", "content": f"목차가 완성됐어! 📚 아래를 확인해봐!"})
                                st.rerun()
                            else:
                                st.error(error_msg)
                        except Exception as e:
                            st.error(f"{error_msg}\n\n(오류 내용: {str(e)[:100]})")

                if "generated_toc" in data:
                    st.markdown("""
                    <div class="result-card">
                        <h3>📋 책 목차가 완성됐어!</h3>
                        <p>이 순서대로 책을 쓸 거야! 어때, 재미있어 보여? 👀</p>
                    </div>
                    """, unsafe_allow_html=True)

                    # 목차를 접을 수 있게 표시 (초등학생도 쉽게)
                    with st.expander("📖 목차 전체 보기 (클릭해서 펼쳐봐!)", expanded=True):
                        st.markdown(data["generated_toc"][:2000] + "..." if len(data["generated_toc"]) > 2000 else data["generated_toc"])

                    st.markdown("---")
                    st.markdown("### 🤔 목차가 마음에 들어?")

                    col_yes, col_no = st.columns(2)
                    with col_yes:
                        if st.button("👍 좋아! 다음으로!", use_container_width=True, type="primary", key="confirm_toc_btn"):
                            history.append({"role": "user", "content": "목차 좋아!"})
                            history.append({"role": "ai", "content": "좋았어! 🎉 이제 진짜 책 내용을 쓸 차례야!"})
                            st.session_state.chat_mode_step = step + 1
                            st.rerun()
                    with col_no:
                        if st.button("🔄 다시 만들어줘", use_container_width=True, key="regenerate_toc_btn"):
                            del data["generated_toc"]
                            if "parsed_toc" in data:
                                del data["parsed_toc"]
                            st.rerun()

            elif action == "generate_draft":
                parsed_toc = data.get("parsed_toc", [])

                if parsed_toc:
                    first_section = parsed_toc[0] if parsed_toc else {"section_num": "1", "section_title": "시작", "part": 1, "part_title": "Part 1"}

                    # 첫 번째 섹션 정보 표시
                    st.markdown(f"""
                    <div class="result-card">
                        <h4>📝 첫 번째로 쓸 이야기</h4>
                        <p><strong>제목:</strong> {first_section.get('section_title', '시작')}</p>
                    </div>
                    """, unsafe_allow_html=True)

                    if st.button("✍️ 첫 번째 이야기 써줘!", use_container_width=True, type="primary", key="gen_draft_btn"):
                        with st.spinner(loading_msg):
                            try:
                                book_info = {
                                    "name": data.get("name", ""),
                                    "topic": data.get("topic", ""),
                                    "target_reader": data.get("target_reader", ""),
                                    "core_message": f"{data.get('topic', '')}에 대한 재미있는 이야기",
                                    "title": data.get("title", ""),
                                    "experience": "",
                                    "tone": "친절하고 따뜻한",
                                }
                                section_info = {
                                    "part_number": first_section.get("part", 1),
                                    "part_title": first_section.get("part_title", "Part 1"),
                                    "section_number": first_section.get("section_num", "1"),
                                    "section_title": first_section.get("section_title", "시작"),
                                    "core_message": "",
                                    "examples": "",
                                }
                                result = generate_draft(book_info, section_info)
                                if result:
                                    data["first_draft"] = result
                                    history.append({"role": "ai", "content": question})
                                    history.append({"role": "user", "content": "첫 번째 이야기 써줘!"})
                                    history.append({"role": "ai", "content": f"첫 번째 이야기가 완성됐어! 📝 아래에서 읽어봐!"})
                                    st.session_state.chat_mode_step = step + 1
                                    st.rerun()
                                else:
                                    st.error(error_msg)
                            except Exception as e:
                                st.error(f"{error_msg}\n\n(오류 내용: {str(e)[:100]})")

                    if "first_draft" in data:
                        st.markdown("""
                        <div class="result-card">
                            <h3>📝 첫 번째 이야기가 완성됐어!</h3>
                            <p>읽어보고 어떤지 말해줘! 😊</p>
                        </div>
                        """, unsafe_allow_html=True)
                        with st.expander("📖 이야기 전체 보기 (클릭해서 펼쳐봐!)", expanded=True):
                            st.markdown(data["first_draft"])
                else:
                    st.warning("앗! 목차가 없어! 😅 '⬅️ 이전으로' 버튼을 눌러서 목차를 먼저 만들어줘!")

        else:
            # 입력 UI (예시 버튼 + 직접 입력)
            examples = current_step_info.get("examples", [])

            if examples:
                st.markdown("### 💡 예시 중에서 골라봐! (클릭하면 바로 입력돼)")

                # 예시가 많으면 2줄로 배치
                if len(examples) <= 4:
                    cols = st.columns(len(examples))
                    for idx, example in enumerate(examples):
                        with cols[idx]:
                            if st.button(f"👆 {example}", key=f"example_{step}_{idx}", use_container_width=True):
                                data[current_step_info["key"]] = example
                                history.append({"role": "ai", "content": question})
                                history.append({"role": "user", "content": example})
                                st.session_state.chat_mode_step = step + 1
                                st.rerun()
                else:
                    # 2줄로 배치
                    row1 = examples[:3]
                    row2 = examples[3:]
                    cols1 = st.columns(len(row1))
                    for idx, example in enumerate(row1):
                        with cols1[idx]:
                            if st.button(f"👆 {example}", key=f"example_{step}_{idx}", use_container_width=True):
                                data[current_step_info["key"]] = example
                                history.append({"role": "ai", "content": question})
                                history.append({"role": "user", "content": example})
                                st.session_state.chat_mode_step = step + 1
                                st.rerun()
                    cols2 = st.columns(len(row2))
                    for idx, example in enumerate(row2):
                        with cols2[idx]:
                            if st.button(f"👆 {example}", key=f"example_{step}_{idx + 3}", use_container_width=True):
                                data[current_step_info["key"]] = example
                                history.append({"role": "ai", "content": question})
                                history.append({"role": "user", "content": example})
                                st.session_state.chat_mode_step = step + 1
                                st.rerun()

            st.markdown("---")
            st.markdown("### ✏️ 아니면 직접 써줘!")
            user_input = st.text_input(
                "입력",
                placeholder=current_step_info.get("placeholder", "여기에 입력해줘!"),
                label_visibility="collapsed",
                key=f"input_{step}"
            )

            if st.button("👍 다음으로 가자!", use_container_width=True, type="primary", key=f"next_{step}"):
                if user_input and user_input.strip():
                    data[current_step_info["key"]] = user_input.strip()
                    history.append({"role": "ai", "content": question})
                    history.append({"role": "user", "content": user_input})
                    st.session_state.chat_mode_step = step + 1
                    st.rerun()
                else:
                    st.warning(f"앗! 아무것도 안 썼어! 😅 {current_step_info.get('placeholder', '뭔가를 입력해줘!')}")

    else:
        # 완료 화면 - 축하 메시지
        st.balloons()

        name = data.get('name', '친구')
        greeting = get_chat_mode_greeting(name)

        st.markdown(f"""
        <div class="chat-bubble chat-bubble-ai" style="font-size: 1.5rem; text-align: center;">
            🎉🎉🎉 축하해, {greeting}! 🎉🎉🎉<br><br>
            네가 진짜 <b>책 쓰기</b>를 시작했어!<br><br>
            이제 너도 <b>작가</b>야! 👏👏👏<br><br>
            정말정말 대단해! 🌟
        </div>
        """, unsafe_allow_html=True)

        # 책 정보 카드
        st.markdown(f"""
        <div class="result-card">
            <h3 style="text-align: center;">📚 {name} 작가님의 첫 번째 책!</h3>
        </div>
        """, unsafe_allow_html=True)

        # 책 정보를 예쁘게 표시
        st.markdown(f"""
        <div style="background: linear-gradient(135deg, #E3F2FD 0%, #BBDEFB 100%); padding: 1.5rem; border-radius: 15px; margin: 1rem 0;">
            <p style="font-size: 1.3rem; margin: 0.5rem 0;"><strong>✍️ 작가:</strong> {data.get('name', '미정')}</p>
            <p style="font-size: 1.3rem; margin: 0.5rem 0;"><strong>📖 주제:</strong> {data.get('topic', '미정')}</p>
            <p style="font-size: 1.3rem; margin: 0.5rem 0;"><strong>👥 독자:</strong> {data.get('target_reader', '미정')}</p>
            <p style="font-size: 1.5rem; margin: 0.5rem 0; color: #1976D2;"><strong>📕 제목:</strong> {data.get('title', '미정')}</p>
        </div>
        """, unsafe_allow_html=True)

        # 첫 번째 이야기 미리보기
        if "first_draft" in data:
            st.markdown("### 📝 네가 쓴 첫 번째 이야기")
            with st.expander("👀 클릭해서 읽어보기", expanded=False):
                st.markdown(data["first_draft"])

        st.markdown("---")

        # 다음 단계 안내
        st.markdown("""
        <div class="result-card">
            <h3>🚀 앞으로 뭘 할 수 있을까?</h3>
            <p>첫 번째 이야기를 썼으니까, 이제 나머지 이야기도 계속 쓸 수 있어!</p>
        </div>
        """, unsafe_allow_html=True)

        col_a, col_b = st.columns(2)
        with col_a:
            if st.button("📝 더 많은 이야기 쓰러 가기!", use_container_width=True, type="primary"):
                # 채팅 모드 데이터를 일반 모드로 완전히 복사
                st.session_state.book_info = {
                    "name": data.get("name", ""),
                    "topic": data.get("topic", ""),
                    "target_reader": data.get("target_reader", ""),
                    "core_message": f"{data.get('topic', '')}에 대한 이야기",
                    "experience": "",
                    "tone": "친절하고 따뜻한",
                    "title": data.get("title", ""),  # 제목도 book_info에 저장
                }
                st.session_state.selected_title = data.get("title", "")

                # 생성된 제목 후보도 복사
                if "generated_titles" in data:
                    st.session_state.generated_titles = data["generated_titles"]

                # 목차 복사
                if "generated_toc" in data:
                    st.session_state.generated_toc = data["generated_toc"]
                    st.session_state.parsed_toc = data.get("parsed_toc", [])

                # 첫 번째 초안 복사
                if "first_draft" in data and st.session_state.parsed_toc:
                    first_section = st.session_state.parsed_toc[0]
                    key = f"{first_section['section_num']}_{first_section['section_title']}"
                    st.session_state.drafts[key] = data["first_draft"]
                    # 현재 섹션 인덱스를 1로 설정 (첫 번째 장 완료했으므로 다음 장으로)
                    st.session_state.current_section_index = 1 if len(st.session_state.parsed_toc) > 1 else 0

                # 채팅 모드 비활성화
                st.session_state.chat_mode_active = False
                st.session_state.previous_mode = "chat"

                # 적절한 단계로 이동 (데이터 상태에 따라)
                if st.session_state.drafts:
                    st.session_state.current_step = 4
                elif st.session_state.parsed_toc:
                    st.session_state.current_step = 4
                elif st.session_state.selected_title:
                    st.session_state.current_step = 3
                else:
                    st.session_state.current_step = 2

                st.success("✅ 좋았어! 일반 모드에서 계속 책을 써보자!")
                st.rerun()

        with col_b:
            if st.button("🔄 다른 책 만들기", use_container_width=True):
                st.session_state.chat_mode_step = 0
                st.session_state.chat_mode_history = []
                st.session_state.chat_mode_data = {}
                st.rerun()

        # 응원 메시지
        st.markdown("""
        <div style="background: #FFF3E0; padding: 1rem; border-radius: 10px; margin-top: 1rem; text-align: center;">
            <p style="font-size: 1.1rem; margin: 0;">
                💪 <b>포기하지 마!</b> 조금씩 쓰다 보면 어느새 책 한 권이 완성될 거야!<br>
                북코치가 항상 응원할게! 화이팅! 🔥
            </p>
        </div>
        """, unsafe_allow_html=True)


# ============================================
# UX 개선 헬퍼 함수들
# ============================================

# 친근한 로딩 메시지 리스트
UX_LOADING_MESSAGES = [
    "열심히 쓰는 중이에요... 조금만 기다려줘!",
    "AI가 열심히 생각하고 있어요...",
    "멋진 글이 만들어지고 있어요!",
    "거의 다 됐어요, 조금만!",
    "창작의 마법이 일어나고 있어요...",
    "좋은 글이 나올 거예요!",
    "AI 작가가 집중하고 있어요...",
]

# 친근한 팁 메시지 리스트
UX_LOADING_TIPS = [
    "글을 쓰는 동안 잠깐 스트레칭을 해보세요!",
    "좋은 책은 천천히 만들어져요.",
    "완벽하지 않아도 괜찮아요. 나중에 고칠 수 있어요!",
    "하루에 조금씩 쓰면 어느새 책 한 권이 완성돼요.",
    "쓰다가 모르는 게 있으면 AI 코치에게 물어보세요!",
]


def render_ux_loading(progress_percent: int = 50, task_name: str = "글 만들기", show_tips: bool = True):
    """향상된 로딩 애니메이션 표시

    Args:
        progress_percent: 진행률 (0-100)
        task_name: 현재 작업 이름
        show_tips: 팁 표시 여부
    """
    import random

    message = random.choice(UX_LOADING_MESSAGES)
    tip = random.choice(UX_LOADING_TIPS) if show_tips else ""

    loading_html = f"""
    <div class="ux-loading-container" id="ux-loading-box">
        <div class="ux-loading-emoji">✨</div>
        <div class="ux-loading-spinner"></div>
        <div class="ux-loading-text">{task_name} 중...</div>
        <div class="ux-loading-subtext">{message}</div>
        <div class="ux-loading-progress">
            <div class="ux-loading-progress-bar" style="width: {progress_percent}%;"></div>
        </div>
        <div class="ux-loading-percent">{progress_percent}%</div>
        {"<div class='ux-loading-tips'>" + tip + "</div>" if show_tips and tip else ""}
    </div>
    """
    return st.markdown(loading_html, unsafe_allow_html=True)


def render_ux_toast(message: str, toast_type: str = "success", icon: str = None):
    """토스트 스타일 알림 표시

    Args:
        message: 표시할 메시지
        toast_type: 타입 (success, error, info, warning)
        icon: 아이콘 (기본값은 타입에 따라 자동 선택)
    """
    icons = {
        "success": "✅",
        "error": "❌",
        "info": "ℹ️",
        "warning": "⚠️"
    }
    selected_icon = icon or icons.get(toast_type, "✅")

    toast_html = f"""
    <div class="ux-toast ux-toast-{toast_type}">
        <span class="ux-toast-icon">{selected_icon}</span>
        <span class="ux-toast-message">{message}</span>
    </div>
    <script>
        setTimeout(function() {{
            var toast = document.querySelector('.ux-toast');
            if (toast) toast.style.display = 'none';
        }}, 3000);
    </script>
    """
    st.markdown(toast_html, unsafe_allow_html=True)


def render_ux_celebration(title: str, message: str, emojis: list = None):
    """화려한 축하 효과 표시

    Args:
        title: 축하 제목
        message: 축하 메시지
        emojis: 표시할 이모지 리스트 (기본값: 파티 이모지)
    """
    if emojis is None:
        emojis = ["🎉", "🎊", "✨"]

    emoji_html = "".join([f'<span class="ux-celebration-emoji">{e}</span>' for e in emojis])

    celebration_html = f"""
    <div class="ux-celebration" id="ux-celebration-box">
        <div class="ux-celebration-content">
            {emoji_html}
            <div class="ux-celebration-title">{title}</div>
            <div class="ux-celebration-message">{message}</div>
        </div>
    </div>
    """
    st.markdown(celebration_html, unsafe_allow_html=True)


def render_ux_breadcrumb(current_step: int):
    """현재 위치를 보여주는 브레드크럼 네비게이션

    Args:
        current_step: 현재 단계 (1-7)
    """
    steps = [
        (1, "정보 입력"),
        (2, "제목 만들기"),
        (3, "목차 만들기"),
        (4, "글 쓰기"),
        (5, "책 소개서"),
        (6, "홍보 페이지"),
        (7, "다운로드"),
    ]

    items_html = []
    for i, (num, name) in enumerate(steps):
        if num < current_step:
            items_html.append(f'<span class="ux-breadcrumb-item completed">✓ {name}</span>')
        elif num == current_step:
            items_html.append(f'<span class="ux-breadcrumb-item active">📍 {name}</span>')
        else:
            items_html.append(f'<span class="ux-breadcrumb-item">{num}. {name}</span>')

        if i < len(steps) - 1:
            items_html.append('<span class="ux-breadcrumb-separator">→</span>')

    breadcrumb_html = f"""
    <div class="ux-breadcrumb" role="navigation" aria-label="현재 위치">
        {''.join(items_html)}
    </div>
    """
    st.markdown(breadcrumb_html, unsafe_allow_html=True)


def render_ux_step_tip(step: int):
    """단계별 팁 표시

    Args:
        step: 현재 단계 (1-7)
    """
    tips = {
        1: """
        <ul>
            <li><strong>이름</strong>은 책 표지에 나와요. 본명이나 필명을 써보세요!</li>
            <li><strong>책 주제</strong>는 한 문장으로 간단히 써도 돼요.</li>
            <li><strong>누가 읽을까?</strong> - 나이, 직업, 관심사를 생각해보세요.</li>
            <li><strong>가장 중요한 이야기</strong>는 책의 핵심 메시지예요!</li>
        </ul>
        """,
        2: """
        <ul>
            <li>제목은 <strong>10자 이내</strong>로 짧고 강렬하게!</li>
            <li>숫자가 들어가면 더 눈에 띄어요. (예: "7일 만에", "3가지 비밀")</li>
            <li>AI가 추천한 제목을 그대로 쓰거나 마음대로 바꿔도 돼요!</li>
        </ul>
        """,
        3: """
        <ul>
            <li>목차는 <strong>책의 설계도</strong>예요. 잘 만들면 글쓰기가 쉬워져요!</li>
            <li>Part 1~5 구조로 40장이 자동 생성돼요.</li>
            <li>마음에 안 드는 부분은 직접 수정해도 돼요!</li>
        </ul>
        """,
        4: """
        <ul>
            <li>한 장씩 천천히 써보세요. <strong>완벽하지 않아도 괜찮아요!</strong></li>
            <li>'AI로 글 만들기' 버튼 하나로 1,500자가 뚝딱!</li>
            <li>직접 쓰고 싶으면 아래 '직접 쓰기'를 사용하세요.</li>
            <li>급하면 '전체 자동 생성'으로 한번에 완성할 수 있어요!</li>
        </ul>
        """,
        5: """
        <ul>
            <li>책 소개서는 <strong>출판사에 보여주는 문서</strong>예요.</li>
            <li>저자 정보를 자세히 쓰면 출판사가 더 관심을 가져요!</li>
            <li>경력, 성과, SNS 팔로워 수 등을 구체적으로 써보세요.</li>
        </ul>
        """,
        6: """
        <ul>
            <li>홍보 페이지는 <strong>책을 알리는 글</strong>이에요.</li>
            <li>웨비나나 강의 정보를 넣으면 더 효과적이에요!</li>
            <li>보너스 혜택을 넣으면 사람들이 더 관심을 가져요.</li>
        </ul>
        """,
        7: """
        <ul>
            <li>드디어 <strong>책이 완성</strong>됐어요! 🎉</li>
            <li>마크다운(.md) 또는 텍스트(.txt) 형식으로 다운로드할 수 있어요.</li>
            <li>다운로드한 파일을 워드나 한글에 붙여넣기해서 편집하세요!</li>
        </ul>
        """,
    }

    tip_content = tips.get(step, "")
    if tip_content:
        tip_html = f"""
        <div class="ux-step-tip">
            <div class="ux-step-tip-content">
                {tip_content}
            </div>
        </div>
        """
        st.markdown(tip_html, unsafe_allow_html=True)


def render_ux_char_counter(text: str, min_chars: int = 0, max_chars: int = 0, target_chars: int = 0):
    """실시간 글자 수 카운터

    Args:
        text: 현재 텍스트
        min_chars: 최소 글자 수 (0이면 체크 안함)
        max_chars: 최대 글자 수 (0이면 체크 안함)
        target_chars: 목표 글자 수 (0이면 표시 안함)
    """
    char_count = len(text.replace(" ", "").replace("\n", "")) if text else 0

    # 상태 결정
    status_class = ""
    status_text = ""
    bar_color = "#1976D2"
    bar_percent = 0

    if target_chars > 0:
        bar_percent = min(100, (char_count / target_chars) * 100)
        if char_count >= target_chars:
            status_class = "success"
            status_text = "✓ 목표 달성!"
            bar_color = "#4CAF50"
        elif char_count >= target_chars * 0.8:
            status_class = "warning"
            status_text = "거의 다 됐어요!"
            bar_color = "#FF9800"
        else:
            status_text = f"목표: {target_chars:,}자"

    if min_chars > 0 and char_count < min_chars:
        status_class = "danger"
        status_text = f"⚠️ 최소 {min_chars}자 필요해요!"
        bar_color = "#f44336"

    if max_chars > 0 and char_count > max_chars:
        status_class = "danger"
        status_text = f"⚠️ 최대 {max_chars}자까지!"
        bar_color = "#f44336"

    counter_html = f"""
    <div class="ux-char-counter">
        <span class="ux-char-count {status_class}">📝 {char_count:,}자</span>
        <div class="ux-char-bar">
            <div class="ux-char-bar-fill" style="width: {bar_percent}%; background: {bar_color};"></div>
        </div>
        <span>{status_text}</span>
    </div>
    """
    st.markdown(counter_html, unsafe_allow_html=True)


def render_ux_faq_button():
    """눈에 띄는 FAQ 버튼"""
    faq_html = """
    <div class="ux-faq-floating" onclick="document.querySelector('[data-testid=\\'stSidebar\\']').scrollIntoView();">
        ❓ 도움이 필요해요?
    </div>
    """
    st.markdown(faq_html, unsafe_allow_html=True)


def render_ux_step_locked(required_step: int, required_action: str):
    """단계 건너뛰기 방지 경고

    Args:
        required_step: 필요한 이전 단계 번호
        required_action: 필요한 행동 설명
    """
    locked_html = f"""
    <div class="ux-step-locked">
        <div class="ux-step-locked-icon">🔒</div>
        <div class="ux-step-locked-text">앗! 아직 이 단계로 올 수 없어요!</div>
        <div class="ux-step-locked-hint">먼저 {required_step}단계에서 {required_action}을(를) 완료해주세요.</div>
    </div>
    """
    st.markdown(locked_html, unsafe_allow_html=True)


def render_ux_nav_buttons(current_step: int, can_go_next: bool = True, next_label: str = None, prev_label: str = None):
    """향상된 이전/다음 네비게이션 버튼

    Args:
        current_step: 현재 단계
        can_go_next: 다음으로 이동 가능 여부
        next_label: 다음 버튼 레이블 (기본값: 자동 생성)
        prev_label: 이전 버튼 레이블 (기본값: 자동 생성)
    """
    step_names = {
        1: "정보 입력",
        2: "제목 만들기",
        3: "목차 만들기",
        4: "글 쓰기",
        5: "책 소개서",
        6: "홍보 페이지",
        7: "다운로드",
    }

    prev_name = step_names.get(current_step - 1, "")
    next_name = step_names.get(current_step + 1, "")

    if prev_label is None:
        prev_label = f"← {current_step - 1}. {prev_name}" if prev_name else "← 이전"
    if next_label is None:
        next_label = f"{current_step + 1}. {next_name} →" if next_name else "다음 →"

    col1, col2 = st.columns(2)

    with col1:
        if current_step > 1:
            if st.button(prev_label, key=f"ux_nav_prev_{current_step}", use_container_width=True):
                st.session_state.current_step = current_step - 1
                st.rerun()

    with col2:
        if current_step < 7:
            if can_go_next:
                if st.button(next_label, key=f"ux_nav_next_{current_step}", use_container_width=True, type="primary"):
                    st.session_state.current_step = current_step + 1
                    st.rerun()
            else:
                st.button(next_label, key=f"ux_nav_next_{current_step}_disabled", use_container_width=True, disabled=True)
                st.caption("👆 위의 필수 항목을 완료해주세요!")


def inject_auto_scroll_script(element_id: str = "ux-scroll-target"):
    """새 콘텐츠 생성 시 자동 스크롤 스크립트 삽입

    Args:
        element_id: 스크롤 대상 요소 ID
    """
    script = f"""
    <script>
        setTimeout(function() {{
            var target = document.getElementById('{element_id}');
            if (target) {{
                target.scrollIntoView({{ behavior: 'smooth', block: 'start' }});
                target.classList.add('ux-scroll-target');
            }}
        }}, 500);
    </script>
    """
    st.markdown(script, unsafe_allow_html=True)


def render_help_and_contact_sections():
    """도움 챗봇과 연락 섹션 - 모달 팝업 형태로 렌더링"""
    current_step = st.session_state.current_step
    book_info = st.session_state.book_info
    student_name = book_info.get("name", "")

    # 도움 챗봇 모달 (활성화된 경우)
    if st.session_state.get("show_help_chatbot", False):
        # 모달 오버레이 시작
        st.markdown("""
        <div class="modal-overlay" id="help-modal-overlay">
            <div class="modal-container">
                <div class="modal-header">
                    <h3>AI 도움 챗봇</h3>
                </div>
                <div class="modal-body">
        """, unsafe_allow_html=True)

        # 닫기 버튼 (Streamlit 버튼)
        col_close1, col_spacer1, col_title1 = st.columns([1, 6, 1])
        with col_close1:
            if st.button("X 닫기", key="close_help_modal", type="secondary"):
                st.session_state.show_help_chatbot = False
                st.rerun()

        # 도움 챗봇 내용
        render_enhanced_chatbot(current_step, book_info)

        st.markdown("""
                </div>
            </div>
        </div>
        """, unsafe_allow_html=True)

    # 연락하기 모달 (활성화된 경우)
    if st.session_state.get("show_contact_section", False):
        # 모달 오버레이 시작
        st.markdown("""
        <div class="modal-overlay" id="contact-modal-overlay">
            <div class="modal-container">
                <div class="modal-header">
                    <h3>선생님께 연락하기</h3>
                </div>
                <div class="modal-body">
        """, unsafe_allow_html=True)

        # 닫기 버튼 (Streamlit 버튼)
        col_close2, col_spacer2, col_title2 = st.columns([1, 6, 1])
        with col_close2:
            if st.button("X 닫기", key="close_contact_modal", type="secondary"):
                st.session_state.show_contact_section = False
                st.rerun()

        # 연락하기 내용
        render_contact_section(student_name, current_step, book_info)

        st.markdown("""
                </div>
            </div>
        </div>
        """, unsafe_allow_html=True)


def main():
    """메인 함수"""
    init_session_state()

    # 접근성: 스킵 네비게이션 링크
    st.markdown('''
    <a href="#main-content" class="skip-link" tabindex="0">
        본문으로 건너뛰기
    </a>
    <div id="live-announcer" class="live-region" aria-live="polite" aria-atomic="true"></div>
    ''', unsafe_allow_html=True)

    # ===== 이전 작업 복구 체크 (앱 시작 시) =====
    # 복구 화면이 표시되면 다른 UI는 표시하지 않음
    if render_recovery_prompt():
        return

    render_sidebar()

    # 플로팅 챗봇 버튼 CSS 추가
    render_floating_chatbot_button()

    # 메인 콘텐츠 영역 시작 (접근성)
    st.markdown('<main id="main-content" role="main" tabindex="-1">', unsafe_allow_html=True)

    # 음성 모드가 활성화된 경우
    if st.session_state.get("voice_mode_active", False):
        st.markdown('<h1 class="main-header" role="banner">마이크 책쓰기 코칭</h1>', unsafe_allow_html=True)
        st.markdown("음성으로 책 내용을 전달하세요!")
        st.markdown('<hr role="separator" aria-hidden="true">', unsafe_allow_html=True)
        render_voice_mode()
        st.markdown('</main>', unsafe_allow_html=True)
        return

    # 유튜브 모드가 활성화된 경우
    if st.session_state.get("youtube_mode_active", False):
        render_youtube_mode()
        st.markdown('</main>', unsafe_allow_html=True)
        return

    # 채팅 모드가 활성화된 경우
    if st.session_state.get("chat_mode_active", False):
        render_chat_mode()
        st.markdown('</main>', unsafe_allow_html=True)
        return

    # 레이아웃
    if st.session_state.show_chatbot:
        col1, col2 = st.columns([2, 1])
        with col1:
            st.markdown('<h1 class="main-header" role="banner">책쓰기 코칭</h1>', unsafe_allow_html=True)
            st.markdown("AI와 함께 6만자 책을 완성하세요!")
            st.markdown('<hr role="separator" aria-hidden="true">', unsafe_allow_html=True)

            steps = {
                1: render_step1,
                2: render_step2,
                3: render_step3,
                4: render_step4,
                5: render_step5,
                6: render_step6,
                7: render_step7,
            }
            current_step = st.session_state.current_step
            if current_step in steps:
                steps[current_step]()

        with col2:
            render_chatbot()
    else:
        st.markdown('<h1 class="main-header" role="banner">책쓰기 코칭</h1>', unsafe_allow_html=True)
        st.markdown("AI와 함께 6만자 책을 완성하세요!")
        st.markdown('<hr role="separator" aria-hidden="true">', unsafe_allow_html=True)

        steps = {
            1: render_step1,
            2: render_step2,
            3: render_step3,
            4: render_step4,
            5: render_step5,
            6: render_step6,
            7: render_step7,
        }
        current_step = st.session_state.current_step
        if current_step in steps:
            steps[current_step]()

    # 도움 챗봇 및 연락 섹션 (모든 레이아웃에서 공통)
    render_help_and_contact_sections()

    # 메인 콘텐츠 영역 종료 (접근성)
    st.markdown('</main>', unsafe_allow_html=True)


if __name__ == "__main__":
    main()
